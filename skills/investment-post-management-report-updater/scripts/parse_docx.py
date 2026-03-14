#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCX 文件解析脚本
功能：解析 Word 文档（DOCX），提取文本内容和结构信息
"""

import argparse
import json
import sys
from typing import Dict, Any, List
from docx import Document
from datetime import datetime


def parse_docx_file(file_path: str) -> Dict[str, Any]:
    """
    解析 DOCX 文件
    
    参数:
        file_path: DOCX 文件路径
    
    返回:
        包含文档内容的字典
    """
    try:
        doc = Document(file_path)
        
        result = {
            'file_name': file_path.split('/')[-1],
            'parse_time': datetime.now().isoformat(),
            'paragraphs': [],
            'tables': [],
            'styles': {},
            'structure': {
                'total_paragraphs': 0,
                'total_tables': 0,
                'heading_levels': {}
            }
        }
        
        # 解析段落
        for idx, para in enumerate(doc.paragraphs):
            para_info = {
                'index': idx,
                'text': para.text.strip(),
                'style': para.style.name if para.style else 'Normal',
                'level': None
            }
            
            # 识别标题级别
            if para.style and para.style.name.startswith('Heading'):
                try:
                    para_info['level'] = int(para.style.name.split()[-1])
                except (ValueError, IndexError):
                    pass
            
            if para_info['text']:  # 只保留非空段落
                result['paragraphs'].append(para_info)
                
                # 统计标题级别
                if para_info['level']:
                    level_key = f"Level_{para_info['level']}"
                    result['structure']['heading_levels'][level_key] = \
                        result['structure']['heading_levels'].get(level_key, 0) + 1
        
        # 解析表格
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # 只保留非空行
                    table_data.append(row_data)
            
            if table_data:
                result['tables'].append({
                    'index': table_idx,
                    'data': table_data,
                    'rows': len(table_data),
                    'cols': len(table_data[0]) if table_data else 0
                })
        
        # 更新统计信息
        result['structure']['total_paragraphs'] = len(result['paragraphs'])
        result['structure']['total_tables'] = len(result['tables'])
        
        return result
        
    except Exception as e:
        return {
            'error': f'解析失败: {str(e)}',
            'file_path': file_path
        }


def extract_text_by_structure(data: Dict[str, Any], level: int = None) -> List[str]:
    """
    根据结构提取文本
    
    参数:
        data: 解析后的文档数据
        level: 指定标题级别（可选）
    
    返回:
        提取的文本列表
    """
    texts = []
    
    for para in data['paragraphs']:
        if level is None or para.get('level') == level:
            texts.append(para['text'])
    
    return texts


def get_full_text(data: Dict[str, Any]) -> str:
    """
    获取文档的完整文本
    
    参数:
        data: 解析后的文档数据
    
    返回:
        完整文本字符串
    """
    paragraphs = [para['text'] for para in data['paragraphs']]
    return '\n\n'.join(paragraphs)


def main():
    parser = argparse.ArgumentParser(description='解析 DOCX 文件')
    parser.add_argument('--file', required=True, help='DOCX 文件路径')
    parser.add_argument('--output', help='输出JSON文件路径（可选）')
    parser.add_argument('--text-only', action='store_true', help='仅输出纯文本内容')
    
    args = parser.parse_args()
    
    # 解析文件
    result = parse_docx_file(args.file)
    
    if 'error' in result:
        print(json.dumps({'error': result['error']}, ensure_ascii=False, indent=2), file=sys.stderr)
        sys.exit(1)
    
    # 如果只要纯文本
    if args.text_only:
        full_text = get_full_text(result)
        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                f.write(full_text)
            print(f'纯文本已保存到: {args.output}')
        else:
            print(full_text)
    else:
        # 输出完整 JSON
        output_json = json.dumps(result, ensure_ascii=False, indent=2)
        if args.output:
            with open(args.output, 'w', encoding='utf-8') as f:
                f.write(output_json)
            print(f'结果已保存到: {args.output}')
        else:
            print(output_json)


if __name__ == '__main__':
    main()
