import sys
from docx import Document

def extract_docx_text(docx_path):
    try:
        doc = Document(docx_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
        return text
    except Exception as e:
        return f"Error extracting DOCX: {str(e)}"

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python extract_docx.py <input-docx-path>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    text = extract_docx_text(docx_path)
    print(text)
