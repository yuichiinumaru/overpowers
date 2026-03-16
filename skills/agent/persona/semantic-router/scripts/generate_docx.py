#!/usr/bin/env python3
"""
Generate DOCX README for semantic-router v2.0
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_readme():
    doc = Document()
    
    # Title
    title = doc.add_heading('Semantic Router', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Version and URL
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('v2.0.0 â€¢ https://clawhub.ai/halfmoon82/semantic-router')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    # Overview
    doc.add_heading('Overview', 1)
    doc.add_paragraph(
        'Semantic Router is an intelligent routing system for OpenClaw that automatically '
        'directs tasks to appropriate model pools based on task type and semantic analysis.'
    )
    
    # What's New in v2.0
    doc.add_heading('ğŸ†• What\'s New in v2.0', 1)
    doc.add_paragraph('Interactive Setup Wizard')
    
    wizard_features = [
        'Step 0: Define your task types',
        'Step 1: Scan your available models',
        'Step 2: Get smart pool recommendations',
        'Step 3: Customize and confirm your setup'
    ]
    for feature in wizard_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Quick Start
    doc.add_heading('Quick Start', 1)
    
    doc.add_heading('Run Setup Wizard (Recommended)', 2)
    doc.add_paragraph('python3 scripts/setup_wizard.py')
    
    doc.add_heading('Run Semantic Check', 2)
    doc.add_paragraph('python3 scripts/semantic_check.py "your message" "current_pool"')
    
    # Three-Pool Architecture
    doc.add_heading('Three-Pool Architecture', 1)
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header
    hdr = table.rows[0].cells
    hdr[0].text = 'Pool'
    hdr[1].text = 'Task Types'
    hdr[2].text = 'Primary Model'
    hdr[3].text = 'Use Case'
    
    # Data
    pools = [
        ('Intelligence', 'Development, Automation, System Ops', 'Code-focused models', 'Complex coding tasks'),
        ('Highspeed', 'Info Retrieval, Search, Coordination', 'Fast/light models', 'Quick queries'),
        ('Humanities', 'Content, Training, Multimodal', 'Balanced models', 'Creative tasks')
    ]
    
    for i, (pool, tasks, model, use) in enumerate(pools, 1):
        row = table.rows[i].cells
        row[0].text = pool
        row[1].text = tasks
        row[2].text = model
        row[3].text = use
    
    doc.add_paragraph()
    
    # Priority Matrix
    doc.add_heading('Priority Matrix', 1)
    
    prio_table = doc.add_table(rows=6, cols=3)
    prio_table.style = 'Light Grid Accent 1'
    
    hdr = prio_table.rows[0].cells
    hdr[0].text = 'Priority'
    hdr[1].text = 'Type'
    hdr[2].text = 'Keywords'
    
    priorities = [
        ('P0', 'Continue', 'ç»§ç»­ã€æ¥ç€ã€åˆšæ‰ã€ä¸‹ä¸€æ­¥'),
        ('P1', 'Development', 'å¼€å‘ã€å†™ä»£ç ã€è°ƒè¯•ã€éƒ¨ç½²'),
        ('P2', 'Query', 'æŸ¥ä¸€ä¸‹ã€æœç´¢ã€æ‰¾ã€å¤©æ°”'),
        ('P3', 'Content', 'å†™æ–‡ç« ã€æ€»ç»“ã€è§£é‡Š'),
        ('P4', 'New Session', 'hiã€åœ¨å—ã€hello')
    ]
    
    for i, (p, t, k) in enumerate(priorities, 1):
        row = prio_table.rows[i].cells
        row[0].text = p
        row[1].text = t
        row[2].text = k
    
    doc.add_paragraph()
    
    # Fallback Chain
    doc.add_heading('Fallback Chain', 1)
    doc.add_paragraph(
        'When primary model fails (429/Timeout/Error), the system automatically '
        'attempts fallback models in order:'
    )
    
    fallback_flow = [
        'Primary Model fails',
        'â†’ Fallback 1 (same pool)',
        'â†’ Fallback 2 (cross-pool)',
        'â†’ All failed â†’ Pause â†’ Report'
    ]
    for step in fallback_flow:
        doc.add_paragraph(step, style='List Bullet')
    
    # Configuration
    doc.add_heading('Configuration Files', 1)
    
    config_files = [
        ('config/pools.json', 'Model pool definitions'),
        ('config/tasks.json', 'Task type keywords'),
        ('scripts/semantic_check.py', 'Core routing script'),
        ('scripts/setup_wizard.py', 'Interactive configuration')
    ]
    
    for file, desc in config_files:
        p = doc.add_paragraph()
        p.add_run(file).bold = True
        p.add_run(f': {desc}')
    
    # Files Section
    doc.add_heading('Files', 1)
    files_list = [
        'scripts/semantic_check.py - Core script with auto-switch support',
        'scripts/setup_wizard.py - Interactive pool configuration (NEW)',
        'config/pools.json - Model pool configuration',
        'config/tasks.json - Task type keywords',
        'references/flow.md - Detailed flow documentation'
    ]
    for f in files_list:
        doc.add_paragraph(f, style='List Bullet')
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Semantic Router v2.0.0 â€¢ OpenClaw')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save
    output_path = '/Users/macmini/.openclaw/workspace/skills/semantic-router/README_v2.docx'
    doc.save(output_path)
    print(f'DOCX README created: {output_path}')

if __name__ == '__main__':
    create_readme()
