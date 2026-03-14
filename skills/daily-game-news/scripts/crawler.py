#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Daily Game News Crawler V4
每日游戏资讯自动抓取与报告生成
直接使用 curl + BeautifulSoup 抓取各网站列表页
"""

import json
import os
import re
from datetime import datetime, timedelta
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import subprocess

# 配置路径
CONFIG_PATH = Path("/home/admin/.openclaw/workspace/configs/news-crawler-config.json")
REPORTS_DIR = Path("/home/admin/.openclaw/workspace/reports/daily-game-news")

# 确保报告目录存在
REPORTS_DIR.mkdir(parents=True, exist_ok=True)


def load_config():
    """加载配置文件"""
    with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)


def get_date_range():
    """获取 24 小时日期范围"""
    today = datetime.now()
    yesterday = today - timedelta(days=1)
    return yesterday.strftime('%Y%m%d'), today.strftime('%Y%m%d')


def format_publish_time(published_str):
    """格式化发布时间"""
    if not published_str:
        return datetime.now().strftime('今天 %H:%M')
    
    try:
        # 处理相对时间
        if '小时前' in published_str:
            hours = int(re.search(r'(\d+) 小时前', published_str).group(1))
            pub_time = datetime.now() - timedelta(hours=hours)
            return pub_time.strftime('今天 %H:%M') if hours < 24 else pub_time.strftime('昨天 %H:%M')
        elif '分钟前' in published_str:
            return datetime.now().strftime('今天 %H:%M')
        elif '今天' in published_str or 'Today' in published_str:
            return datetime.now().strftime('今天 %H:%M')
        elif '昨天' in published_str or 'Yesterday' in published_str:
            return datetime.now().strftime('昨天 %H:%M')
        
        # 尝试解析绝对时间
        for fmt in ['%Y-%m-%d %H:%M', '%Y-%m-%d', '%m-%d %H:%M', '%b %d, %Y', '%B %d, %Y']:
            try:
                pub_time = datetime.strptime(published_str[:20], fmt)
                now = datetime.now()
                if pub_time.date() == now.date():
                    return pub_time.strftime('今天 %H:%M')
                elif pub_time.date() == (now - timedelta(days=1)).date():
                    return pub_time.strftime('昨天 %H:%M')
                else:
                    return pub_time.strftime('%m-%d %H:%M')
            except:
                continue
        
        return datetime.now().strftime('今天 %H:%M')
    except:
        return datetime.now().strftime('今天 %H:%M')


def fetch_html(url, timeout=60, use_web_fetch=False):
    """使用 curl 或 web_fetch 抓取 HTML"""
    try:
        # 游民星空和 GameSpot 使用 web_fetch（绕过 Cloudflare）
        if use_web_fetch:
            print(f"    使用 web_fetch 抓取（超时={timeout}s）...")
            # 直接调用 web_fetch 工具（通过 Python API）
            import httpx
            response = httpx.get(
                f'https://r.jina.ai/{url}',
                headers={'X-Return-Format': 'markdown'},
                timeout=timeout
            )
            if response.status_code == 200:
                content = response.text
                print(f"    ✓ web_fetch 成功，返回 {len(content)} 字节")
                return content
            print(f"    ⚠️ web_fetch 失败：{response.status_code}")
            return None
        
        # 其他网站使用 curl
        result = subprocess.run(
            ['curl', '-s', '-A', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
             '--max-time', str(timeout), '-L', '--compressed', url],
            capture_output=True, text=True, timeout=timeout + 5
        )
        if result.returncode == 0:
            return result.stdout
    except Exception as e:
        print(f"    ⚠️ 抓取失败：{e}")
    return None


def parse_gcores(html):
    """解析机核 GCORES"""
    articles = []
    soup = BeautifulSoup(html, 'lxml')
    
    # 机核的文章卡片 - 使用正确的选择器
    for item in soup.select('a.news')[:30]:
        href = item.get('href', '')
        title_elem = item.select_one('h3')
        title = title_elem.get_text(strip=True) if title_elem else item.get('title', '')
        time_elem = item.select_one('.news_meta span')
        pub_time = time_elem.get_text(strip=True) if time_elem else ''
        
        if href and title and len(title) > 5:
            full_url = f'https://www.gcores.com{href}' if href.startswith('/') else href
            articles.append({
                'title': title.strip(),
                'url': full_url,
                'published': format_publish_time(pub_time),
                'content': ''
            })
    
    return articles[:10]


def parse_gamersky(html):
    """解析游民星空 - 使用 web_fetch 提取的 markdown"""
    articles = []
    
    if not html or len(html) < 100:
        print(f"    ⚠️ HTML 内容过短：{len(html) if html else 0} 字节")
        return articles
    
    # jina.ai 返回的格式：
    # *   [标题](URL "title") [![Image...](image)](URL) 摘要 时间
    import re
    # 匹配 [标题](URL) 格式，URL 包含 /news/ 和 .shtml
    pattern = r'\[([^\]]+)\]\((https://www\.gamersky\.com/news/[^)]+\.shtml)'
    
    matches = re.findall(pattern, html)
    print(f"    正则匹配到 {len(matches)} 篇新闻文章")
    
    count = 0
    seen_urls = set()
    for title, url in matches:
        # 去重
        if url in seen_urls:
            continue
        seen_urls.add(url)
        
        if title and len(title) > 5:
            # 尝试从同一行提取时间（查找 URL 后面的日期）
            pub_time = ''
            # 在 html 中查找包含该 URL 的行
            for line in html.split('\n'):
                if url in line:
                    # 从该行提取时间
                    time_match = re.search(r'(\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2})', line)
                    if time_match:
                        pub_time = time_match.group(1)
                    break
            
            articles.append({
                'title': title.strip(),
                'url': url,
                'published': format_publish_time(pub_time),
                'content': ''
            })
            count += 1
            if count >= 10:
                break
    
    print(f"    ✓ 成功解析 {len(articles)} 篇文章")
    return articles[:10]


def parse_chuapp(html):
    """解析触乐"""
    articles = []
    soup = BeautifulSoup(html, 'lxml')
    
    for item in soup.select('article, .post, a[href*="/article/"]')[:30]:
        if item.name == 'a':
            href = item.get('href', '')
            title = item.get('title', item.get_text(strip=True))
        else:
            link = item.select_one('a')
            if not link:
                continue
            href = link.get('href', '')
            title = link.get('title', link.get_text(strip=True))
        
        if href and title and len(title) > 5 and '.html' in href:
            articles.append({
                'title': title.strip(),
                'url': href if href.startswith('http') else f'https://www.chuapp.com{href}',
                'published': '',
                'content': ''
            })
    
    return articles[:10]


def parse_yystv(html):
    """解析游研社"""
    articles = []
    soup = BeautifulSoup(html, 'lxml')
    
    # 游研社的文章结构：<li class="articles-item"><a class="articles-link">
    for item in soup.select('li.articles-item')[:30]:
        link = item.select_one('a.articles-link')
        if not link:
            continue
            
        href = link.get('href', '')
        title_elem = link.select_one('h2.articles-title')
        title = title_elem.get_text(strip=True) if title_elem else link.get('title', '')
        
        if href and title and len(title) > 5:
            # 提取发布时间
            time_elem = item.select_one('.article-meta time')
            pub_time = time_elem.get_text(strip=True) if time_elem else ''
            
            articles.append({
                'title': title.strip(),
                'url': f'https://www.yystv.cn{href}' if href.startswith('/') else href,
                'published': format_publish_time(pub_time),
                'content': ''
            })
    
    return articles[:10]


def parse_ign(html):
    """解析 IGN"""
    articles = []
    soup = BeautifulSoup(html, 'lxml')
    
    for item in soup.select('article a, [class*="article"] a, a[href*="/articles/"], a[href*="/news/"]')[:30]:
        href = item.get('href', '')
        title = item.get('title', item.get_text(strip=True))
        
        if href and title and len(title) > 5:
            if '/articles/' in href or '/news/' in href:
                if href.startswith('/'):
                    href = f'https://www.ign.com{href}'
                articles.append({
                    'title': clean_title(title.strip()),
                    'url': href,
                    'published': '',
                    'content': ''
                })
    
    return articles[:10]


def parse_gamespot(html):
    """解析 GameSpot - 使用 jina.ai 提取的 markdown"""
    articles = []
    
    # jina.ai 返回的格式包含文章链接和标题
    # 查找 gamespot.com/articles 或 gamespot.com/news 的链接
    import re
    pattern = r'-\s+https://www\.gamespot\.com/(articles|news)/[^\s]+\s*\n\s*\n?\s*([^\n]+?)(?=\n\s*-|\Z)'
    
    matches = re.findall(pattern, html, re.DOTALL)
    for url_path, title in matches[:15]:
        if title and len(title) > 5:
            # 提取完整 URL
            url_match = re.search(r'-\s+(https://www\.gamespot\.com/[^\s]+)', html)
            if url_match:
                full_url = url_match.group(1)
                articles.append({
                    'title': clean_title(title.strip()),
                    'url': full_url,
                    'published': '',
                    'content': ''
                })
    
    # 备用方案：查找所有 gamespot 链接
    if not articles:
        for line in html.split('\n'):
            if 'gamespot.com/articles' in line or 'gamespot.com/news' in line:
                url_match = re.search(r'(https://www\.gamespot\.com/[^\s]+)', line)
                if url_match:
                    url = url_match.group(1)
                    # 尝试从同一行或下一行提取标题
                    title = line.replace(url, '').strip('- ').strip()
                    if title and len(title) > 5:
                        articles.append({
                            'title': clean_title(title),
                            'url': url,
                            'published': '',
                            'content': ''
                        })
    
    return articles[:10]


def parse_gamedeveloper(html):
    """解析 GameDeveloper"""
    articles = []
    soup = BeautifulSoup(html, 'lxml')
    
    # GameDeveloper 的文章结构：<div class="ListPreview-TitleWrapper">
    for item in soup.select('.ListPreview-TitleWrapper')[:30]:
        link = item.select_one('a[href]')
        if not link:
            continue
            
        href = link.get('href', '')
        title = link.get_text(strip=True)
        
        if href and title and len(title) > 5:
            # 查找父级的时间信息
            parent = item.find_parent('li') or item.find_parent('div')
            time_elem = parent.select_one('.CardTime, time') if parent else None
            pub_time = time_elem.get_text(strip=True) if time_elem else ''
            
            articles.append({
                'title': clean_title(title.strip()),
                'url': href if href.startswith('http') else f'https://www.gamedeveloper.com{href}',
                'published': format_publish_time(pub_time),
                'content': ''
            })
    
    return articles[:10]


def parse_gamelook(html):
    """解析 GameLook"""
    articles = []
    soup = BeautifulSoup(html, 'lxml')
    
    # GameLook 的文章结构：<li><a href="..." title="...">
    for item in soup.select('li a, .news-list a')[:30]:
        href = item.get('href', '')
        title = item.get('title', item.get_text(strip=True))
        
        if href and title and len(title) > 5 and '/2026' in href:
            # 提取时间
            parent = item.find_parent('li')
            time_elem = parent.select_one('.date, time') if parent else None
            pub_time = time_elem.get_text(strip=True) if time_elem else ''
            
            articles.append({
                'title': title.strip(),
                'url': href if href.startswith('http') else f'http://www.gamelook.com.cn{href}',
                'published': format_publish_time(pub_time),
                'content': ''
            })
    
    return articles[:10]


def parse_youxituoluo(html):
    """解析游戏陀螺"""
    articles = []
    soup = BeautifulSoup(html, 'lxml')
    
    # 游戏陀螺的文章结构：通过 data-url 属性
    for item in soup.select('p.desc.text_justify')[:30]:
        href = item.get('data-url', '')
        title = item.get_text(strip=True)
        
        if href and title and len(title) > 5:
            # 查找相邻的时间信息
            parent = item.find_parent('li') or item.find_parent('div')
            time_elem = parent.select_one('.status, time') if parent else None
            pub_time = time_elem.get_text(strip=True) if time_elem else ''
            
            articles.append({
                'title': title.strip(),
                'url': f'https://www.youxituoluo.com{href}' if href.startswith('/') else href,
                'published': format_publish_time(pub_time),
                'content': ''
            })
    
    return articles[:10]


def clean_title(title):
    """清理标题中的脏数据"""
    # 移除常见的前缀/后缀模式
    patterns = [
        r'^\d+d\s*',  # 移除 "1d ", "2d " 等
        r'^\d+h\s*',  # 移除 "1h ", "2h " 等
        r'\s*-\s*IGN$',
        r'\s*\|\s*IGN$',
        r'\s*-\s*GameSpot$',
        r'\s*_\s*GameSpot$',
    ]
    
    for pattern in patterns:
        title = re.sub(pattern, '', title, flags=re.IGNORECASE)
    
    return title.strip()


def fetch_website(config, date_start, date_end):
    """抓取单个网站"""
    site_id = config['id']
    site_name = config['name']
    base_url = config['base_url']
    
    articles = []
    
    for section in config.get('分区', []):
        limit = section.get('筛选数量', 2)
        section_name = section.get('name', '')
        
        print(f"  - 抓取 {site_name} - {section_name}...")
        
        # 根据网站 ID 选择解析器
        parsers = {
            'gcores': parse_gcores,
            'gamersky': parse_gamersky,
            'chuapp': parse_chuapp,
            'yystv': parse_yystv,
            'ign': parse_ign,
            'gamespot': parse_gamespot,
            'gamedeveloper': parse_gamedeveloper,
            'gamelook': parse_gamelook,
            'youxituoluo': parse_youxituoluo,
        }
        
        parser = parsers.get(site_id)
        if not parser:
            print(f"    ⚠️ 未知网站类型，使用通用解析")
            continue
        
        # 抓取 HTML（游民星空和 GameSpot 使用 jina.ai 绕过防护）
        use_web_fetch = (site_id in ['gamersky', 'gamespot'])
        html = fetch_html(base_url, use_web_fetch=use_web_fetch)
        if not html:
            print(f"    ⚠️ HTML 抓取失败")
            continue
        
        # 解析文章
        articles = parser(html)
        print(f"    ✓ 解析到 {len(articles)} 篇文章")
        
        # 限制数量
        articles = articles[:limit]
    
    return articles


def classify_article(title, content, config):
    """分类文章"""
    text = (title + ' ' + content).lower()
    rules = config.get('分类规则', {}).get('规则', [])
    
    for rule in rules:
        condition = rule.get('条件', '')
        category = rule.get('分类', '其他')
        
        if ' OR ' in condition:
            for part in condition.split(' OR '):
                part = part.strip().strip("'\"")
                if '|' in part:
                    if any(kw.lower() in text for kw in part.split('|')):
                        return category
                elif part.lower() in text:
                    return category
        elif ' AND ' in condition:
            if all(
                any(kw.lower() in text for kw in part.split('|')) if '|' in part
                else part.lower().strip("'\"") in text
                for part in condition.split(' AND ')
            ):
                return category
        else:
            if '|' in condition:
                if any(kw.lower().strip("'\"") in text for kw in condition.split('|')):
                    return category
            elif condition.lower().strip("'\"") in text:
                return category
    
    return '其他'


def generate_markdown_report(articles_by_category, date_str, config):
    """生成 Markdown 报告"""
    category_emoji = {
        '头条要闻': '🔥',
        '新品动态': '🎮',
        '厂商动态': '🏢',
        '行业数据': '📊',
        '值得关注': '⭐',
        '投融资': '💰',
        '其他': '📁'
    }
    
    category_order = ['头条要闻', '新品动态', '厂商动态', '行业数据', '值得关注', '投融资', '其他']
    
    report = f"""# 📰 每日游戏资讯报告

**报告日期**: {date_str}  
**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  
**数据来源**: {', '.join([site['name'] for site in config.get('网站配置', [])])}

---

"""
    
    total = 0
    for category in category_order:
        articles = articles_by_category.get(category, [])
        if articles:
            total += len(articles)
            emoji = category_emoji.get(category, '📁')
            report += f"## {emoji} {category}\n\n"
            report += "| 发布源 | 时间 | 文章标题 | 链接 |\n"
            report += "|--------|------|----------|------|\n"
            
            for article in articles:
                title_display = article['title'][:40] + '...' if len(article['title']) > 40 else article['title']
                published = article['published'] or '今天'
                report += f"| {article['source']} | {published} | {title_display} | [查看详情]({article['url']}) |\n"
            
            report += "\n---\n\n"
    
    report += f"\n**报告结束** - 共抓取 {total} 篇文章\n"
    return report


def generate_word_report(articles_by_category, date_str, config):
    """生成 Word 报告"""
    category_emoji = {
        '头条要闻': '🔥',
        '新品动态': '🎮',
        '厂商动态': '🏢',
        '行业数据': '📊',
        '值得关注': '⭐',
        '投融资': '💰',
        '其他': '📁'
    }
    
    category_order = ['头条要闻', '新品动态', '厂商动态', '行业数据', '值得关注', '投融资', '其他']
    
    doc = Document()
    title = doc.add_heading('📰 每日游戏资讯报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'报告日期：{date_str}')
    doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'数据来源：{", ".join([site["name"] for site in config.get("网站配置", [])])}')
    doc.add_paragraph()
    
    total = 0
    for category in category_order:
        articles = articles_by_category.get(category, [])
        if articles:
            total += len(articles)
            emoji = category_emoji.get(category, '📁')
            
            heading = doc.add_heading(f'{emoji} {category}', level=1)
            heading.runs[0].bold = True
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            headers = ['发布源', '时间', '文章标题', '链接']
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            
            for article in articles:
                row = table.add_row().cells
                row[0].text = article['source']
                row[1].text = article['published'] or '今天'
                row[2].text = article['title'][:40] + '...' if len(article['title']) > 40 else article['title']
                row[3].text = article['url']
    
    doc.add_paragraph()
    doc.add_paragraph(f'本报告共抓取 {total} 篇文章').bold = True
    
    file_path = REPORTS_DIR / f'daily-game-news-{date_str}.docx'
    doc.save(file_path)
    return file_path


def generate_text_report(articles_by_category, date_str, config):
    """生成文本报告"""
    category_emoji = {
        '头条要闻': '🔥',
        '新品动态': '🎮',
        '厂商动态': '🏢',
        '行业数据': '📊',
        '值得关注': '⭐',
        '投融资': '💰',
        '其他': '📁'
    }
    
    category_order = ['头条要闻', '新品动态', '厂商动态', '行业数据', '值得关注', '投融资', '其他']
    
    lines = ["=" * 60, "📰 每日游戏资讯报告", "=" * 60]
    lines.append(f"报告日期：{date_str}")
    lines.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"数据来源：{', '.join([site['name'] for site in config.get('网站配置', [])])}")
    lines.append("")
    
    total = 0
    for category in category_order:
        articles = articles_by_category.get(category, [])
        if articles:
            total += len(articles)
            emoji = category_emoji.get(category, '📁')
            lines.append(f"\n{emoji} {category}")
            lines.append("-" * 40)
            
            for i, article in enumerate(articles, 1):
                lines.append(f"\n{i}. {article['title']}")
                lines.append(f"   发布源：{article['source']}")
                lines.append(f"   时间：{article['published'] or '今天'}")
                lines.append(f"   链接：{article['url']}")
    
    lines.extend(["\n" + "=" * 60, f"本报告共抓取 {total} 篇文章", "=" * 60])
    
    file_path = REPORTS_DIR / f'daily-game-news-{date_str}.txt'
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(lines))
    
    return file_path


def main():
    """主函数"""
    print("=" * 60)
    print("📰 Daily Game News Crawler V4")
    print("每日游戏资讯自动抓取（直接爬取 + BeautifulSoup）")
    print("=" * 60)
    
    print("\n📋 加载配置文件...")
    config = load_config()
    print(f"  ✓ 配置文件：{CONFIG_PATH}")
    print(f"  ✓ 配置网站数量：{len(config.get('网站配置', []))}")
    
    date_start, date_end = get_date_range()
    date_str = datetime.now().strftime('%Y-%m-%d')
    print(f"\n📅 抓取日期范围：{date_start} - {date_end}")
    print(f"  报告日期：{date_str}")
    
    print("\n🕷️ 开始抓取文章...")
    all_articles = []
    
    for site_config in config.get('网站配置', []):
        site_name = site_config.get('name', 'Unknown')
        priority = site_config.get('抓取优先级', 2)
        
        if priority == 1:
            print(f"\n【{site_name}】")
            articles = fetch_website(site_config, date_start, date_end)
            
            for article in articles:
                category = classify_article(article['title'], article.get('content', ''), config)
                article['category'] = category
                article['source'] = site_name
            
            all_articles.extend(articles)
        else:
            print(f"\n⏭️ 跳过 {site_name} (优先级：{priority})")
    
    print(f"\n✅ 共抓取 {len(all_articles)} 篇文章")
    
    print("\n🏷️ 分类整理...")
    articles_by_category = {}
    for article in all_articles:
        category = article.get('category', '其他')
        if category not in articles_by_category:
            articles_by_category[category] = []
        articles_by_category[category].append(article)
        print(f"  - {article['title'][:30]}... → {category}")
    
    print("\n📝 生成报告...")
    
    markdown_report = generate_markdown_report(articles_by_category, date_str, config)
    print("  ✓ Markdown 报告生成完成")
    
    word_path = generate_word_report(articles_by_category, date_str, config)
    print(f"  ✓ Word 报告已保存：{word_path}")
    
    text_path = generate_text_report(articles_by_category, date_str, config)
    print(f"  ✓ 文本报告已保存：{text_path}")
    
    print("\n" + "=" * 60)
    print("📊 报告预览")
    print("=" * 60)
    print(markdown_report[:2000])
    
    print("\n" + "=" * 60)
    print("✅ 报告生成完成！")
    print("=" * 60)
    print(f"\n📁 Word 报告位置：{word_path}")
    print(f"📁 文本报告位置：{text_path}")
    
    return {
        'date': date_str,
        'total_articles': len(all_articles),
        'categories': articles_by_category,
        'word_path': str(word_path)
    }


if __name__ == '__main__':
    main()
