#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Daily Game News Crawler V2
每日游戏资讯自动抓取与报告生成（混合策略版）

抓取策略（优先级从高到低）：
1. RSS 订阅（稳定、快速、不易被反爬）
2. 直接爬取网站 HTML（需要解析规则）
3. SearXNG 搜索（作为备用方案）
"""

import json
import os
import re
import time
import hashlib
from datetime import datetime, timedelta
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 配置路径
CONFIG_PATH = Path("/home/admin/.openclaw/workspace/configs/news-crawler-config.json")
REPORTS_DIR = Path("/home/admin/.openclaw/workspace/reports/daily-game-news")
CACHE_DIR = Path("/home/admin/.openclaw/workspace/reports/daily-game-news/.cache")

# 确保目录存在
REPORTS_DIR.mkdir(parents=True, exist_ok=True)
CACHE_DIR.mkdir(parents=True, exist_ok=True)

# 缓存配置
CACHE_EXPIRY = 3600  # 缓存过期时间（秒）


def load_config():
    """加载配置文件"""
    with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)


def get_date_range():
    """获取 24 小时日期范围"""
    today = datetime.now()
    yesterday = today - timedelta(days=1)
    return yesterday.strftime('%Y%m%d'), today.strftime('%Y%m%d')


def format_publish_time(published_str):
    """格式化发布时间为具体时分格式
    
    Args:
        published_str: ISO 格式时间字符串或任意格式
        
    Returns:
        格式化后的时间字符串，如 "03-07 13:20" 或 "今天 13:20"
    """
    if not published_str:
        return datetime.now().strftime('今天 %H:%M')
    
    try:
        # 尝试解析 ISO 格式 (2026-03-07T13:20:00Z)
        if 'T' in published_str:
            # 处理时区
            pub_str = published_str.replace('Z', '+00:00')
            # Python 3.6 不支持 +00:00 格式，需要特殊处理
            if '+' in pub_str:
                pub_str = pub_str.split('+')[0]
            elif pub_str.endswith('-00:00'):
                pub_str = pub_str[:-6]
            
            pub_time = datetime.fromisoformat(pub_str)
        else:
            # 尝试其他常见格式
            for fmt in ['%Y-%m-%d %H:%M:%S', '%Y-%m-%d %H:%M', '%Y-%m-%d']:
                try:
                    pub_time = datetime.strptime(published_str[:len(fmt.replace('%','').replace('-','')).replace(' ','')], fmt.replace('%Y','%Y').replace('%m','%m').replace('%d','%d').replace('%H','%H').replace('%M','%M').replace('%S','%S'))
                    break
                except:
                    continue
            else:
                # 所有格式都失败，返回今天
                return datetime.now().strftime('今天 %H:%M')
        
        # 计算时间差
        now = datetime.now()
        diff = now - pub_time
        
        # 如果是今天
        if diff.days == 0 and pub_time.date() == now.date():
            return pub_time.strftime('今天 %H:%M')
        # 如果是昨天
        elif diff.days == 1 or (pub_time.date() == (now - timedelta(days=1)).date()):
            return pub_time.strftime('昨天 %H:%M')
        # 其他日期显示月 - 日 时分
        else:
            return pub_time.strftime('%m-%d %H:%M')
            
    except Exception as e:
        # 解析失败，返回今天
        return datetime.now().strftime('今天 %H:%M')


def get_cache_key(site_id, method):
    """生成缓存键"""
    date_str = datetime.now().strftime('%Y%m%d')
    key_str = f"{site_id}:{method}:{date_str}"
    return hashlib.md5(key_str.encode()).hexdigest()[:16]


def load_from_cache(site_id, method):
    """从缓存加载数据"""
    cache_key = get_cache_key(site_id, method)
    cache_file = CACHE_DIR / f"{cache_key}.json"
    
    if cache_file.exists():
        try:
            with open(cache_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                # 检查缓存是否过期
                if datetime.now().timestamp() - data.get('timestamp', 0) < CACHE_EXPIRY:
                    return data.get('articles', [])
        except:
            pass
    return []


def save_to_cache(site_id, method, articles):
    """保存数据到缓存"""
    cache_key = get_cache_key(site_id, method)
    cache_file = CACHE_DIR / f"{cache_key}.json"
    
    try:
        with open(cache_file, 'w', encoding='utf-8') as f:
            json.dump({
                'timestamp': datetime.now().timestamp(),
                'articles': articles
            }, f, ensure_ascii=False)
    except:
        pass


def fetch_with_rss(site_config):
    """使用 RSS 订阅抓取文章"""
    site_id = site_config['id']
    site_name = site_config['name']
    
    # 定义各网站的 RSS 地址
    rss_feeds = {
        'gcores': 'https://www.gcores.com/rss',
        'gamersky': 'https://rss.gamersky.com/news/',
        'yystv': 'https://www.yystv.cn/feed',
        'chuapp': 'https://www.chuapp.com/?feed=rss2',
        'ign': 'https://www.ign.com/rss',
        'gamelook': 'http://www.gamelook.com.cn/feed',
    }
    
    rss_url = rss_feeds.get(site_id)
    if not rss_url:
        return []
    
    print(f"  📡 尝试 RSS 抓取：{site_name}")
    
    try:
        import subprocess
        # 使用 uv run 执行 RSS 抓取脚本
        script_path = os.path.join(os.path.dirname(__file__), 'rss_fetcher.py')
        
        result = subprocess.run(
            ['uv', 'run', '--with', 'feedparser', script_path, rss_url, '--json'],
            capture_output=True, text=True, timeout=60,
            cwd=os.path.dirname(script_path)
        )
        
        if result.returncode == 0:
            data = json.loads(result.stdout)
            articles = []
            limit = site_config.get('分区', [{}])[0].get('筛选数量', 2)
            
            for item in data.get('entries', [])[:limit * 2]:
                articles.append({
                    'source': site_name,
                    'title': item.get('title', ''),
                    'url': item.get('link', ''),
                    'published': item.get('published', '') or item.get('updated', ''),
                    'content': item.get('summary', '')[:200],
                    'method': 'rss'
                })
            
            print(f"    ✓ RSS 抓取成功：{len(articles)} 篇")
            return articles
        else:
            print(f"  ⚠️ RSS 抓取失败：{result.stderr[:200]}")
    except Exception as e:
        print(f"  ⚠️ RSS 抓取异常：{e}")
    
    return []


def fetch_direct(site_config):
    """直接爬取网站 HTML"""
    site_id = site_config['id']
    site_name = site_config['name']
    base_url = site_config['base_url']
    
    print(f"  🕷️ 尝试直接爬取：{site_name}")
    
    try:
        import subprocess
        script_path = os.path.join(os.path.dirname(__file__), 'direct_fetcher.py')
        
        result = subprocess.run(
            ['uv', 'run', '--with', 'requests', '--with', 'beautifulsoup4', 
             script_path, site_id, base_url],
            capture_output=True, text=True, timeout=60,
            cwd=os.path.dirname(script_path)
        )
        
        if result.returncode == 0:
            data = json.loads(result.stdout)
            articles = []
            limit = site_config.get('分区', [{}])[0].get('筛选数量', 2)
            
            for item in data.get('articles', [])[:limit]:
                articles.append({
                    'source': site_name,
                    'title': item.get('title', ''),
                    'url': item.get('url', ''),
                    'published': item.get('published', ''),
                    'content': item.get('content', '')[:200],
                    'method': 'direct'
                })
            
            print(f"    ✓ 直接爬取成功：{len(articles)} 篇")
            return articles
        else:
            print(f"  ⚠️ 直接爬取失败：{result.stderr[:200]}")
    except Exception as e:
        print(f"  ⚠️ 直接爬取异常：{e}")
    
    return []


def fetch_with_searxng(site, date_start, date_end, limit=10):
    """使用 SearXNG 搜索抓取文章（备用方案）"""
    import subprocess
    
    # 构建搜索查询
    search_query = f'site:{site}'
    
    try:
        script_path = os.path.join(os.path.dirname(__file__), '../../searxng/scripts/searxng.py')
        script_path = os.path.abspath(script_path)
        
        result = subprocess.run(
            ['uv', 'run', script_path,
             'search', search_query, '-n', str(limit), '--format', 'json'],
            capture_output=True, text=True, timeout=60,
            env={**os.environ, 'SEARXNG_URL': 'http://localhost:8080'},
            cwd=os.path.dirname(script_path)
        )
        
        if result.returncode == 0:
            data = json.loads(result.stdout)
            results = data.get('results', [])
            
            # 检查是否有引擎可用
            unresponsive = data.get('unresponsive_engines', [])
            if len(unresponsive) >= 3 and len(results) == 0:
                print(f"  ⚠️ SearXNG 所有引擎都不可用")
                return []
            
            # 手动过滤 24 小时内的文章
            filtered = []
            today = datetime.now().date()
            yesterday = today - timedelta(days=1)
            
            for item in results:
                pub_date = item.get('publishedDate', '')
                if pub_date:
                    try:
                        if 'T' in pub_date:
                            article_date = datetime.fromisoformat(pub_date.replace('Z', '+00:00').split('+')[0]).date()
                        else:
                            article_date = datetime.strptime(pub_date[:10], '%Y-%m-%d').date()
                        
                        if article_date >= yesterday:
                            filtered.append(item)
                    except:
                        filtered.append(item)
                else:
                    filtered.append(item)
            
            articles = []
            for item in filtered[:limit]:
                articles.append({
                    'source': '',  # 后面会填充
                    'title': item.get('title', ''),
                    'url': item.get('url', ''),
                    'published': item.get('publishedDate', ''),
                    'content': item.get('content', '')[:200],
                    'method': 'searxng'
                })
            
            return articles
        else:
            print(f"  ⚠️ SearXNG 错误：{result.stderr[:200]}")
    except subprocess.TimeoutExpired:
        print(f"  ⚠️ SearXNG 超时")
    except Exception as e:
        print(f"  ⚠️ SearXNG 搜索失败：{e}")
    return []


def fetch_website(config, date_start, date_end):
    """根据配置抓取单个网站（混合策略）"""
    site_id = config['id']
    site_name = config['name']
    base_url = config['base_url']
    
    articles = []
    
    # 检查缓存
    cached = load_from_cache(site_id, 'mixed')
    if cached:
        print(f"  ♻️ 使用缓存：{site_name} ({len(cached)} 篇)")
        return cached
    
    # 策略 1: RSS 订阅
    articles = fetch_with_rss(config)
    
    # 策略 2: 直接爬取（如果 RSS 失败）
    if len(articles) == 0:
        articles = fetch_direct(config)
    
    # 策略 3: SearXNG（备用方案）
    if len(articles) == 0:
        # 从配置文件获取搜索域名
        search_domains = {
            'gcores': 'gcores.com/articles',
            'gamersky': 'gamersky.com/news',
            'youxituoluo': 'youxituoluo.com',
            'chuapp': 'chuapp.com',
            'yystv': 'yystv.cn/p',
            'gamelook': 'gamelook.com.cn',
            'ign': 'ign.com/articles',
            'gamespot': 'gamespot.com/news',
            'gamedeveloper': 'gamedeveloper.com'
        }
        
        search_domain = search_domains.get(site_id, site_id)
        limit = config.get('分区', [{}])[0].get('筛选数量', 2)
        
        print(f"  🔍 使用 SearXNG 备用：{site_name}")
        results = fetch_with_searxng(search_domain, date_start, date_end, limit * 2)
        
        for item in results[:limit]:
            item['source'] = site_name
            articles.append(item)
    
    # 保存到缓存
    if len(articles) > 0:
        save_to_cache(site_id, 'mixed', articles)
    
    return articles


def classify_article(title, content, config):
    """根据配置自动分类文章"""
    text = (title + ' ' + content).lower()
    
    # 从配置中读取分类规则
    rules = config.get('分类规则', {}).get('规则', [])
    
    for rule in rules:
        condition = rule.get('条件', '')
        category = rule.get('分类', '其他')
        
        # 简单解析条件（支持 OR 和 AND）
        if ' OR ' in condition:
            parts = condition.split(' OR ')
            for part in parts:
                part = part.strip().strip("'\"")
                if '|' in part:
                    keywords = part.split('|')
                    for kw in keywords:
                        if kw.lower().strip() in text:
                            return category
                elif part.lower() in text:
                    return category
        elif ' AND ' in condition:
            parts = condition.split(' AND ')
            match_all = True
            for part in parts:
                part = part.strip().strip("'\"")
                if '|' in part:
                    keywords = part.split('|')
                    if not any(kw.lower().strip() in text for kw in keywords):
                        match_all = False
                        break
                elif part.lower() not in text:
                    match_all = False
                    break
            if match_all:
                return category
        else:
            if '|' in condition:
                keywords = condition.split('|')
                for kw in keywords:
                    if kw.lower().strip().strip("'\"") in text:
                        return category
            elif condition.lower().strip("'\"") in text:
                return category
    
    return '其他'


def generate_markdown_report(articles_by_category, date_str, config):
    """生成 Markdown 格式报告"""
    category_emoji = {
        '头条要闻': '🔥',
        '新品动态': '🎮',
        '厂商动态': '🏢',
        '行业数据': '📊',
        '值得关注': '⭐',
        '投融资': '💰',
        '其他': '📁'
    }
    
    category_order = ['头条要闻', '新品动态', '厂商动态', '行业数据', '值得关注', '投融资', '其他']
    
    report = f"""# 📰 每日游戏资讯报告

**报告日期**: {date_str}  
**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  
**数据来源**: {', '.join([site['name'] for site in config.get('网站配置', [])])}

---

"""
    
    total_articles = 0
    
    for category in category_order:
        articles = articles_by_category.get(category, [])
        if articles:
            total_articles += len(articles)
            emoji = category_emoji.get(category, '📁')
            report += f"## {emoji} {category}\n\n"
            report += "| 发布源 | 时间 | 文章标题 | 链接 |\n"
            report += "|--------|------|----------|------|\n"
            
            for article in articles:
                title = article['title'][:40] + '...' if len(article['title']) > 40 else article['title']
                published = article['published'] or '今天'
                report += f"| {article['source']} | {published} | {title} | [查看详情]({article['url']}) |\n"
            
            report += "\n---\n\n"
    
    report += f"\n**报告结束** - 共抓取 {total_articles} 篇文章\n"
    
    return report


def generate_word_report(articles_by_category, date_str, config):
    """生成 Word 格式报告 (.docx)"""
    category_emoji = {
        '头条要闻': '🔥',
        '新品动态': '🎮',
        '厂商动态': '🏢',
        '行业数据': '📊',
        '值得关注': '⭐',
        '投融资': '💰',
        '其他': '📁'
    }
    
    category_order = ['头条要闻', '新品动态', '厂商动态', '行业数据', '值得关注', '投融资', '其他']
    
    doc = Document()
    
    title = doc.add_heading('📰 每日游戏资讯报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'报告日期：{date_str}')
    doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'数据来源：{", ".join([site["name"] for site in config.get("网站配置", [])])}')
    doc.add_paragraph()
    
    total_articles = 0
    
    for category in category_order:
        articles = articles_by_category.get(category, [])
        if articles:
            total_articles += len(articles)
            emoji = category_emoji.get(category, '📁')
            
            heading = doc.add_heading(f'{emoji} {category}', level=1)
            heading.runs[0].bold = True
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            header_cells = table.rows[0].cells
            headers = ['发布源', '时间', '文章标题', '链接']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            for article in articles:
                row = table.add_row().cells
                row[0].text = article['source']
                published = article['published'] or '今天'
                row[1].text = published
                title_text = article['title'][:40] + '...' if len(article['title']) > 40 else article['title']
                row[2].text = title_text
                row[3].text = article['url']
            
            doc.add_paragraph()
    
    doc.add_paragraph()
    doc.add_paragraph(f'本报告共抓取 {total_articles} 篇文章').bold = True
    
    file_path = REPORTS_DIR / f'daily-game-news-{date_str}.docx'
    doc.save(file_path)
    
    return file_path


def generate_text_report(articles_by_category, date_str, config):
    """生成文本格式报告"""
    category_emoji = {
        '头条要闻': '🔥',
        '新品动态': '🎮',
        '厂商动态': '🏢',
        '行业数据': '📊',
        '值得关注': '⭐',
        '投融资': '💰',
        '其他': '📁'
    }
    
    category_order = ['头条要闻', '新品动态', '厂商动态', '行业数据', '值得关注', '投融资', '其他']
    
    report_lines = []
    report_lines.append("=" * 60)
    report_lines.append("📰 每日游戏资讯报告")
    report_lines.append("=" * 60)
    report_lines.append(f"报告日期：{date_str}")
    report_lines.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report_lines.append(f"数据来源：{', '.join([site['name'] for site in config.get('网站配置', [])])}")
    report_lines.append("")
    
    total_articles = 0
    
    for category in category_order:
        articles = articles_by_category.get(category, [])
        if articles:
            total_articles += len(articles)
            emoji = category_emoji.get(category, '📁')
            
            report_lines.append(f"\n{emoji} {category}")
            report_lines.append("-" * 40)
            
            for i, article in enumerate(articles, 1):
                report_lines.append(f"\n{i}. {article['title']}")
                report_lines.append(f"   发布源：{article['source']}")
                published = article['published'] or '今天'
                report_lines.append(f"   时间：{published}")
                report_lines.append(f"   链接：{article['url']}")
                if article.get('content'):
                    report_lines.append(f"   摘要：{article['content']}...")
    
    report_lines.append("\n" + "=" * 60)
    report_lines.append(f"本报告共抓取 {total_articles} 篇文章")
    report_lines.append("=" * 60)
    
    report_text = "\n".join(report_lines)
    file_path = REPORTS_DIR / f'daily-game-news-{date_str}.txt'
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(report_text)
    
    return file_path


def main():
    """主函数"""
    print("=" * 60)
    print("📰 Daily Game News Crawler V2")
    print("每日游戏资讯自动抓取（混合策略版）")
    print("=" * 60)
    
    print("\n📋 加载配置文件...")
    config = load_config()
    print(f"  ✓ 配置文件：{CONFIG_PATH}")
    print(f"  ✓ 配置网站数量：{len(config.get('网站配置', []))}")
    
    date_start, date_end = get_date_range()
    date_str = datetime.now().strftime('%Y-%m-%d')
    print(f"\n📅 抓取日期范围：{date_start} - {date_end}")
    print(f"  报告日期：{date_str}")
    
    print("\n🕷️ 开始抓取文章...")
    print("📊 抓取策略：RSS → 直接爬取 → SearXNG（备用）\n")
    
    all_articles = []
    
    for site_config in config.get('网站配置', []):
        site_name = site_config.get('name', 'Unknown')
        priority = site_config.get('抓取优先级', 2)
        
        if priority == 1:
            print(f"\n【{site_name}】")
            articles = fetch_website(site_config, date_start, date_end)
            
            for article in articles:
                category = classify_article(article['title'], article.get('content', ''), config)
                article['category'] = category
            
            all_articles.extend(articles)
            
            # 添加请求间隔，避免被反爬
            time.sleep(1)
        else:
            print(f"\n⏭️ 跳过 {site_name} (优先级：{priority})")
    
    print(f"\n✅ 共抓取 {len(all_articles)} 篇文章")
    
    print("\n🏷️ 分类整理...")
    articles_by_category = {}
    for article in all_articles:
        category = article.get('category', '其他')
        if category not in articles_by_category:
            articles_by_category[category] = []
        articles_by_category[category].append(article)
        print(f"  - {article['title'][:30]}... → {category}")
    
    print("\n📝 生成报告...")
    
    markdown_report = generate_markdown_report(articles_by_category, date_str, config)
    print("  ✓ Markdown 报告生成完成")
    
    word_path = generate_word_report(articles_by_category, date_str, config)
    print(f"  ✓ Word 报告已保存：{word_path}")
    
    text_path = generate_text_report(articles_by_category, date_str, config)
    print(f"  ✓ 文本报告已保存：{text_path}")
    
    print("\n" + "=" * 60)
    print("📊 报告预览")
    print("=" * 60)
    print(markdown_report[:2000])
    
    print("\n" + "=" * 60)
    print("✅ 报告生成完成！")
    print("=" * 60)
    print(f"\n📁 Word 报告位置：{word_path}")
    print(f"📁 文本报告位置：{text_path}")
    print(f"\n💡 获取报告指令:")
    print(f"   read {word_path}")
    
    return {
        'date': date_str,
        'total_articles': len(all_articles),
        'categories': articles_by_category,
        'word_path': str(word_path)
    }


if __name__ == '__main__':
    main()
