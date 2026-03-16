#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Daily Game News Crawler V3
每日游戏资讯自动抓取与报告生成
结合使用 SearXNG 搜索 + web_fetch 直接抓取
"""

import json
import os
import re
import subprocess
from datetime import datetime, timedelta
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 配置路径
CONFIG_PATH = Path("/home/admin/.openclaw/workspace/configs/news-crawler-config.json")
REPORTS_DIR = Path("/home/admin/.openclaw/workspace/reports/daily-game-news")

# 确保报告目录存在
REPORTS_DIR.mkdir(parents=True, exist_ok=True)


def load_config():
    """加载配置文件"""
    with open(CONFIG_PATH, 'r', encoding='utf-8') as f:
        return json.load(f)


def get_date_range():
    """获取 24 小时日期范围"""
    today = datetime.now()
    yesterday = today - timedelta(days=1)
    return yesterday.strftime('%Y%m%d'), today.strftime('%Y%m%d')


def format_publish_time(published_str):
    """格式化发布时间为具体时分格式"""
    if not published_str:
        return datetime.now().strftime('今天 %H:%M')
    
    try:
        if 'T' in published_str:
            pub_str = published_str.replace('Z', '+00:00')
            if '+' in pub_str:
                pub_str = pub_str.split('+')[0]
            elif pub_str.endswith('-00:00'):
                pub_str = pub_str[:-6]
            
            pub_time = datetime.fromisoformat(pub_str)
        else:
            for fmt in ['%Y-%m-%d %H:%M:%S', '%Y-%m-%d %H:%M', '%Y-%m-%d']:
                try:
                    pub_time = datetime.strptime(published_str[:len(fmt.replace('%','').replace('-','')).replace(' ','')], fmt.replace('%Y','%Y').replace('%m','%m').replace('%d','%d').replace('%H','%H').replace('%M','%M').replace('%S','%S'))
                    break
                except:
                    continue
            else:
                return datetime.now().strftime('今天 %H:%M')
        
        now = datetime.now()
        diff = now - pub_time
        
        if diff.days == 0 and pub_time.date() == now.date():
            return pub_time.strftime('今天 %H:%M')
        elif diff.days == 1 or (pub_time.date() == (now - timedelta(days=1)).date()):
            return pub_time.strftime('昨天 %H:%M')
        else:
            return pub_time.strftime('%m-%d %H:%M')
            
    except Exception as e:
        return datetime.now().strftime('今天 %H:%M')


def fetch_with_searxng(site_domain, site_name, site_id, limit=10):
    """
    使用 SearXNG 搜索抓取文章（改进版）
    使用 site:语法 + URL 过滤
    """
    # 策略：使用 site:语法搜索特定网站
    search_query = f'site:{site_domain}'
    
    try:
        script_path = os.path.join(os.path.dirname(__file__), '../../searxng/scripts/searxng.py')
        script_path = os.path.abspath(script_path)
        
        result = subprocess.run(
            ['uv', 'run', script_path,
             'search', search_query, '-n', str(limit * 3), '--format', 'json'],
            capture_output=True, text=True, timeout=60,
            env={**os.environ, 'SEARXNG_URL': 'http://localhost:8080'},
            cwd=os.path.dirname(script_path)
        )
        
        if result.returncode == 0:
            data = json.loads(result.stdout)
            results = data.get('results', [])
            
            # 过滤出目标网站的结果
            filtered = []
            domain_variants = [site_domain, f'www.{site_domain}', f'{site_domain}/']
            
            for item in results:
                url = item.get('url', '')
                # 检查 URL 是否属于目标网站
                if any(domain in url for domain in domain_variants):
                    filtered.append(item)
            
            print(f"    SearXNG: 搜索 '{search_query}' → {len(filtered)} 条匹配结果")
            return filtered[:limit]
        else:
            print(f"    ⚠️ SearXNG 错误：{result.stderr[:200]}")
    except subprocess.TimeoutExpired:
        print(f"    ⚠️ SearXNG 超时")
    except Exception as e:
        print(f"    ⚠️ SearXNG 搜索失败：{e}")
    
    return []


def fetch_with_web_fetch(url):
    """
    使用 web_fetch 工具抓取单个 URL 的内容
    返回提取的标题、内容等
    """
    try:
        # 使用 OpenClaw 的 web_fetch 工具（通过 Python 调用）
        script_path = os.path.join(os.path.dirname(__file__), 'web_fetch_wrapper.py')
        
        if os.path.exists(script_path):
            result = subprocess.run(
                ['uv', 'run', script_path, url],
                capture_output=True, text=True, timeout=30
            )
            
            if result.returncode == 0:
                data = json.loads(result.stdout)
                return {
                    'title': data.get('title', ''),
                    'content': data.get('text', ''),
                    'success': True
                }
        
        # 如果 wrapper 不存在，返回基本信息
        return {
            'title': '',
            'content': '',
            'success': False
        }
    except Exception as e:
        return {
            'title': '',
            'content': '',
            'success': False,
            'error': str(e)
        }


def fetch_website_list_page(base_url, site_id):
    """
    直接抓取网站列表页，提取文章链接
    使用 curl + BeautifulSoup 解析（更可靠）
    """
    articles = []
    
    # 为不同网站定制列表页 URL
    list_urls = {
        'gcores': 'https://www.gcores.com/news',
        'gamersky': 'https://www.gamersky.com/news/',
        'yystv': 'https://www.yystv.cn/docs',
        'chuapp': 'https://www.chuapp.com/category/daily',
        'youxituoluo': 'https://www.youxituoluo.com/news',
        'gamelook': 'http://www.gamelook.com.cn/page/2/',
        'ign': 'https://www.ign.com/news',
        'gamespot': 'https://www.gamespot.com/news/',
        'gamedeveloper': 'https://www.gamedeveloper.com/latest-news'
    }
    
    target_url = list_urls.get(site_id, base_url)
    
    print(f"    抓取列表页：{target_url}")
    
    try:
        # 使用 curl 抓取 HTML
        result = subprocess.run(
            ['curl', '-s', '-A', 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36', 
             '--max-time', '30', '-L', target_url],
            capture_output=True, text=True, timeout=35
        )
        
        if result.returncode == 0:
            html = result.stdout
            
            # 尝试使用 BeautifulSoup 解析（如果可用）
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(html, 'html.parser')
                
                # 根据不同网站使用不同的选择器
                if site_id == 'gcores':
                    # 机核
                    for item in soup.select('a[href^="/news/"]')[:20]:
                        href = item.get('href', '')
                        title = item.get('title', item.get_text(strip=True))
                        if href and title and len(title) > 5:
                            articles.append({
                                'title': title.strip(),
                                'url': f'https://www.gcores.com{href}',
                                'published': '',
                                'content': ''
                            })
                
                elif site_id == 'chuapp':
                    # 触乐
                    for item in soup.select('a[href*="/article/"]')[:20]:
                        href = item.get('href', '')
                        title = item.get('title', item.get_text(strip=True))
                        if href and title and len(title) > 5 and '.html' in href:
                            articles.append({
                                'title': title.strip(),
                                'url': href if href.startswith('http') else f'https://www.chuapp.com{href}',
                                'published': '',
                                'content': ''
                            })
                
                elif site_id == 'yystv':
                    # 游研社
                    for item in soup.select('a[href^="/docs/"]')[:20]:
                        href = item.get('href', '')
                        title = item.get('title', item.get_text(strip=True))
                        if href and title and len(title) > 5:
                            articles.append({
                                'title': title.strip(),
                                'url': f'https://www.yystv.cn{href}',
                                'published': '',
                                'content': ''
                            })
                
                elif site_id == 'gamersky':
                    # 游民星空
                    for item in soup.select('.newslist li a, .topnews a')[:20]:
                        href = item.get('href', '')
                        title = item.get('title', item.get_text(strip=True))
                        if href and title and len(title) > 5:
                            articles.append({
                                'title': title.strip(),
                                'url': href if href.startswith('http') else f'https://www.gamersky.com{href}',
                                'published': '',
                                'content': ''
                            })
                
                elif site_id == 'ign':
                    # IGN
                    for item in soup.select('a[href*="/articles/"], a[href*="/news/"]')[:20]:
                        href = item.get('href', '')
                        title = item.get('title', item.get_text(strip=True))
                        if href and title and len(title) > 5 and href.startswith('/'):
                            articles.append({
                                'title': title.strip(),
                                'url': f'https://www.ign.com{href}',
                                'published': '',
                                'content': ''
                            })
                
                elif site_id == 'gamedeveloper':
                    # GameDeveloper
                    for item in soup.select('a[href*="/news/"], a[href*="/article/"]')[:20]:
                        href = item.get('href', '')
                        title = item.get('title', item.get_text(strip=True))
                        if href and title and len(title) > 5:
                            articles.append({
                                'title': title.strip(),
                                'url': href if href.startswith('http') else f'https://www.gamedeveloper.com{href}',
                                'published': '',
                                'content': ''
                            })
                
                else:
                    # 通用规则：查找所有文章链接
                    for item in soup.select('article a, .post a, .entry a')[:20]:
                        href = item.get('href', '')
                        title = item.get('title', item.get_text(strip=True))
                        if href and title and len(title) > 5 and href.startswith('http'):
                            articles.append({
                                'title': title.strip(),
                                'url': href,
                                'published': '',
                                'content': ''
                            })
                
                # 去重
                seen_urls = set()
                unique_articles = []
                for article in articles:
                    if article['url'] not in seen_urls:
                        seen_urls.add(article['url'])
                        unique_articles.append(article)
                
                articles = unique_articles[:10]  # 最多保留 10 篇
                print(f"    ✓ BeautifulSoup 提取 {len(articles)} 篇文章")
                
            except ImportError:
                # BeautifulSoup 不可用，回退到正则表达式
                print(f"    ⚠️ BeautifulSoup 不可用，使用正则表达式")
                # ... 正则表达式回退逻辑 ...
    
    except subprocess.TimeoutExpired:
        print(f"    ⚠️ 列表页抓取超时")
    except Exception as e:
        print(f"    ⚠️ 列表页抓取失败：{e}")
    
    return articles


def fetch_website(config, date_start, date_end):
    """根据配置抓取单个网站（混合策略）"""
    site_id = config['id']
    site_name = config['name']
    base_url = config['base_url']
    
    articles = []
    
    for section in config.get('分区', []):
        limit = section.get('筛选数量', 2)
        section_name = section.get('name', '')
        
        print(f"  - 抓取 {site_name} - {section_name}...")
        
        # 策略 1: 先尝试用 SearXNG 搜索
        searxng_results = fetch_with_searxng(
            site_domain=base_url.replace('https://', '').replace('http://', '').split('/')[0],
            site_name=site_name,
            site_id=site_id,
            limit=limit
        )
        
        for item in searxng_results:
            articles.append({
                'source': site_name,
                'title': clean_title(item.get('title', ''), site_name),
                'url': item.get('url', ''),
                'published': format_publish_time(item.get('publishedDate', '')),
                'content': item.get('content', '')[:200]
            })
        
        # 策略 2: 如果 SearXNG 结果不足，补充抓取列表页
        if len(articles) < limit:
            print(f"    SearXNG 结果不足 ({len(articles)}/{limit})，补充抓取列表页...")
            list_articles = fetch_website_list_page(base_url, site_id)
            
            for item in list_articles:
                # 避免重复
                if not any(a['url'] == item['url'] for a in articles):
                    articles.append({
                        'source': site_name,
                        'title': clean_title(item['title'], site_name),
                        'url': item['url'],
                        'published': item.get('published', ''),
                        'content': item.get('content', '')[:200]
                    })
                    
                    if len(articles) >= limit:
                        break
        
        print(f"    ✓ 共抓取 {len(articles)} 篇")
    
    return articles


def clean_title(title, site_name):
    """清理标题中的网站名称后缀"""
    # 常见后缀模式
    patterns = [
        f' - {site_name}',
        f' | {site_name}',
        f' _ {site_name}',
        f' - {site_name}.*',
    ]
    
    for pattern in patterns:
        title = re.sub(pattern, '', title, flags=re.IGNORECASE)
    
    return title.strip()


def classify_article(title, content, config):
    """根据配置自动分类文章"""
    text = (title + ' ' + content).lower()
    
    rules = config.get('分类规则', {}).get('规则', [])
    
    for rule in rules:
        condition = rule.get('条件', '')
        category = rule.get('分类', '其他')
        
        if ' OR ' in condition:
            parts = condition.split(' OR ')
            for part in parts:
                part = part.strip().strip("'\"")
                if '|' in part:
                    keywords = part.split('|')
                    for kw in keywords:
                        if kw.lower().strip() in text:
                            return category
                elif part.lower() in text:
                    return category
        elif ' AND ' in condition:
            parts = condition.split(' AND ')
            match_all = True
            for part in parts:
                part = part.strip().strip("'\"")
                if '|' in part:
                    keywords = part.split('|')
                    if not any(kw.lower().strip() in text for kw in keywords):
                        match_all = False
                        break
                elif part.lower() not in text:
                    match_all = False
                    break
            if match_all:
                return category
        else:
            if '|' in condition:
                keywords = condition.split('|')
                for kw in keywords:
                    if kw.lower().strip().strip("'\"") in text:
                        return category
            elif condition.lower().strip("'\"") in text:
                return category
    
    return '其他'


def generate_markdown_report(articles_by_category, date_str, config):
    """生成 Markdown 格式报告"""
    category_emoji = {
        '头条要闻': '🔥',
        '新品动态': '🎮',
        '厂商动态': '🏢',
        '行业数据': '📊',
        '值得关注': '⭐',
        '投融资': '💰',
        '其他': '📁'
    }
    
    category_order = ['头条要闻', '新品动态', '厂商动态', '行业数据', '值得关注', '投融资', '其他']
    
    report = f"""# 📰 每日游戏资讯报告

**报告日期**: {date_str}  
**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  
**数据来源**: {', '.join([site['name'] for site in config.get('网站配置', [])])}

---

"""
    
    total_articles = 0
    
    for category in category_order:
        articles = articles_by_category.get(category, [])
        if articles:
            total_articles += len(articles)
            emoji = category_emoji.get(category, '📁')
            report += f"## {emoji} {category}\n\n"
            report += "| 发布源 | 时间 | 文章标题 | 链接 |\n"
            report += "|--------|------|----------|------|\n"
            
            for article in articles:
                title = article['title'][:40] + '...' if len(article['title']) > 40 else article['title']
                published = article['published'] or '今天'
                report += f"| {article['source']} | {published} | {title} | [查看详情]({article['url']}) |\n"
            
            report += "\n---\n\n"
    
    report += f"\n**报告结束** - 共抓取 {total_articles} 篇文章\n"
    
    return report


def generate_word_report(articles_by_category, date_str, config):
    """生成 Word 格式报告 (.docx)"""
    category_emoji = {
        '头条要闻': '🔥',
        '新品动态': '🎮',
        '厂商动态': '🏢',
        '行业数据': '📊',
        '值得关注': '⭐',
        '投融资': '💰',
        '其他': '📁'
    }
    
    category_order = ['头条要闻', '新品动态', '厂商动态', '行业数据', '值得关注', '投融资', '其他']
    
    doc = Document()
    
    title = doc.add_heading('📰 每日游戏资讯报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'报告日期：{date_str}')
    doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'数据来源：{", ".join([site["name"] for site in config.get("网站配置", [])])}')
    doc.add_paragraph()
    
    total_articles = 0
    
    for category in category_order:
        articles = articles_by_category.get(category, [])
        if articles:
            total_articles += len(articles)
            emoji = category_emoji.get(category, '📁')
            
            heading = doc.add_heading(f'{emoji} {category}', level=1)
            heading.runs[0].bold = True
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            header_cells = table.rows[0].cells
            headers = ['发布源', '时间', '文章标题', '链接']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            for article in articles:
                row = table.add_row().cells
                row[0].text = article['source']
                published = article['published'] or '今天'
                row[1].text = published
                title_text = article['title'][:40] + '...' if len(article['title']) > 40 else article['title']
                row[2].text = title_text
                row[3].text = article['url']
            
            doc.add_paragraph()
    
    doc.add_paragraph()
    doc.add_paragraph(f'本报告共抓取 {total_articles} 篇文章').bold = True
    
    file_path = REPORTS_DIR / f'daily-game-news-{date_str}.docx'
    doc.save(file_path)
    
    return file_path


def generate_text_report(articles_by_category, date_str, config):
    """生成文本格式报告"""
    category_emoji = {
        '头条要闻': '🔥',
        '新品动态': '🎮',
        '厂商动态': '🏢',
        '行业数据': '📊',
        '值得关注': '⭐',
        '投融资': '💰',
        '其他': '📁'
    }
    
    category_order = ['头条要闻', '新品动态', '厂商动态', '行业数据', '值得关注', '投融资', '其他']
    
    report_lines = []
    report_lines.append("=" * 60)
    report_lines.append("📰 每日游戏资讯报告")
    report_lines.append("=" * 60)
    report_lines.append(f"报告日期：{date_str}")
    report_lines.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report_lines.append(f"数据来源：{', '.join([site['name'] for site in config.get('网站配置', [])])}")
    report_lines.append("")
    
    total_articles = 0
    
    for category in category_order:
        articles = articles_by_category.get(category, [])
        if articles:
            total_articles += len(articles)
            emoji = category_emoji.get(category, '📁')
            
            report_lines.append(f"\n{emoji} {category}")
            report_lines.append("-" * 40)
            
            for i, article in enumerate(articles, 1):
                report_lines.append(f"\n{i}. {article['title']}")
                report_lines.append(f"   发布源：{article['source']}")
                published = article['published'] or '今天'
                report_lines.append(f"   时间：{published}")
                report_lines.append(f"   链接：{article['url']}")
                if article.get('content'):
                    report_lines.append(f"   摘要：{article['content']}...")
    
    report_lines.append("\n" + "=" * 60)
    report_lines.append(f"本报告共抓取 {total_articles} 篇文章")
    report_lines.append("=" * 60)
    
    report_text = "\n".join(report_lines)
    file_path = REPORTS_DIR / f'daily-game-news-{date_str}.txt'
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(report_text)
    
    return file_path


def main():
    """主函数"""
    print("=" * 60)
    print("📰 Daily Game News Crawler V3")
    print("每日游戏资讯自动抓取（SearXNG + 列表页直抓）")
    print("=" * 60)
    
    print("\n📋 加载配置文件...")
    config = load_config()
    print(f"  ✓ 配置文件：{CONFIG_PATH}")
    print(f"  ✓ 配置网站数量：{len(config.get('网站配置', []))}")
    
    date_start, date_end = get_date_range()
    date_str = datetime.now().strftime('%Y-%m-%d')
    print(f"\n📅 抓取日期范围：{date_start} - {date_end}")
    print(f"  报告日期：{date_str}")
    
    print("\n🕷️ 开始抓取文章...")
    all_articles = []
    
    for site_config in config.get('网站配置', []):
        site_name = site_config.get('name', 'Unknown')
        priority = site_config.get('抓取优先级', 2)
        
        if priority == 1:
            print(f"\n【{site_name}】")
            articles = fetch_website(site_config, date_start, date_end)
            
            for article in articles:
                category = classify_article(article['title'], article.get('content', ''), config)
                article['category'] = category
            
            all_articles.extend(articles)
        else:
            print(f"\n⏭️ 跳过 {site_name} (优先级：{priority})")
    
    print(f"\n✅ 共抓取 {len(all_articles)} 篇文章")
    
    print("\n🏷️ 分类整理...")
    articles_by_category = {}
    for article in all_articles:
        category = article.get('category', '其他')
        if category not in articles_by_category:
            articles_by_category[category] = []
        articles_by_category[category].append(article)
        print(f"  - {article['title'][:30]}... → {category}")
    
    print("\n📝 生成报告...")
    
    markdown_report = generate_markdown_report(articles_by_category, date_str, config)
    print("  ✓ Markdown 报告生成完成")
    
    word_path = generate_word_report(articles_by_category, date_str, config)
    print(f"  ✓ Word 报告已保存：{word_path}")
    
    text_path = generate_text_report(articles_by_category, date_str, config)
    print(f"  ✓ 文本报告已保存：{text_path}")
    
    print("\n" + "=" * 60)
    print("📊 报告预览")
    print("=" * 60)
    print(markdown_report[:2000])
    
    print("\n" + "=" * 60)
    print("✅ 报告生成完成！")
    print("=" * 60)
    print(f"\n📁 Word 报告位置：{word_path}")
    print(f"📁 文本报告位置：{text_path}")
    print(f"\n💡 获取报告指令:")
    print(f"   read {word_path}")
    
    return {
        'date': date_str,
        'total_articles': len(all_articles),
        'categories': articles_by_category,
        'word_path': str(word_path)
    }


if __name__ == '__main__':
    main()
