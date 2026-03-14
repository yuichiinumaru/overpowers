#!/usr/bin/env python3
"""
Word 文档读取器
支持 .docx 和 .doc 格式的 Word 文档解析
"""

import argparse
import json
import os
import sys
import re
import traceback
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    
try:
    import subprocess
    SUBPROCESS_AVAILABLE = True
except ImportError:
    SUBPROCESS_AVAILABLE = False

class WordReader:
    """Word 文档读取器"""
    
    def __init__(self, file_path):
        self.file_path = Path(file_path)
        self.document = None
        self.format_type = None
        self.encoding = 'utf-8'
        
        # 检查文件是否存在
        if not self.file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 检查文件扩展名
        if self.file_path.suffix.lower() not in ['.docx', '.doc']:
            raise ValueError(f"不支持的文件格式: {self.file_path.suffix}")
    
    def read_docx(self):
        """读取 .docx 格式文档"""
        if not DOCX_AVAILABLE:
            raise Exception("缺少 python-docx 库。请安装：pip3 install python-docx")
        
        try:
            self.document = Document(str(self.file_path))
            self.format_type = 'docx'
            return True
        except Exception as e:
            raise Exception(f"读取 .docx 文件失败: {str(e)}")
    
    def read_doc(self):
        """读取 .doc 格式文档（使用 antiword）"""
        if not SUBPROCESS_AVAILABLE:
            raise Exception("缺少 subprocess 模块")
        
        try:
            # 检查 antiword 是否可用
            result = subprocess.run(['which', 'antiword'], 
                                  capture_output=True, text=True)
            if result.returncode != 0:
                raise Exception("antiword 未安装。请安装 antiword: Ubuntu/Debian: sudo apt-get install antiword; macOS: brew install antiword")
            
            # 使用 antiword 转换
            result = subprocess.run(['antiword', str(self.file_path)], 
                                  capture_output=True, text=True, encoding='utf-8')
            
            if result.returncode != 0:
                raise Exception(f"antiword 转换失败: {result.stderr}")
            
            # 创建临时文档对象
            class TempDocument:
                def __init__(self, text):
                    self.text = text
                    self.paragraphs = [TempParagraph(p) for p in text.split('\n') if p.strip()]
            
            class TempParagraph:
                def __init__(self, text):
                    self.text = text
            
            self.document = TempDocument(result.stdout)
            self.format_type = 'doc'
            return True
        except Exception as e:
            raise Exception(f"读取 .doc 文件失败: {str(e)}")
    
    def read_metadata(self):
        """读取文档元数据"""
        metadata = {
            'filename': self.file_path.name,
            'size': f"{self.file_path.stat().st_size} bytes",
            'created': datetime.fromtimestamp(self.file_path.stat().st_ctime).isoformat(),
            'modified': datetime.fromtimestamp(self.file_path.stat().st_mtime).isoformat()
        }
        
        if self.format_type == 'docx' and hasattr(self.document, 'core_properties'):
            props = self.document.core_properties
            metadata.update({
                'title': getattr(props, 'title', ''),
                'author': getattr(props, 'author', ''),
                'subject': getattr(props, 'subject', ''),
                'keywords': getattr(props, 'keywords', ''),
                'comments': getattr(props, 'comments', ''),
                'application': getattr(props, 'application', ''),
                'category': getattr(props, 'category', '')
            })
        
        return metadata
    
    def extract_text(self):
        """提取文档文本"""
        text_content = []
        
        if self.format_type == 'docx':
            # 提取段落文本
            for para in self.document.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # 提取表格文本
            for table in self.document.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    table_text.append(' | '.join(row_text))
                text_content.append('\n'.join(table_text))
        
        else:  # doc 格式
            text_content = [para.text for para in self.document.paragraphs if para.text.strip()]
        
        return '\n\n'.join(text_content)
    
    def extract_tables(self):
        """提取表格数据"""
        tables = []
        
        if self.format_type == 'docx':
            for i, table in enumerate(self.document.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append({
                    'id': i + 1,
                    'rows': len(table.rows),
                    'columns': len(table.columns) if table.rows else 0,
                    'data': table_data
                })
        
        return tables
    
    def extract_images(self):
        """提取图片信息"""
        images = []
        
        if self.format_type == 'docx':
            try:
                # 获取文档中的关系
                part = self.document.part
                image_parts = part.related_parts
                
                for rel in part.relationships:
                    if rel.reltype == RT.IMAGE:
                        image_data = image_parts[rel.rId]._blob
                        image_info = {
                            'id': rel.rId,
                            'filename': f"image_{rel.rId}.{rel.target_ref.split('.')[-1]}",
                            'size': f"{len(image_data)} bytes"
                        }
                        images.append(image_info)
            except:
                # 图片提取可能失败，忽略错误
                pass
        
        return images
    
    def extract_all(self):
        """提取所有内容"""
        result = {
            'metadata': self.read_metadata(),
            'format': self.format_type,
            'text': self.extract_text(),
            'tables': self.extract_tables(),
            'images': self.extract_images()
        }
        return result
    
    def to_markdown(self, extract_type='all'):
        """转换为 Markdown 格式"""
        if extract_type == 'text':
            return self.extract_text()
        
        result = self.extract_all()
        md_content = []
        
        # 标题
        md_content.append(f"# {result['metadata']['filename']}")
        md_content.append("")
        
        # 元数据
        metadata = result['metadata']
        if metadata.get('title'):
            md_content.append(f"**标题**：{metadata['title']}")
        if metadata.get('author'):
            md_content.append(f"**作者**：{metadata['author']}")
        md_content.append(f"**文件大小**：{metadata['size']}")
        md_content.append(f"**创建时间**：{metadata['created']}")
        md_content.append(f"**修改时间**：{metadata['modified']}")
        md_content.append("")
        
        # 文本内容
        if result['text']:
            md_content.append("## 正文内容")
            md_content.append("")
            md_content.append(result['text'])
            md_content.append("")
        
        # 表格
        if result['tables']:
            md_content.append("## 表格内容")
            md_content.append("")
            for table in result['tables']:
                md_content.append(f"### 表格 {table['id']} ({table['rows']}行 x {table['columns']}列)")
                md_content.append("")
                # 转换为 Markdown 表格
                for row in table['data']:
                    md_row = " | ".join([str(cell) for cell in row])
                    md_content.append(f"| {md_row} |")
                md_content.append("")
        
        # 图片
        if result['images']:
            md_content.append("## 图片列表")
            md_content.append("")
            for img in result['images']:
                md_content.append(f"- **{img['filename']}** ({img['size']})")
            md_content.append("")
        
        return '\n'.join(md_content)
    
    def to_text(self, extract_type='all'):
        """转换为纯文本格式"""
        if extract_type == 'text':
            return self.extract_text()
        
        result = self.extract_all()
        text_content = []
        
        # 标题和元数据
        text_content.append(f"文件：{result['metadata']['filename']}")
        text_content.append("=" * 50)
        text_content.append("")
        
        for key, value in result['metadata'].items():
            if value and key not in ['filename', 'size', 'created', 'modified']:
                text_content.append(f"{key}：{value}")
        
        text_content.append("")
        
        # 文本内容
        if result['text']:
            text_content.append("正文内容：")
            text_content.append("-" * 20)
            text_content.append(result['text'])
            text_content.append("")
        
        # 表格
        if result['tables']:
            text_content.append("表格内容：")
            text_content.append("-" * 20)
            for table in result['tables']:
                text_content.append(f"表格 {table['id']}:")
                for row in table['data']:
                    text_content.append("  " + " | ".join([str(cell) for cell in row]))
                text_content.append("")
        
        return '\n'.join(text_content)

def main():
    parser = argparse.ArgumentParser(description='读取 Word 文档')
    parser.add_argument('path', help='文档路径或目录路径（批量模式）')
    parser.add_argument('--format', choices=['json', 'text', 'markdown'], 
                       default='text', help='输出格式')
    parser.add_argument('--extract', choices=['text', 'tables', 'images', 'metadata', 'all'], 
                       default='all', help='提取内容类型')
    parser.add_argument('--batch', action='store_true', help='批量处理模式')
    parser.add_argument('--output', help='输出文件路径')
    parser.add_argument('--encoding', default='utf-8', help='文本编码')
    
    args = parser.parse_args()
    
    try:
        if args.batch:
            # 批量处理模式
            path = Path(args.path)
            if not path.is_dir():
                print("错误：批量模式需要指定目录路径")
                sys.exit(1)
            
            # 查找所有 Word 文档
            word_files = []
            for ext in ['.docx', '.doc']:
                word_files.extend(path.glob(f"**/*{ext}"))
            
            if not word_files:
                print("未找到 Word 文档")
                sys.exit(0)
            
            print(f"找到 {len(word_files)} 个 Word 文档")
            
            results = {}
            for file_path in word_files:
                print(f"正在处理: {file_path}")
                try:
                    reader = WordReader(file_path)
                    if file_path.suffix.lower() == '.docx':
                        reader.read_docx()
                    else:
                        reader.read_doc()
                    
                    if args.format == 'json':
                        content = reader.extract_all()
                    elif args.format == 'markdown':
                        content = reader.to_markdown(args.extract)
                    else:
                        content = reader.to_text(args.extract)
                    
                    results[str(file_path)] = {
                        'filename': file_path.name,
                        'content': content,
                        'status': 'success'
                    }
                    
                except Exception as e:
                    results[str(file_path)] = {
                        'filename': file_path.name,
                        'error': str(e),
                        'status': 'failed'
                    }
            
            # 保存结果
            if args.output:
                with open(args.output, 'w', encoding='utf-8') as f:
                    json.dump(results, f, ensure_ascii=False, indent=2)
                print(f"结果已保存到: {args.output}")
            else:
                print(json.dumps(results, ensure_ascii=False, indent=2))
        
        else:
            # 单文件处理模式
            reader = WordReader(args.path)
            
            # 根据文件类型读取
            if args.path.lower().endswith('.docx'):
                reader.read_docx()
            else:
                reader.read_doc()
            
            # 根据格式输出
            if args.format == 'json':
                content = reader.extract_all()
            elif args.format == 'markdown':
                content = reader.to_markdown(args.extract)
            else:
                content = reader.to_text(args.extract)
            
            # 输出结果
            if args.output:
                with open(args.output, 'w', encoding=args.encoding) as f:
                    f.write(content)
                print(f"结果已保存到: {args.output}")
            else:
                print(content)
    
    except Exception as e:
        print(f"错误: {str(e)}", file=sys.stderr)
        if '--debug' in sys.argv or '-d' in sys.argv:
            traceback.print_exc()
        sys.exit(1)

if __name__ == '__main__':
    main()