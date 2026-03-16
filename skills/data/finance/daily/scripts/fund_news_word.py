#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基金新闻Word文档生成脚本
用于「过去七天」和「指定日期范围」查询
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime, timedelta
import re


def set_font(run, font_name='等线', font_size=11):
    """设置中文字体"""
    run.font.size = Pt(font_size)
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)


def should_exclude(title, content):
    """
    判断是否应该排除该新闻
    返回 (是否排除, 排除原因)
    """
    text = f"{title} {content}"
    
    # 排除ETF龙虎榜
    if '龙虎榜' in title:
        return True, 'ETF龙虎榜'
    
    # 排除「x只ETF获融资净买入」类
    if 'ETF获融资净买入' in text or re.search(r'\d+只ETF', text):
        return True, 'ETF融资净买入'
    
    # 排除「x月x日资金净流入」类
    if '资金净流入' in text or '资金净流出' in text:
        return True, '资金净流入/流出'
    
    # 排除私募相关（标题含关键词）
    private_keywords = ['私募', '私募基金', '私募机构', '私募产品']
    for kw in private_keywords:
        if kw in title:
            return True, '私募相关'
    
    return False, None


def generate_word_document(news_data, start_date, end_date, output_path):
    """
    生成基金新闻Word文档
    
    参数:
        news_data: 按日期分组的新闻数据
            {日期: [{title, content, url, source, time}, ...]}
        start_date: 起始日期 (datetime对象)
        end_date: 截止日期 (datetime对象)
        output_path: 输出文件路径
    
    返回:
        output_path 或 None(无新闻时)
    """
    source_order = ['证券时报', '中国证券报·中证网', '证券日报']
    
    # 生成日期范围
    date_list = []
    current_date = start_date
    while current_date <= end_date:
        date_list.append(current_date)
        current_date += timedelta(days=1)
    
    # 统计有效新闻数
    total_news = 0
    
    doc = Document()
    is_first_date = True
    
    for date in date_list:
        date_str = f"{date.year}.{date.month}.{date.day}"
        
        # 新日期新起一页
        if not is_first_date:
            doc.add_page_break()
        is_first_date = False
        
        # 日期段落
        date_para = doc.add_paragraph(date_str)
        for run in date_para.runs:
            set_font(run)
        
        doc.add_paragraph()
        
        # 该日期的新闻
        if date_str in news_data:
            for source in source_order:
                source_news = [n for n in news_data[date_str] if n['source'] == source]
                source_news.sort(key=lambda x: x.get('time', '00:00'))
                
                has_news = False
                for news in source_news:
                    exclude, reason = should_exclude(news['title'], news['content'])
                    if exclude:
                        continue
                    
                    has_news = True
                    total_news += 1
                    
                    # 段落1: 新闻正文
                    content_para = doc.add_paragraph(style='List Bullet')
                    
                    title_run = content_para.add_run(news['title'])
                    title_run.bold = True
                    set_font(title_run)
                    
                    sep_run = content_para.add_run('。')
                    sep_run.bold = False
                    set_font(sep_run)
                    
                    content_run = content_para.add_run(news['content'])
                    content_run.bold = False
                    set_font(content_run)
                    
                    # 段落2: 链接
                    link_para = doc.add_paragraph()
                    link_run = link_para.add_run(news['url'])
                    set_font(link_run)
                    
                    # 段落3: 来源
                    source_para = doc.add_paragraph()
                    source_run = source_para.add_run(f"{news['source']}，{news['title']} {news['source']}")
                    set_font(source_run)
                    
                    doc.add_paragraph()
                
                if not has_news and source_news:
                    simple_source = source.replace('·中证网', '')
                    note_para = doc.add_paragraph()
                    note_run = note_para.add_run(f"{date_str}「{simple_source}无符合规则新闻」")
                    set_font(note_run)
        else:
            for source in source_order:
                simple_source = source.replace('·中证网', '')
                note_para = doc.add_paragraph()
                note_run = note_para.add_run(f"{date_str}「{simple_source}无符合规则新闻」")
                set_font(note_run)
    
    if total_news == 0:
        return None
    
    doc.save(output_path)
    return output_path


def get_filename(start_date, end_date):
    """生成文件名"""
    start_str = start_date.strftime('%Y%m%d')
    end_str = end_date.strftime('%Y%m%d')
    return f"{start_str}-{end_str}基金新闻.docx"


if __name__ == '__main__':
    # 测试示例
    from datetime import datetime
    
    news_data = {
        '2026.3.10': [
            {
                'title': 'ETF更名正加快重塑指数化投资底层生态',
                'content': '进入3月份，一场影响全市场ETF产品的"更名变革"正在全面提速。',
                'url': 'http://www.stcn.com/article/detail/3670939.html',
                'source': '证券时报',
                'time': '08:41'
            }
        ]
    }
    
    start_date = datetime(2026, 3, 10)
    end_date = datetime(2026, 3, 10)
    output_path = f'/tmp/{get_filename(start_date, end_date)}'
    
    result = generate_word_document(news_data, start_date, end_date, output_path)
    if result:
        print(f"Word文档已生成：{result}")
    else:
        print("该时间段无有效基金新闻")
