#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
投后管理报告生成脚本
功能：基于模板和分析内容，生成新季度的投后管理报告
"""

import argparse
import json
import sys
import os
from typing import Dict, Any
from docx import Document
from docx.shared import Pt
from datetime import datetime


def generate_report(
    template_path: str,
    financial_data: Dict[str, Any],
    financial_analysis: str,
    business_update: str,
    industry_update: str,
    output_path: str
) -> bool:
    """
    生成新的投后管理报告
    
    参数:
        template_path: 上季度报告模板路径
        financial_data: 财务数据（JSON 字符串或字典）
        financial_analysis: 财务分析文本
        business_update: 经营情况更新文本
        industry_update: 行业分析文本
        output_path: 输出文件路径
    
    返回:
        是否成功
    """
    try:
        # 加载模板文档
        doc = Document(template_path)
        
        # 如果 financial_data 是字符串，解析为字典
        if isinstance(financial_data, str):
            financial_data = json.loads(financial_data)
        
        # 标记已更新过的章节，避免重复更新
        updated_sections = set()
        
        # 遍历所有段落，寻找需要更新的章节
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().lower()
            
            # 更新财务数据章节
            if '财务数据' in text or '财务概况' in text or '财务指标' in text:
                if 'financial' not in updated_sections:
                    # 在财务章节后插入新的财务分析
                    if financial_analysis:
                        new_para = para.insert_paragraph_before(financial_analysis)
                        new_para.style = para.style
                    updated_sections.add('financial')
            
            # 更新公司经营情况章节
            elif '经营情况' in text or '公司经营' in text or '业务进展' in text:
                if 'business' not in updated_sections:
                    if business_update:
                        # 尝试更新该章节的内容
                        if i + 1 < len(doc.paragraphs):
                            next_para = doc.paragraphs[i + 1]
                            # 如果下一段不是标题，替换内容
                            if not next_para.style.name.startswith('Heading'):
                                next_para.text = business_update
                                updated_sections.add('business')
                    updated_sections.add('business')
            
            # 更新行业情况章节
            elif '行业' in text or '竞争格局' in text or '市场环境' in text:
                if 'industry' not in updated_sections:
                    if industry_update:
                        if i + 1 < len(doc.paragraphs):
                            next_para = doc.paragraphs[i + 1]
                            if not next_para.style.name.startswith('Heading'):
                                next_para.text = industry_update
                                updated_sections.add('industry')
                    updated_sections.add('industry')
        
        # 在文档开头添加更新时间戳
        timestamp = datetime.now().strftime('%Y年%m月%d日')
        timestamp_para = doc.paragraphs[0].insert_paragraph_before(
            f'报告更新时间：{timestamp}'
        )
        timestamp_para.style = doc.styles['Normal']
        
        # 保存文档
        doc.save(output_path)
        
        return True
        
    except Exception as e:
        print(f'生成报告失败: {str(e)}', file=sys.stderr)
        return False


def main():
    parser = argparse.ArgumentParser(description='生成投后管理报告')
    parser.add_argument('--template', required=True, help='上季度报告模板路径（DOCX）')
    parser.add_argument('--financial-data', required=True, help='财务数据（JSON字符串或文件路径）')
    parser.add_argument('--financial-analysis', required=True, help='财务分析文本')
    parser.add_argument('--business-update', required=True, help='经营情况更新文本')
    parser.add_argument('--industry-update', required=True, help='行业分析更新文本')
    parser.add_argument('--output', required=True, help='输出报告路径')
    
    args = parser.parse_args()
    
    # 处理财务数据参数
    financial_data = args.financial_data
    if os.path.isfile(args.financial_data):
        # 如果是文件路径，读取文件内容
        with open(args.financial_data, 'r', encoding='utf-8') as f:
            financial_data = f.read()
    
    # 生成报告
    success = generate_report(
        template_path=args.template,
        financial_data=financial_data,
        financial_analysis=args.financial_analysis,
        business_update=args.business_update,
        industry_update=args.industry_update,
        output_path=args.output
    )
    
    if success:
        print(f'报告已成功生成: {args.output}')
        sys.exit(0)
    else:
        sys.exit(1)


if __name__ == '__main__':
    main()
