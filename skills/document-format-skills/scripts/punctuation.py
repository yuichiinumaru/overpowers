#!/usr/bin/env python3
"""
标点符号修复 v5
- 修复引号处理bug：使用明确的Unicode转义序列
- 正确处理省略号和句号
"""

import re
import sys
from docx import Document

# 中文标点（使用Unicode转义确保正确）
LEFT_DOUBLE_QUOTE = '\u201c'   # " 左双引号
RIGHT_DOUBLE_QUOTE = '\u201d'  # " 右双引号
LEFT_SINGLE_QUOTE = '\u2018'   # ' 左单引号
RIGHT_SINGLE_QUOTE = '\u2019'  # ' 右单引号

# 基本替换映射
REPLACEMENTS = {
    "(": "（",
    ")": "）",
    ":": "：",
    ";": "；",
    "?": "？",
    "!": "！",
}


def has_chinese(text):
    """检查是否包含中文"""
    return bool(re.search(r"[\u4e00-\u9fff]", text))


def fix_text(text):
    """修复文本中的标点"""
    if not text:
        return text

    result = text

    # ===== 第一步：处理省略号（必须在句号之前）=====
    result = re.sub(r"\.{2,}", "……", result)
    result = re.sub(r"。{2,}", "……", result)

    # ===== 第二步：处理破折号 =====
    result = re.sub(r"--+", "——", result)
    result = re.sub(r"—(?!—)", "——", result)

    # ===== 第三步：基本标点替换（只在有中文的文本中）=====
    if has_chinese(result):
        for en, cn in REPLACEMENTS.items():
            result = result.replace(en, cn)

    # ===== 第四步：逗号特殊处理 =====
    result = re.sub(r"([\u4e00-\u9fff]),", r"\1，", result)
    result = re.sub(r",([\u4e00-\u9fff])", r"，\1", result)

    # ===== 第五步：句号特殊处理 =====
    result = re.sub(r"([\u4e00-\u9fff])\.(\s|$)", r"\1。\2", result)

    # ===== 第六步：双引号处理 =====
    # 需要处理的双引号字符（使用Unicode确保正确）
    double_quote_chars = [
        '"',        # U+0022 ASCII直引号
        '\u201c',   # U+201C 左双引号 "
        '\u201d',   # U+201D 右双引号 "
        '\u201e',   # U+201E 双低引号 „
        '\u201f',   # U+201F 双高反引号 ‟
        '\u300c',   # U+300C 日文左引号 「
        '\u300d',   # U+300D 日文右引号 」
    ]

    # 统一替换为临时标记
    temp_result = result
    for q in double_quote_chars:
        temp_result = temp_result.replace(q, "\x00")

    # 配对处理
    if "\x00" in temp_result:
        chars = list(temp_result)
        quote_count = 0
        for i, c in enumerate(chars):
            if c == "\x00":
                # 偶数位置用左引号，奇数位置用右引号
                if quote_count % 2 == 0:
                    chars[i] = LEFT_DOUBLE_QUOTE  # "
                else:
                    chars[i] = RIGHT_DOUBLE_QUOTE  # "
                quote_count += 1
        result = "".join(chars)

    # ===== 第七步：单引号处理 =====
    single_quote_chars = [
        "'",        # U+0027 ASCII单引号
        '\u2018',   # U+2018 左单引号 '
        '\u2019',   # U+2019 右单引号 '
        '\u201a',   # U+201A 单低引号 ‚
        '\u201b',   # U+201B 单高反引号 ‛
    ]

    temp_result = result
    for q in single_quote_chars:
        temp_result = temp_result.replace(q, "\x01")

    if "\x01" in temp_result:
        chars = list(temp_result)
        quote_count = 0
        for i, c in enumerate(chars):
            if c == "\x01":
                # 偶数位置用左引号，奇数位置用右引号
                if quote_count % 2 == 0:
                    chars[i] = LEFT_SINGLE_QUOTE  # '
                else:
                    chars[i] = RIGHT_SINGLE_QUOTE  # '
                quote_count += 1
        result = "".join(chars)

    return result


def process_paragraph(para):
    """处理段落 - 合并所有run的文本，修复后重新分配"""
    full_text = para.text
    if not full_text.strip():
        return False

    fixed_text = fix_text(full_text)

    if fixed_text == full_text:
        return False

    runs = para.runs
    if not runs:
        return False

    # 把修复后的文本放到第一个run，清空其他run
    first_run = runs[0]
    first_run.text = fixed_text

    for run in runs[1:]:
        run.text = ""

    return True


def process_document(input_path, output_path):
    """处理文档"""
    print(f"Reading: {input_path}")
    doc = Document(input_path)

    changes = 0

    # 处理段落
    for i, para in enumerate(doc.paragraphs):
        if process_paragraph(para):
            changes += 1
            preview = para.text[:50] + "..." if len(para.text) > 50 else para.text
            print(f"  Para {i + 1}: {preview}")

    # 处理表格
    table_changes = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if process_paragraph(para):
                        table_changes += 1

    if table_changes > 0:
        print(f"  Tables: {table_changes} cells fixed")

    print()
    print(f"Total: {changes} paragraphs + {table_changes} table cells fixed")
    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python punctuation.py input.docx output.docx")
        sys.exit(1)

    process_document(sys.argv[1], sys.argv[2])
