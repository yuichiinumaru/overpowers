#!/usr/bin/env python3
"""
文档格式统一 v5
修复问题：
- 标题检测更全面
- 主送机关顶格
- 落款右对齐
- 清除斜体、下划线、颜色
- 特殊段落处理（附件、特此说明等）

公文标准：
- 页边距：上37mm，下35mm，左28mm，右26mm
- 主标题：居中，二号（22pt），方正小标宋简体
- 主送机关：顶格，三号仿宋
- 正文：三号仿宋GB2312，首行缩进2字符，行距28磅
- 一级标题："一、" 三号黑体，首行缩进2字符
- 二级标题："（一）" 三号楷体GB2312，首行缩进2字符
- 三级标题："1." 三号仿宋GB2312，首行缩进2字符
- 四级标题："（1）" 三号仿宋GB2312，首行缩进2字符
- 落款：右对齐，三号仿宋
- 附件：顶格，三号仿宋
"""

import sys
import re
from docx import Document
from docx.shared import Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 字号对照：二号=22pt，三号=16pt，小四=12pt
# 2字符缩进 = 2 × 16pt = 32pt（三号字）

PRESETS = {
    'official': {
        'name': '公文格式',
        'page': {'top': 3.7, 'bottom': 3.5, 'left': 2.8, 'right': 2.6},
        # 主标题：二号方正小标宋简体，居中
        'title': {
            'font_cn': '方正小标宋简体',
            'font_en': 'Times New Roman',
            'size': 22,  # 二号
            'bold': False,
            'align': 'center',
            'indent': 0,
        },
        # 主送机关：三号仿宋，顶格
        'recipient': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 0,  # 顶格
        },
        # 一级标题：三号黑体，"一、"，首行缩进2字符
        'heading1': {
            'font_cn': '黑体',
            'font_en': 'Times New Roman',
            'size': 16,  # 三号
            'bold': False,
            'align': 'left',
            'indent': 32,  # 2字符缩进
        },
        # 二级标题：三号楷体GB2312，"（一）"，首行缩进2字符
        'heading2': {
            'font_cn': '楷体_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 32,
        },
        # 三级标题：三号仿宋GB2312，"1."，首行缩进2字符
        'heading3': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 32,
        },
        # 四级标题：三号仿宋GB2312，"（1）"，首行缩进2字符
        'heading4': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 32,
        },
        # 正文：三号仿宋GB2312，首行缩进2字符（32pt），行距28磅
        'body': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'justify',
            'indent': 32,  # 2字符 = 2×16pt
            'line_spacing': 28,
        },
        # 落款单位：三号仿宋，右对齐
        'signature': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'right',
            'indent': 0,
        },
        # 落款日期：三号仿宋，右对齐
        'date': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'right',
            'indent': 0,
        },
        # 附件行：三号仿宋，格式同正文（首行缩进2字符）
        'attachment': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'justify',
            'indent': 32,  # 同正文，首行缩进2字符
        },
        # 结束语（特此说明/通知等）：三号仿宋，首行缩进
        'closing': {
            'font_cn': '仿宋_GB2312',
            'font_en': 'Times New Roman',
            'size': 16,
            'bold': False,
            'align': 'left',
            'indent': 32,
        },
    },
    'academic': {
        'name': '学术论文格式',
        'page': {'top': 2.5, 'bottom': 2.5, 'left': 2.5, 'right': 2.5},
        'title': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 18, 'bold': True, 'align': 'center', 'indent': 0},
        'recipient': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 0},
        'heading1': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 15, 'bold': True, 'align': 'left', 'indent': 0},
        'heading2': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 14, 'bold': True, 'align': 'left', 'indent': 0},
        'heading3': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 0},
        'heading4': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 0},
        'body': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'justify', 'indent': 24, 'line_spacing': None},
        'signature': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'right', 'indent': 0},
        'date': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'right', 'indent': 0},
        'attachment': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 0},
        'closing': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 12, 'bold': False, 'align': 'left', 'indent': 24},
    },
    'legal': {
        'name': '法律文书格式',
        'page': {'top': 3.0, 'bottom': 2.5, 'left': 3.0, 'right': 2.5},
        'title': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 22, 'bold': True, 'align': 'center', 'indent': 0},
        'recipient': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'heading1': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'heading2': {'font_cn': '黑体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'heading3': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'heading4': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'body': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'justify', 'indent': 28, 'line_spacing': None},
        'signature': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'right', 'indent': 0},
        'date': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'right', 'indent': 0},
        'attachment': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 0},
        'closing': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': 14, 'bold': False, 'align': 'left', 'indent': 28},
    },
}


def remove_background(doc):
    """移除页面背景颜色"""
    body = doc._body._body
    document = body.getparent()
    for elem in list(document):
        tag_name = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag_name == 'background':
            document.remove(elem)
    
    for para in doc.paragraphs:
        pPr = para._p.get_or_add_pPr()
        shd = pPr.find(qn('w:shd'))
        if shd is not None:
            pPr.remove(shd)
        for run in para.runs:
            run.font.highlight_color = None
            rPr = run._r.get_or_add_rPr()
            shd = rPr.find(qn('w:shd'))
            if shd is not None:
                rPr.remove(shd)


def detect_para_type(text, index, total, alignment, all_texts):
    """
    检测段落类型
    返回: 'title', 'recipient', 'heading1', 'heading2', 'heading3', 'heading4', 
          'body', 'signature', 'date', 'attachment', 'closing'
    
    参数:
        text: 段落文本
        index: 段落索引
        total: 总段落数
        alignment: 原始对齐方式
        all_texts: 所有非空段落的文本列表，用于上下文判断
    """
    text = text.strip()
    if not text:
        return 'empty'
    
    # ===== 一级标题："一、" "二、" 等 =====
    if re.match(r'^[一二三四五六七八九十]+、', text):
        return 'heading1'
    
    # ===== 二级标题："（一）" "（二）" 等 =====
    if re.match(r'^（[一二三四五六七八九十]+）', text):
        return 'heading2'
    if re.match(r'^\([一二三四五六七八九十]+\)', text):
        return 'heading2'
    
    # ===== 三级标题："1." "2." 等 =====
    if re.match(r'^\d+\.\s*\S', text) and len(text) < 60:
        return 'heading3'
    
    # ===== 四级标题："（1）" "（2）" 等 =====
    if re.match(r'^（\d+）', text) and len(text) < 60:
        return 'heading4'
    if re.match(r'^\(\d+\)', text) and len(text) < 60:
        return 'heading4'
    
    # ===== 主送机关：XXX： 或 XXX: =====
    # 通常在文档开头几段，以冒号结尾，且较短
    if re.match(r'^[\u4e00-\u9fff]+[：:]$', text) and len(text) < 20:
        return 'recipient'
    
    # ===== 附件行 =====
    if re.match(r'^附件[：:]\s*', text):
        return 'attachment'
    if re.match(r'^附件\d*[：:．.\s]', text):
        return 'attachment'
    if re.match(r'^附件$', text):
        return 'attachment'
    
    # ===== 结束语 =====
    closing_patterns = [
        r'^特此(说明|通知|报告|函复|函告|批复|公告|通报)。?$',
        r'^此致$',
        r'^敬礼[！!]?$',
        r'^以上(报告|意见|方案).{0,10}$',
        r'^妥否.{0,10}$',
        r'^请.{0,15}(批示|审批|审议|指示|核准)。?$',
    ]
    for pattern in closing_patterns:
        if re.match(pattern, text):
            return 'closing'
    
    # ===== 落款日期 =====
    # 支持多种日期格式
    date_patterns = [
        r'^\d{4}年\d{1,2}月\d{1,2}日$',
        r'^\d{4}\.\d{1,2}\.\d{1,2}$',
        r'^\d{4}/\d{1,2}/\d{1,2}$',
        r'^\d{4}-\d{1,2}-\d{1,2}$',
        r'^二[○〇零oO0][一二三四五六七八九零〇○oO0]{2}年.{1,3}月.{1,3}日$',
    ]
    for pattern in date_patterns:
        if re.match(pattern, text):
            return 'date'
    
    # ===== 落款单位 =====
    # 判断逻辑：在文档后部，短文本，且下一段是日期或者是文档末尾附近
    if index >= total - 10 and len(text) < 30:
        # 检查是否像单位名称
        if re.search(r'(公司|局|委|部|厅|院|所|中心|办公室|集团|银行|学校|大学|医院)$', text):
            return 'signature'
        # 或者检查下文是否有日期
        remaining_texts = all_texts[all_texts.index(text)+1:] if text in all_texts else []
        for next_text in remaining_texts[:3]:
            for pattern in date_patterns:
                if re.match(pattern, next_text.strip()):
                    return 'signature'
    
    # ===== 主标题 =====
    # 判断条件：在前5段，且满足以下条件之一
    if index < 5:
        # 1. 明确的公文标题模式
        title_patterns = [
            r'^关于.+的(通知|报告|请示|函|意见|决定|公告|通报|批复|说明|方案|总结|汇报|复函|答复|建议)$',
            r'^.{2,30}(通知|报告|请示|函|意见|决定|公告|通报|批复|工作方案|工作总结|实施方案|管理办法|暂行规定)$',
        ]
        for pattern in title_patterns:
            if re.match(pattern, text):
                return 'title'
        
        # 2. 较长的标题（20-80字符），不以标点结尾
        if 15 < len(text) < 80 and not re.search(r'[。！？，、；：]$', text):
            # 排除以序号开头的
            if not re.match(r'^[一二三四五六七八九十\d（(]', text):
                return 'title'
        
        # 3. 居中的短文本（原本就是居中的）
        if alignment == WD_ALIGN_PARAGRAPH.CENTER and len(text) < 60:
            return 'title'
    
    # ===== 其他都是正文 =====
    return 'body'


def set_font(run, font_cn, font_en, size, bold=False):
    """
    设置字体，同时清除原有格式（斜体、下划线、颜色）
    """
    # 基本字体设置
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    
    # 清除斜体
    run.font.italic = False
    
    # 清除下划线
    run.font.underline = False
    
    # 清除颜色（设置为黑色）
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # 清除删除线
    run.font.strike = False
    run.font.double_strike = False
    
    # 清除上下标
    run.font.subscript = False
    run.font.superscript = False
    
    # 设置中文字体
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)


def format_paragraph(para, fmt, para_type, line_spacing_pt=28):
    """格式化段落"""
    pf = para.paragraph_format
    
    # 对齐方式
    align_map = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    pf.alignment = align_map.get(fmt.get('align', 'justify'), WD_ALIGN_PARAGRAPH.JUSTIFY)
    
    # 段落左缩进清零（重要：确保"文本之前缩进"为0）
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    
    # 首行缩进
    indent = fmt.get('indent', 0)
    if indent > 0:
        pf.first_line_indent = Pt(indent)
    else:
        pf.first_line_indent = Pt(0)
    
    # 行距（固定值28磅，或使用配置）
    ls = fmt.get('line_spacing', line_spacing_pt)
    if ls:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(ls)
    else:
        pf.line_spacing = 1.5
    
    # 段前段后
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    
    # 字体
    for run in para.runs:
        set_font(run, fmt['font_cn'], fmt['font_en'], fmt['size'], fmt.get('bold', False))


def add_page_number(doc):
    """添加页码（底端居中）"""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # 清空现有内容
        for para in footer.paragraphs:
            para.clear()
        
        # 添加页码
        if footer.paragraphs:
            para = footer.paragraphs[0]
        else:
            para = footer.add_paragraph()
        
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加页码域
        run = para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        run2 = para.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        run2._r.append(instrText)
        
        run3 = para.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._r.append(fldChar2)


def format_document(input_path, output_path, preset_name='official'):
    """格式化文档"""
    if preset_name not in PRESETS:
        print(f'Unknown preset: {preset_name}')
        print(f'Available: {", ".join(PRESETS.keys())}')
        sys.exit(1)
    
    preset = PRESETS[preset_name]
    print(f'Preset: {preset["name"]}')
    print(f'Input: {input_path}')
    
    doc = Document(input_path)
    total_paras = len(doc.paragraphs)
    
    # 收集所有非空段落文本，用于上下文判断
    all_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # 1. 移除背景
    print('1. Removing background...')
    remove_background(doc)
    
    # 2. 设置页面边距
    print('2. Setting page margins...')
    page = preset['page']
    for section in doc.sections:
        section.top_margin = Cm(page['top'])
        section.bottom_margin = Cm(page['bottom'])
        section.left_margin = Cm(page['left'])
        section.right_margin = Cm(page['right'])
    
    # 3. 格式化段落
    print('3. Formatting paragraphs...')
    stats = {
        'title': 0, 'recipient': 0, 'heading1': 0, 'heading2': 0, 
        'heading3': 0, 'heading4': 0, 'body': 0, 'signature': 0, 
        'date': 0, 'attachment': 0, 'closing': 0
    }
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        para_type = detect_para_type(
            text, i, total_paras, 
            para.paragraph_format.alignment,
            all_texts
        )
        
        # 选择对应的格式
        fmt_key = para_type if para_type in preset else 'body'
        fmt = preset.get(fmt_key, preset['body'])
        
        format_paragraph(para, fmt, para_type)
        stats[para_type] = stats.get(para_type, 0) + 1
        
        # 打印处理信息
        preview = text[:35] + '...' if len(text) > 35 else text
        print(f'   [{para_type:10}] {preview}')
    
    # 4. 处理表格
    print('4. Formatting tables...')
    body_fmt = preset['body']
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        for run in para.runs:
                            set_font(run, body_fmt['font_cn'], body_fmt['font_en'], body_fmt['size'])
                        para.paragraph_format.first_line_indent = Pt(0)
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                        para.paragraph_format.line_spacing = Pt(28)
    
    # 5. 添加页码
    print('5. Adding page numbers...')
    add_page_number(doc)
    
    # 保存
    doc.save(output_path)
    
    print()
    print('=' * 50)
    print('Statistics:')
    for k, v in stats.items():
        if v > 0:
            print(f'  {k}: {v}')
    print(f'Output: {output_path}')


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: python formatter.py input.docx output.docx [--preset official|academic|legal]')
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    preset = 'official'
    if '--preset' in sys.argv:
        idx = sys.argv.index('--preset')
        if idx + 1 < len(sys.argv):
            preset = sys.argv[idx + 1]
    
    format_document(input_file, output_file, preset)
