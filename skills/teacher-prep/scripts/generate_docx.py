#!/usr/bin/env python3
"""
生成课后练习题 Word 文档 - 支持多种课文类型
用法: python generate_docx.py <课文名> <课文类型> [备课资料文件路径]
"""

import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='SimSun', size=12, bold=False):
    """设置中文字体"""
    font = run.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def create_exercises_poem(lesson_name, doc):
    """生成古诗类练习题"""
    
    # 一、基础积累
    heading = doc.add_heading('一、基础积累', level=1)
    for run in heading.runs:
        set_chinese_font(run, 'SimHei', 14, True)
    
    p = doc.add_paragraph()
    run = p.add_run(f'1. 默写《{lesson_name}》')
    set_chinese_font(run, 'SimSun', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('（请在此处默写全诗）')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('2. 给下列加点字注音')
    set_chinese_font(run, 'SimSun', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('（    ）（    ）（    ）（    ）')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('3. 看拼音写词语')
    set_chinese_font(run, 'SimSun', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('（      ）（      ）（      ）')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    # 二、理解感悟
    heading = doc.add_heading('二、理解感悟', level=1)
    for run in heading.runs:
        set_chinese_font(run, 'SimHei', 14, True)
    
    p = doc.add_paragraph()
    run = p.add_run('4. 解释下列词语的意思')
    set_chinese_font(run, 'SimSun', 12)
    
    for i in range(1, 4):
        p = doc.add_paragraph()
        run = p.add_run(f'（{i}）______________________________')
        set_chinese_font(run, 'SimSun', 11)
        p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('5. 写出诗句的意思')
    set_chinese_font(run, 'SimSun', 12)
    
    for i in range(1, 3):
        p = doc.add_paragraph()
        run = p.add_run(f'（{i}）______________________________________________')
        set_chinese_font(run, 'SimSun', 11)
        p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('6. 这首诗表达了诗人怎样的情感？')
    set_chinese_font(run, 'SimSun', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('________________________________________________________________')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)

def create_exercises_modern(lesson_name, doc):
    """生成现代文类练习题"""
    
    # 一、基础积累
    heading = doc.add_heading('一、基础积累', level=1)
    for run in heading.runs:
        set_chinese_font(run, 'SimHei', 14, True)
    
    p = doc.add_paragraph()
    run = p.add_run('1. 看拼音写词语')
    set_chinese_font(run, 'SimSun', 12)
    
    for i in range(1, 4):
        p = doc.add_paragraph()
        run = p.add_run(f'（{i}）（      ）（      ）（      ）')
        set_chinese_font(run, 'SimSun', 11)
        p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('2. 给加点字选择正确的读音，打"√"')
    set_chinese_font(run, 'SimSun', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('（    ）  （    ）  （    ）  （    ）')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('3. 比一比，再组词')
    set_chinese_font(run, 'SimSun', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('（    ）——（    ）  （    ）——（    ）')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    # 二、理解感悟
    heading = doc.add_heading('二、理解感悟', level=1)
    for run in heading.runs:
        set_chinese_font(run, 'SimHei', 14, True)
    
    p = doc.add_paragraph()
    run = p.add_run('4. 根据课文内容填空')
    set_chinese_font(run, 'SimSun', 12)
    
    for i in range(1, 4):
        p = doc.add_paragraph()
        run = p.add_run(f'（{i}）______________________________________________')
        set_chinese_font(run, 'SimSun', 11)
        p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('5. 读句子，回答问题')
    set_chinese_font(run, 'SimSun', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('（1）________________________________________________________________')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('________________________________________________________________')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('6. 概括课文主要内容')
    set_chinese_font(run, 'SimSun', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('________________________________________________________________')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('________________________________________________________________')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)

def create_exercises_fable(lesson_name, doc):
    """生成寓言/童话类练习题"""
    
    # 一、基础积累
    heading = doc.add_heading('一、基础积累', level=1)
    for run in heading.runs:
        set_chinese_font(run, 'SimHei', 14, True)
    
    p = doc.add_paragraph()
    run = p.add_run('1. 看拼音写词语')
    set_chinese_font(run, 'SimSun', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('（      ）（      ）（      ）（      ）')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('2. 给下列词语选择正确的解释')
    set_chinese_font(run, 'SimSun', 12)
    
    for i in range(1, 4):
        p = doc.add_paragraph()
        run = p.add_run(f'（{i}）__________________（    ）')
        set_chinese_font(run, 'SimSun', 11)
        p.paragraph_format.left_indent = Inches(0.3)
    
    # 二、理解感悟
    heading = doc.add_heading('二、理解感悟', level=1)
    for run in heading.runs:
        set_chinese_font(run, 'SimHei', 14, True)
    
    p = doc.add_paragraph()
    run = p.add_run('3. 根据故事内容填空')
    set_chinese_font(run, 'SimSun', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('故事的起因：______________________________________________')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('故事的经过：______________________________________________')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('故事的结果：______________________________________________')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('4. 这个故事告诉我们什么道理？')
    set_chinese_font(run, 'SimSun', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('________________________________________________________________')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('5. 你喜欢故事中的谁？为什么？')
    set_chinese_font(run, 'SimSun', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('________________________________________________________________')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)

def create_exercises(lesson_name, lesson_type, content_file=None):
    """生成课后练习题 Word 文档"""
    
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # 标题
    title = doc.add_heading(f'《{lesson_name}》课后练习', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_chinese_font(run, 'SimHei', 18, True)
    
    # 根据课文类型生成不同练习题
    if lesson_type in ['古诗', '古诗词']:
        create_exercises_poem(lesson_name, doc)
    elif lesson_type in ['寓言', '童话']:
        create_exercises_fable(lesson_name, doc)
    else:
        create_exercises_modern(lesson_name, doc)
    
    # 三、拓展提升（通用）
    heading = doc.add_heading('三、拓展提升', level=1)
    for run in heading.runs:
        set_chinese_font(run, 'SimHei', 14, True)
    
    p = doc.add_paragraph()
    run = p.add_run('7. 联系生活实际，谈谈你的感受或收获')
    set_chinese_font(run, 'SimSun', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('________________________________________________________________')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('________________________________________________________________')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    p = doc.add_paragraph()
    run = p.add_run('8. 创意表达（画一画、写一写、演一演）')
    set_chinese_font(run, 'SimSun', 12)
    
    p = doc.add_paragraph()
    run = p.add_run('________________________________________________________________')
    set_chinese_font(run, 'SimSun', 11)
    p.paragraph_format.left_indent = Inches(0.3)
    
    # 分页：参考答案
    doc.add_page_break()
    
    title = doc.add_heading(f'《{lesson_name}》参考答案', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_chinese_font(run, 'SimHei', 18, True)
    
    p = doc.add_paragraph()
    run = p.add_run('（教师根据实际教学情况填写参考答案）')
    set_chinese_font(run, 'SimSun', 12)
    
    # 保存
    output_file = f"练习题_{lesson_name}.docx"
    doc.save(output_file)
    print(f"练习题已生成: {output_file}")
    return output_file

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("用法: python generate_docx.py <课文名> <课文类型> [备课资料文件路径]")
        print("课文类型: 古诗/现代文/寓言/童话/说明文")
        sys.exit(1)
    
    lesson_name = sys.argv[1]
    lesson_type = sys.argv[2]
    content_file = sys.argv[3] if len(sys.argv) > 3 else None
    
    create_exercises(lesson_name, lesson_type, content_file)
