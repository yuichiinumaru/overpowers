#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
图片处理与图文混排模块
支持图片插入、环绕方式设置、批量图文排版等功能
"""

import os
import re
from pathlib import Path
from typing import List, Tuple, Optional, Union
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement, parse_xml
from PIL import Image


class ImageWrapStyle:
    """图片环绕方式"""
    INLINE = 'inline'           # 嵌入型
    SQUARE = 'square'           # 四周型
    TIGHT = 'tight'             # 紧密型
    THROUGH = 'through'         # 穿越型
    TOP_BOTTOM = 'topBottom'    # 上下型
    BEHIND = 'behind'           # 衬于文字下方
    IN_FRONT = 'inFront'        # 浮于文字上方


class ImageHandler:
    """图片处理器"""
    
    def __init__(self, doc: Document = None):
        self.doc = doc
        self.default_width = Inches(4)
        self.default_height = None  # 保持比例
        
    def set_document(self, doc: Document):
        """设置文档对象"""
        self.doc = doc
    
    def insert_image(self, 
                     image_path: str,
                     paragraph=None,
                     width: Inches = None,
                     height: Inches = None,
                     wrap_style: str = ImageWrapStyle.INLINE,
                     align: str = 'left',
                     caption: str = None) -> bool:
        """
        插入图片到文档
        
        Args:
            image_path: 图片路径
            paragraph: 指定段落（可选）
            width: 图片宽度
            height: 图片高度
            wrap_style: 环绕方式
            align: 对齐方式 (left/center/right)
            caption: 图片标题（可选）
        
        Returns:
            是否成功
        """
        try:
            if not os.path.exists(image_path):
                print(f"图片不存在: {image_path}")
                return False
            
            # 如果没有指定段落，创建新段落
            if paragraph is None:
                paragraph = self.doc.add_paragraph()
            
            # 设置对齐
            if align == 'center':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 添加图片
            run = paragraph.add_run()
            
            # 设置图片尺寸
            img_width = width if width else self.default_width
            img_height = height if height else self.default_height
            
            if img_height:
                inline_shape = run.add_picture(image_path, width=img_width, height=img_height)
            else:
                inline_shape = run.add_picture(image_path, width=img_width)
            
            # 设置环绕方式
            if wrap_style != ImageWrapStyle.INLINE:
                self._set_image_wrap(inline_shape, wrap_style)
            
            # 添加标题
            if caption:
                caption_para = self.doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption_para.add_run(caption)
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            return True
            
        except Exception as e:
            print(f"插入图片失败: {e}")
            return False
    
    def _set_image_wrap(self, inline_shape, wrap_style: str):
        """设置图片环绕方式"""
        try:
            # 获取图片的 anchor 元素
            anchor = inline_shape._inline.getparent()
            
            # 创建新的 anchor 元素（浮动图片）
            # 这需要在 XML 层面操作
            # 简化处理：使用 python-docx 的底层 API
            
            # 获取或创建 wrap 元素
            inline = inline_shape._inline
            graphic = inline.graphic
            graphicData = graphic.graphicData
            pic = graphicData.pic
            
            # 创建 wrap 元素
            wrap = OxmlElement('wp:wrap' + wrap_style.capitalize())
            if wrap_style in ['square', 'tight', 'through']:
                wrap.set(qn('wrapText'), 'bothSides')
            
            # 这里需要更复杂的 XML 操作来真正改变环绕方式
            # 目前 python-docx 对环绕方式的支持有限
            
        except Exception as e:
            print(f"设置环绕方式失败: {e}")
    
    def insert_images_grid(self, 
                          image_paths: List[str],
                          cols: int = 2,
                          width: Inches = None,
                          captions: List[str] = None):
        """
        以网格形式插入多张图片
        
        Args:
            image_paths: 图片路径列表
            cols: 列数
            width: 单张图片宽度
            captions: 图片标题列表（可选）
        """
        if not image_paths:
            return
        
        # 计算图片宽度
        if width is None:
            # 默认每行图片总宽度为 6 英寸
            page_width = Inches(6)
            width = page_width / cols
        
        # 创建表格布局
        rows = (len(image_paths) + cols - 1) // cols
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # 隐藏表格边框
        for row in table.rows:
            for cell in row.cells:
                cell._element.get_or_add_tcPr().append(
                    parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
                )
        
        # 填充图片
        for idx, image_path in enumerate(image_paths):
            if not os.path.exists(image_path):
                continue
            
            row = idx // cols
            col = idx % cols
            cell = table.rows[row].cells[col]
            
            # 清空单元格
            cell.text = ''
            
            # 插入图片
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = paragraph.add_run()
            try:
                run.add_picture(image_path, width=width)
                
                # 添加标题
                if captions and idx < len(captions):
                    caption_para = cell.add_paragraph()
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.add_run(captions[idx])
                    caption_run.font.size = Pt(9)
                    caption_run.font.name = '宋体'
                    caption_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            except Exception as e:
                print(f"插入图片失败 {image_path}: {e}")
    
    def create_text_image_layout(self,
                                  text: str,
                                  image_path: str,
                                  layout: str = 'left',
                                  image_width: Inches = Inches(2.5)):
        """
        创建图文混排布局
        
        Args:
            text: 文本内容
            image_path: 图片路径
            layout: 布局方式 (left/right/top/bottom)
            image_width: 图片宽度
        """
        if layout == 'left':
            # 图片左，文字右
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # 隐藏边框
            for row in table.rows:
                for cell in row.cells:
                    cell._element.get_or_add_tcPr().append(
                        parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
                    )
            
            # 左侧图片
            left_cell = table.rows[0].cells[0]
            left_cell.text = ''
            left_para = left_cell.paragraphs[0]
            left_run = left_para.add_run()
            if os.path.exists(image_path):
                left_run.add_picture(image_path, width=image_width)
            
            # 右侧文字
            right_cell = table.rows[0].cells[1]
            right_cell.text = text
            
        elif layout == 'right':
            # 图片右，文字左
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # 隐藏边框
            for row in table.rows:
                for cell in row.cells:
                    cell._element.get_or_add_tcPr().append(
                        parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
                    )
            
            # 左侧文字
            left_cell = table.rows[0].cells[0]
            left_cell.text = text
            
            # 右侧图片
            right_cell = table.rows[0].cells[1]
            right_cell.text = ''
            right_para = right_cell.paragraphs[0]
            right_run = right_para.add_run()
            if os.path.exists(image_path):
                right_run.add_picture(image_path, width=image_width)
                
        elif layout == 'top':
            # 图片上，文字下
            if os.path.exists(image_path):
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(image_path, width=image_width)
            
            self.doc.add_paragraph(text)
            
        elif layout == 'bottom':
            # 文字上，图片下
            self.doc.add_paragraph(text)
            
            if os.path.exists(image_path):
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(image_path, width=image_width)
    
    def resize_image(self, image_path: str, output_path: str, 
                     max_width: int = None, max_height: int = None,
                     quality: int = 95) -> bool:
        """
        调整图片尺寸
        
        Args:
            image_path: 原图路径
            output_path: 输出路径
            max_width: 最大宽度（像素）
            max_height: 最大高度（像素）
            quality: 图片质量（1-100）
        
        Returns:
            是否成功
        """
        try:
            with Image.open(image_path) as img:
                # 转换为 RGB（处理 PNG 等格式）
                if img.mode in ('RGBA', 'LA', 'P'):
                    img = img.convert('RGB')
                
                # 计算新尺寸
                width, height = img.size
                ratio = 1.0
                
                if max_width and width > max_width:
                    ratio = min(ratio, max_width / width)
                
                if max_height and height > max_height:
                    ratio = min(ratio, max_height / height)
                
                if ratio < 1.0:
                    new_width = int(width * ratio)
                    new_height = int(height * ratio)
                    img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
                
                # 保存
                img.save(output_path, quality=quality, optimize=True)
                return True
                
        except Exception as e:
            print(f"调整图片尺寸失败: {e}")
            return False
    
    def batch_resize(self, image_paths: List[str], 
                     output_dir: str,
                     max_width: int = 800,
                     max_height: int = 600) -> List[str]:
        """
        批量调整图片尺寸
        
        Args:
            image_paths: 图片路径列表
            output_dir: 输出目录
            max_width: 最大宽度
            max_height: 最大高度
        
        Returns:
            输出文件路径列表
        """
        os.makedirs(output_dir, exist_ok=True)
        output_paths = []
        
        for image_path in image_paths:
            if not os.path.exists(image_path):
                continue
            
            filename = os.path.basename(image_path)
            output_path = os.path.join(output_dir, f"resized_{filename}")
            
            if self.resize_image(image_path, output_path, max_width, max_height):
                output_paths.append(output_path)
        
        return output_paths


class MarkdownImageProcessor:
    """Markdown 图片处理器 - 处理 Markdown 中的图片语法"""
    
    def __init__(self, doc: Document = None, base_path: str = None):
        self.doc = doc
        self.base_path = base_path or os.getcwd()
        self.image_handler = ImageHandler(doc)
        
    def set_document(self, doc: Document):
        """设置文档对象"""
        self.doc = doc
        self.image_handler.set_document(doc)
    
    def process_markdown_images(self, md_content: str, default_width: Inches = Inches(4)):
        """
        处理 Markdown 内容中的图片，插入到 Word 文档
        
        Args:
            md_content: Markdown 内容
            default_width: 默认图片宽度
        """
        # 匹配 Markdown 图片语法 ![alt](path "title")
        image_pattern = r'!\[([^\]]*)\]\(([^)"\s]+)(?:\s+"([^"]*)")?\)'
        
        lines = md_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            
            # 检查是否是单独的图片行
            match = re.match(r'^!\[([^\]]*)\]\(([^)"\s]+)(?:\s+"([^"]*)")?\)\s*$', line.strip())
            if match:
                alt_text = match.group(1)
                image_path = match.group(2)
                title = match.group(3)
                
                # 解析图片路径
                if not os.path.isabs(image_path):
                    image_path = os.path.join(self.base_path, image_path)
                
                # 插入图片
                caption = title if title else alt_text
                self.image_handler.insert_image(
                    image_path,
                    width=default_width,
                    align='center',
                    caption=caption if caption else None
                )
                
                i += 1
                continue
            
            # 检查行内图片
            if '![' in line:
                # 处理行内图片（简化处理，直接插入图片）
                parts = re.split(image_pattern, line)
                if len(parts) > 1:
                    # 有图片，创建段落
                    paragraph = self.doc.add_paragraph()
                    
                    # 提取所有图片匹配
                    matches = re.findall(image_pattern, line)
                    for alt, img_path, title in matches:
                        if not os.path.isabs(img_path):
                            img_path = os.path.join(self.base_path, img_path)
                        
                        self.image_handler.insert_image(
                            img_path,
                            paragraph=paragraph,
                            width=Inches(2),
                            align='left'
                        )
                    
                    i += 1
                    continue
            
            i += 1
    
    def extract_images_from_markdown(self, md_content: str) -> List[dict]:
        """
        从 Markdown 内容中提取图片信息
        
        Args:
            md_content: Markdown 内容
        
        Returns:
            图片信息列表 [{'alt': '', 'path': '', 'title': ''}]
        """
        image_pattern = r'!\[([^\]]*)\]\(([^)"\s]+)(?:\s+"([^"]*)")?\)'
        matches = re.findall(image_pattern, md_content)
        
        images = []
        for alt, path, title in matches:
            images.append({
                'alt': alt,
                'path': path,
                'title': title
            })
        
        return images


class BatchImageLayout:
    """批量图文排版器"""
    
    def __init__(self, doc: Document = None):
        self.doc = doc
        self.image_handler = ImageHandler(doc)
    
    def set_document(self, doc: Document):
        """设置文档对象"""
        self.doc = doc
        self.image_handler.set_document(doc)
    
    def create_gallery(self, 
                       items: List[dict],
                       layout: str = 'grid',
                       cols: int = 2):
        """
        创建图文画廊
        
        Args:
            items: 图文项列表 [{'text': '', 'image': '', 'title': ''}]
            layout: 布局方式 (grid/list)
            cols: 网格列数
        """
        if layout == 'grid':
            # 网格布局
            self._create_grid_gallery(items, cols)
        else:
            # 列表布局
            self._create_list_gallery(items)
    
    def _create_grid_gallery(self, items: List[dict], cols: int):
        """创建网格画廊"""
        rows = (len(items) + cols - 1) // cols
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        # 隐藏边框
        for row in table.rows:
            for cell in row.cells:
                cell._element.get_or_add_tcPr().append(
                    parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
                )
        
        for idx, item in enumerate(items):
            row = idx // cols
            col = idx % cols
            cell = table.rows[row].cells[col]
            cell.text = ''
            
            # 添加标题
            if item.get('title'):
                title_para = cell.add_paragraph()
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_para.add_run(item['title'])
                title_run.bold = True
                title_run.font.size = Pt(11)
                title_run.font.name = '宋体'
                title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 添加图片
            if item.get('image') and os.path.exists(item['image']):
                img_para = cell.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_run = img_para.add_run()
                img_run.add_picture(item['image'], width=Inches(2.5))
            
            # 添加文字
            if item.get('text'):
                text_para = cell.add_paragraph()
                text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                text_run = text_para.add_run(item['text'])
                text_run.font.size = Pt(10)
                text_run.font.name = '宋体'
                text_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def _create_list_gallery(self, items: List[dict]):
        """创建列表画廊"""
        for item in items:
            # 添加标题
            if item.get('title'):
                title_para = self.doc.add_paragraph()
                title_run = title_para.add_run(item['title'])
                title_run.bold = True
                title_run.font.size = Pt(12)
                title_run.font.name = '宋体'
                title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 添加图片（左）和文字（右）
            if item.get('image') and os.path.exists(item['image']):
                self.image_handler.create_text_image_layout(
                    text=item.get('text', ''),
                    image_path=item['image'],
                    layout='left',
                    image_width=Inches(2)
                )
            else:
                # 只有文字
                if item.get('text'):
                    self.doc.add_paragraph(item['text'])
            
            # 添加空行
            self.doc.add_paragraph()


# 便捷函数
def insert_image_to_doc(docx_path: str, image_path: str, 
                        output_path: str = None,
                        width: Inches = None,
                        caption: str = None) -> str:
    """
    向 Word 文档插入图片
    
    Args:
        docx_path: Word 文档路径
        image_path: 图片路径
        output_path: 输出路径（可选）
        width: 图片宽度
        caption: 图片标题
    
    Returns:
        输出文件路径
    """
    doc = Document(docx_path)
    handler = ImageHandler(doc)
    
    handler.insert_image(image_path, width=width, caption=caption)
    
    if output_path is None:
        output_path = docx_path
    
    doc.save(output_path)
    return output_path


def create_image_grid(docx_path: str, image_paths: List[str],
                     cols: int = 2, width: Inches = None) -> str:
    """
    创建图片网格
    
    Args:
        docx_path: 输出 Word 文档路径
        image_paths: 图片路径列表
        cols: 列数
        width: 单张图片宽度
    
    Returns:
        输出文件路径
    """
    doc = Document()
    handler = ImageHandler(doc)
    
    handler.insert_images_grid(image_paths, cols=cols, width=width)
    
    doc.save(docx_path)
    return docx_path


if __name__ == '__main__':
    # 测试代码
    print("图片处理模块已加载")
    print("支持的环绕方式:", [attr for attr in dir(ImageWrapStyle) if not attr.startswith('_')])
