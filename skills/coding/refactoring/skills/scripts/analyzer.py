#!/usr/bin/env python3
"""格式诊断模块 v2 - 修复列表项误报"""

import re
import sys
from collections import defaultdict
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 不需要首行缩进的模式（只有主标题不需要缩进）
NO_INDENT_PATTERNS = [
    r'^附件[：:]',                      # 附件：
    r'^联系人[：:]',                    # 联系人：
    r'^抄送[：:]',                      # 抄送：
    r'^主送[：:]',                      # 主送：
]


def is_no_indent_para(text, alignment):
    """检查是否是不需要首行缩进的段落（只有居中的主标题不需要）"""
    # 居中的短文本（主标题）
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return True
    # 特殊行
    for pattern in NO_INDENT_PATTERNS:
        if re.match(pattern, text.strip()):
            return True
    return False


def analyze_punctuation(doc):
    """分析标点符号问题"""
    issues = []
    
    patterns = [
        ('英文括号', r'[\(\)]'),
        ('英文引号', r'["\']'),
        ('英文冒号', r'(?<=[^\d\s]):(?=[^\d/\\])'),
        ('英文逗号', r'(?<=[^\d]),(?=[^\d])'),
        ('英文分号', r';'),
        ('英文问号', r'\?'),
        ('英文叹号', r'!'),
    ]
    
    # 省略号：2个及以上连续的点（不是省略号格式）
    ellipsis_pattern = r'\.{2,}'
    # 破折号：连续的-
    dash_pattern = r'--+'
    # 句末英文句号：中文后面的单独句点（不是省略号的一部分）
    # 使用负向前瞻确保不是连续句点的一部分
    period_pattern = r'(?<=[\u4e00-\u9fff])\.(?!\.)'
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        
        # 只在包含中文的段落中检查
        if not re.search(r'[\u4e00-\u9fff]', text):
            continue
        
        for name, pattern in patterns:
            for match in re.finditer(pattern, text):
                issues.append({
                    'para': i + 1,
                    'type': name,
                    'char': match.group(),
                })
        
        # 检查省略号（连续多个点）
        for match in re.finditer(ellipsis_pattern, text):
            issues.append({'para': i + 1, 'type': '不规范省略号', 'char': match.group()})
        
        # 检查破折号
        for match in re.finditer(dash_pattern, text):
            issues.append({'para': i + 1, 'type': '不规范破折号', 'char': match.group()})
        
        # 检查句末英文句号（中文后面的单独句点）
        for match in re.finditer(period_pattern, text):
            issues.append({'para': i + 1, 'type': '英文句号', 'char': match.group()})
    
    return issues


def analyze_numbering(doc):
    """分析序号问题"""
    issues = []
    
    numbering_patterns = {
        'chinese_1': r'^[一二三四五六七八九十]+、',
        'chinese_2': r'^（[一二三四五六七八九十]+）',
        'arabic_dot': r'^\d+\.',
        'arabic_comma': r'^\d+、',
        'arabic_paren': r'^\d+[）\)]',
        'arabic_paren_full': r'^（\d+）',
    }
    
    found_styles = defaultdict(list)
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        for style_name, pattern in numbering_patterns.items():
            if re.match(pattern, text):
                found_styles[style_name].append(i + 1)
                break
    
    # 检查阿拉伯数字序号风格是否统一
    arabic_styles = [k for k in found_styles if k.startswith('arabic')]
    if len(arabic_styles) > 1:
        issues.append({
            'type': '序号格式不统一',
            'detail': f"同时存在: {', '.join(arabic_styles)}",
        })
    
    return issues


def analyze_paragraph_format(doc):
    """分析段落格式问题"""
    issues = []
    
    indent_issues = []
    line_spacing_values = defaultdict(list)
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # 跳过空段落和短段落（可能是标题）
        if not text or len(text) < 10:
            continue
        
        alignment = para.paragraph_format.alignment
        
        # 跳过不需要缩进的段落（居中的主标题等）
        if is_no_indent_para(text, alignment):
            continue
        
        pf = para.paragraph_format
        
        # 检查首行缩进
        indent = pf.first_line_indent
        if indent is None or indent == Pt(0) or (hasattr(indent, 'pt') and indent.pt == 0):
            indent_issues.append(i + 1)
        
        # 记录行距
        if pf.line_spacing is not None:
            line_spacing_values[str(pf.line_spacing)].append(i + 1)
    
    if indent_issues:
        issues.append({
            'type': '缺少首行缩进',
            'paras': indent_issues
        })
    
    if len(line_spacing_values) > 1:
        issues.append({
            'type': '行距不统一',
            'detail': f"存在 {len(line_spacing_values)} 种不同行距",
        })
    
    return issues


def analyze_font(doc):
    """分析字体问题"""
    issues = []
    
    font_names = set()
    font_sizes = set()
    
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        for run in para.runs:
            if run.font.name:
                font_names.add(run.font.name)
            if run.font.size:
                font_sizes.add(str(run.font.size))
    
    if len(font_names) > 4:
        issues.append({
            'type': '字体种类过多',
            'detail': f"检测到 {len(font_names)} 种字体: {', '.join(list(font_names)[:5])}..."
        })
    
    if len(font_sizes) > 4:
        issues.append({
            'type': '字号不统一',
            'detail': f"检测到 {len(font_sizes)} 种字号"
        })
    
    return issues


def print_report(results):
    """打印诊断报告"""
    print('=' * 50)
    print('           格式诊断报告')
    print('=' * 50)
    print()
    
    total = 0
    
    # 标点问题
    punct = results['punctuation']
    if punct:
        by_type = defaultdict(list)
        for issue in punct:
            by_type[issue['type']].append(issue['para'])
        
        print(f"【标点问题】共 {len(punct)} 处")
        for issue_type, paras in by_type.items():
            unique_paras = sorted(set(paras))
            if len(unique_paras) > 5:
                para_str = f"第{unique_paras[0]}、{unique_paras[1]}...{unique_paras[-1]}段"
            else:
                para_str = f"第{', '.join(map(str, unique_paras))}段"
            print(f"  - {issue_type}: {para_str}")
        print()
        total += len(punct)
    
    # 序号问题
    num = results['numbering']
    if num:
        print(f"【序号问题】共 {len(num)} 处")
        for issue in num:
            print(f"  - {issue['type']}: {issue.get('detail', '')}")
        print()
        total += len(num)
    
    # 段落问题
    para = results['paragraph']
    if para:
        print(f"【段落问题】共 {len(para)} 处")
        for issue in para:
            if issue['type'] == '缺少首行缩进':
                paras = issue['paras']
                if len(paras) > 5:
                    para_str = f"第{paras[0]}、{paras[1]}...等{len(paras)}段"
                else:
                    para_str = f"第{', '.join(map(str, paras))}段"
                print(f"  - {issue['type']}: {para_str}")
            else:
                print(f"  - {issue['type']}: {issue.get('detail', '')}")
        print()
        total += len(para)
    
    # 字体问题
    font = results['font']
    if font:
        print(f"【字体问题】共 {len(font)} 处")
        for issue in font:
            print(f"  - {issue['type']}: {issue.get('detail', '')}")
        print()
        total += len(font)
    
    # 总结
    print('-' * 50)
    if total == 0:
        print('OK 未发现明显格式问题')
    else:
        print(f'共发现 {total} 处格式问题')
        print()
        print('建议：')
        if results['punctuation']:
            print('  - 运行 punctuation.py 修复标点问题')
        if results['paragraph'] or results['font']:
            print('  - 运行 formatter.py 统一段落和字体格式')
    print()


def main():
    if len(sys.argv) < 2:
        print('Usage: python analyzer.py input.docx [--json]')
        sys.exit(1)
    
    input_file = sys.argv[1]
    print(f'Analyzing: {input_file}')
    print()
    
    doc = Document(input_file)
    
    results = {
        'punctuation': analyze_punctuation(doc),
        'numbering': analyze_numbering(doc),
        'paragraph': analyze_paragraph_format(doc),
        'font': analyze_font(doc)
    }
    
    if '--json' in sys.argv:
        import json
        print(json.dumps(results, ensure_ascii=False, indent=2))
    else:
        print_report(results)


if __name__ == '__main__':
    main()
