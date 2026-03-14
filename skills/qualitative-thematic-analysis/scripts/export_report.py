#!/usr/bin/env python3
"""
export_report.py — 将主题分析结果导出为 DOCX 报告
用法：python3 export_report.py <分析结果JSON路径> <输出DOCX路径>

JSON格式：
{
  "title": "研究报告标题",
  "research_question": "研究问题",
  "data_overview": "资料概述",
  "themes": [
    {
      "name": "主题名称",
      "definition": "主题定义",
      "sub_themes": [
        {
          "name": "子主题名称",
          "description": "描述",
          "quotes": ["引文1", "引文2"]
        }
      ]
    }
  ],
  "summary": "分析发现总结",
  "limitations": "局限与反思"
}
"""
import sys
import json
import os
from datetime import datetime

def export_report(data, output_path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("需要安装 python-docx：pip install python-docx")
        sys.exit(1)

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    # 标题
    title = doc.add_heading(data.get("title", "质性研究主题分析报告"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 元信息
    doc.add_paragraph(f"分析日期：{datetime.now().strftime('%Y年%m月%d日')}")
    doc.add_paragraph(f"研究问题：{data.get('research_question', '')}")
    doc.add_paragraph()

    # 一、资料概述
    doc.add_heading("一、资料概述", 1)
    doc.add_paragraph(data.get("data_overview", ""))

    # 二、主题分析结果
    doc.add_heading("二、主题分析结果", 1)

    themes = data.get("themes", [])
    for i, theme in enumerate(themes, 1):
        doc.add_heading(f"主题{i}：{theme['name']}", 2)

        # 主题定义
        definition_para = doc.add_paragraph()
        run = definition_para.add_run("主题定义：")
        run.bold = True
        definition_para.add_run(theme.get("definition", ""))

        # 子主题
        for sub in theme.get("sub_themes", []):
            doc.add_heading(sub["name"], 3)
            doc.add_paragraph(sub.get("description", ""))

            # 引文
            for quote in sub.get("quotes", []):
                q_para = doc.add_paragraph(style="Quote")
                q_para.add_run(f""{quote}"")

    # 三、分析发现小结
    doc.add_heading("三、分析发现小结", 1)
    doc.add_paragraph(data.get("summary", ""))

    # 四、研究局限与反思
    doc.add_heading("四、研究局限与反思", 1)
    doc.add_paragraph(data.get("limitations", ""))

    doc.save(output_path)
    print(f"✅ 分析报告已保存：{output_path}")

def main():
    if len(sys.argv) < 3:
        print("用法：python3 export_report.py <分析结果JSON路径> <输出DOCX路径>")
        sys.exit(1)

    json_path = os.path.expanduser(sys.argv[1])
    output_path = os.path.expanduser(sys.argv[2])

    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    export_report(data, output_path)

if __name__ == "__main__":
    main()
