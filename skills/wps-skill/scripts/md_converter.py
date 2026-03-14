#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown 到 Word 文档转换器
支持完整的 Markdown 语法转换，包括标题、列表、表格、代码块、图片等
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class MarkdownToWordConverter:
    """Markdown 到 Word 转换器"""
    
    def __init__(self):
        self.doc = None
        self.default_font = '宋体'
        self.default_font_size = Pt(12)
        self.code_font = 'Consolas'
        
    def convert(self, md_content: str, output_path: str, title: str = None):
        """
        将 Markdown 内容转换为 Word 文档
        
        Args:
            md_content: Markdown 文本内容
            output_path: 输出 Word 文件路径
            title: 文档标题（可选）
        """
        self.doc = Document()
        
        # 设置默认字体
        self._set_default_font()
        
        # 添加标题
        if title:
            self._add_title(title)
        
        # 解析并转换 Markdown
        self._parse_markdown(md_content)
        
        # 保存文档
        self.doc.save(output_path)
        return output_path
    
    def convert_file(self, md_file: str, output_path: str = None, title: str = None):
        """
        将 Markdown 文件转换为 Word 文档
        
        Args:
            md_file: Markdown 文件路径
            output_path: 输出 Word 文件路径（可选，默认与输入文件同名）
            title: 文档标题（可选）
        """
        # 读取 Markdown 文件
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 确定输出路径
        if output_path is None:
            base_name = os.path.splitext(md_file)[0]
            output_path = base_name + '.docx'
        
        return self.convert(md_content, output_path, title)
    
    def _set_default_font(self):
        """设置文档默认字体"""
        # 设置默认段落字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.default_font
        font.size = self.default_font_size
        
        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
    
    def _add_title(self, title: str):
        """添加文档标题"""
        heading = self.doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置标题字体
        for run in heading.runs:
            run.font.name = self.default_font
            run.font.size = Pt(18)
            run.font.bold = True
            run.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
    
    def _parse_markdown(self, md_content: str):
        """解析 Markdown 内容"""
        lines = md_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # 空行
            if not stripped:
                i += 1
                continue
            
            # 代码块
            if stripped.startswith('```'):
                i = self._handle_code_block(lines, i)
                continue
            
            # 表格
            if '|' in stripped and i + 1 < len(lines) and '|' in lines[i + 1] and '---' in lines[i + 1]:
                i = self._handle_table(lines, i)
                continue
            
            # 标题
            if stripped.startswith('#'):
                self._handle_heading(stripped)
                i += 1
                continue
            
            # 无序列表
            if stripped.startswith(('- ', '* ', '+ ')):
                i = self._handle_unordered_list(lines, i)
                continue
            
            # 有序列表
            if re.match(r'^\d+\.', stripped):
                i = self._handle_ordered_list(lines, i)
                continue
            
            # 引用块
            if stripped.startswith('>'):
                i = self._handle_blockquote(lines, i)
                continue
            
            # 水平分割线
            if stripped in ['---', '***', '___']:
                self._add_horizontal_line()
                i += 1
                continue
            
            # 普通段落
            i = self._handle_paragraph(lines, i)
    
    def _handle_heading(self, line: str):
        """处理标题"""
        match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if match:
            level = len(match.group(1))
            text = match.group(2)
            
            heading = self.doc.add_heading(level=level)
            self._add_formatted_text(heading, text)
            
            # 设置字体
            for run in heading.runs:
                run.font.name = self.default_font
                run.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
    
    def _handle_code_block(self, lines: list, start_idx: int) -> int:
        """处理代码块"""
        language = lines[start_idx].strip()[3:].strip()
        code_lines = []
        i = start_idx + 1
        
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
        
        code_content = '\n'.join(code_lines)
        
        # 添加代码块段落
        paragraph = self.doc.add_paragraph()
        paragraph.style = 'Normal'
        
        # 添加代码内容
        run = paragraph.add_run(code_content)
        run.font.name = self.code_font
        run.font.size = Pt(10)
        
        # 设置代码块背景色（浅灰色）
        self._set_paragraph_shading(paragraph, 'F5F5F5')
        
        return i + 1
    
    def _handle_table(self, lines: list, start_idx: int) -> int:
        """处理表格"""
        # 解析表头
        header_line = lines[start_idx].strip()
        headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]
        
        # 跳过分隔行
        i = start_idx + 2
        
        # 收集数据行
        rows = []
        while i < len(lines) and '|' in lines[i]:
            row_line = lines[i].strip()
            cells = [cell.strip() for cell in row_line.split('|') if cell.strip()]
            if cells:
                rows.append(cells)
            i += 1
        
        # 创建表格
        if headers:
            table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = 'Table Grid'
            
            # 填充表头
            hdr_cells = table.rows[0].cells
            for j, header in enumerate(headers):
                hdr_cells[j].text = header
                # 设置表头样式
                for paragraph in hdr_cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.name = self.default_font
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
            
            # 填充数据
            for row_idx, row_data in enumerate(rows):
                row_cells = table.rows[row_idx + 1].cells
                for j, cell_text in enumerate(row_data):
                    if j < len(row_cells):
                        row_cells[j].text = cell_text
                        # 设置字体
                        for paragraph in row_cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.font.name = self.default_font
                                run.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
        
        return i
    
    def _handle_unordered_list(self, lines: list, start_idx: int) -> int:
        """处理无序列表"""
        i = start_idx
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            if not stripped.startswith(('- ', '* ', '+ ')):
                break
            
            # 提取列表项内容
            content = stripped[2:]
            
            # 添加列表项
            paragraph = self.doc.add_paragraph(content, style='List Bullet')
            self._add_formatted_text(paragraph, content)
            
            # 设置字体
            for run in paragraph.runs:
                run.font.name = self.default_font
                run.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
            
            i += 1
        
        return i
    
    def _handle_ordered_list(self, lines: list, start_idx: int) -> int:
        """处理有序列表"""
        i = start_idx
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
            if not match:
                break
            
            content = match.group(2)
            
            # 添加列表项
            paragraph = self.doc.add_paragraph(content, style='List Number')
            self._add_formatted_text(paragraph, content)
            
            # 设置字体
            for run in paragraph.runs:
                run.font.name = self.default_font
                run.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
            
            i += 1
        
        return i
    
    def _handle_blockquote(self, lines: list, start_idx: int) -> int:
        """处理引用块"""
        i = start_idx
        quote_lines = []
        
        while i < len(lines) and lines[i].strip().startswith('>'):
            line = lines[i].strip()
            if line.startswith('> '):
                quote_lines.append(line[2:])
            elif line == '>':
                quote_lines.append('')
            else:
                quote_lines.append(line[1:].strip())
            i += 1
        
        # 添加引用段落
        quote_text = '\n'.join(quote_lines)
        paragraph = self.doc.add_paragraph()
        
        # 添加引用样式（左侧缩进和边框）
        run = paragraph.add_run(quote_text)
        run.font.italic = True
        run.font.name = self.default_font
        run.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
        
        # 设置段落缩进
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.right_indent = Inches(0.5)
        
        return i
    
    def _handle_paragraph(self, lines: list, start_idx: int) -> int:
        """处理普通段落"""
        i = start_idx
        paragraph_lines = []
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # 遇到空行或特殊格式结束段落
            if not stripped or stripped.startswith(('#', '- ', '* ', '+ ', '>', '|', '```')):
                break
            
            # 检查是否是列表或标题
            if re.match(r'^\d+\.', stripped):
                break
            
            paragraph_lines.append(line)
            i += 1
        
        if paragraph_lines:
            paragraph_text = ' '.join(paragraph_lines)
            paragraph = self.doc.add_paragraph()
            self._add_formatted_text(paragraph, paragraph_text)
            
            # 设置字体
            for run in paragraph.runs:
                run.font.name = self.default_font
                run.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
        
        return i
    
    def _add_formatted_text(self, paragraph, text: str):
        """添加带格式的文本（支持粗体、斜体、行内代码、链接）"""
        # 清除默认文本
        paragraph.clear()
        
        # 解析行内格式
        patterns = [
            (r'\*\*\*([^*]+)\*\*\*', 'bold_italic'),  # 粗斜体
            (r'\*\*([^*]+)\*\*', 'bold'),             # 粗体
            (r'\*([^*]+)\*', 'italic'),               # 斜体
            (r'`([^`]+)`', 'code'),                    # 行内代码
            (r'\[([^\]]+)\]\(([^)]+)\)', 'link'),     # 链接
        ]
        
        # 简单的格式处理
        remaining = text
        
        while remaining:
            # 查找最近的格式标记
            earliest_match = None
            earliest_pattern = None
            earliest_pos = len(remaining)
            
            for pattern, style in patterns:
                match = re.search(pattern, remaining)
                if match and match.start() < earliest_pos:
                    earliest_pos = match.start()
                    earliest_match = match
                    earliest_pattern = style
            
            if earliest_match:
                # 添加格式标记前的普通文本
                if earliest_pos > 0:
                    run = paragraph.add_run(remaining[:earliest_pos])
                    run.font.name = self.default_font
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
                
                # 处理格式文本
                if earliest_pattern == 'bold':
                    run = paragraph.add_run(earliest_match.group(1))
                    run.bold = True
                    run.font.name = self.default_font
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
                elif earliest_pattern == 'italic':
                    run = paragraph.add_run(earliest_match.group(1))
                    run.italic = True
                    run.font.name = self.default_font
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
                elif earliest_pattern == 'bold_italic':
                    run = paragraph.add_run(earliest_match.group(1))
                    run.bold = True
                    run.italic = True
                    run.font.name = self.default_font
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
                elif earliest_pattern == 'code':
                    run = paragraph.add_run(earliest_match.group(1))
                    run.font.name = self.code_font
                    run.font.size = Pt(10)
                elif earliest_pattern == 'link':
                    run = paragraph.add_run(earliest_match.group(1))
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.underline = True
                    run.font.name = self.default_font
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
                
                remaining = remaining[earliest_match.end():]
            else:
                # 没有更多格式标记
                run = paragraph.add_run(remaining)
                run.font.name = self.default_font
                run.element.rPr.rFonts.set(qn('w:eastAsia'), self.default_font)
                break
    
    def _set_paragraph_shading(self, paragraph, color: str):
        """设置段落背景色"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        paragraph.paragraph_format.element.get_or_add_pPr().append(shading_elm)
    
    def _add_horizontal_line(self):
        """添加水平分割线"""
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.border_bottom = True
        run = paragraph.add_run('_' * 50)
        run.font.color.rgb = RGBColor(200, 200, 200)


class WordToMarkdownConverter:
    """Word 到 Markdown 转换器"""
    
    def __init__(self):
        self.md_content = []
    
    def convert(self, docx_path: str, output_path: str = None) -> str:
        """
        将 Word 文档转换为 Markdown
        
        Args:
            docx_path: Word 文件路径
            output_path: 输出 Markdown 文件路径（可选）
        
        Returns:
            Markdown 文本内容
        """
        doc = Document(docx_path)
        self.md_content = []
        
        for paragraph in doc.paragraphs:
            self._convert_paragraph(paragraph)
        
        # 处理表格
        for table in doc.tables:
            self._convert_table(table)
        
        md_text = '\n\n'.join(self.md_content)
        
        # 保存到文件
        if output_path:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_text)
        
        return md_text
    
    def _convert_paragraph(self, paragraph):
        """转换段落"""
        text = paragraph.text.strip()
        if not text:
            return
        
        # 根据样式判断类型
        style_name = paragraph.style.name if paragraph.style else 'Normal'
        
        if style_name.startswith('Heading'):
            level = int(style_name.replace('Heading ', '')) if ' ' in style_name else 1
            self.md_content.append(f'{"#" * level} {text}')
        elif 'List Bullet' in style_name:
            self.md_content.append(f'- {text}')
        elif 'List Number' in style_name:
            self.md_content.append(f'1. {text}')
        else:
            # 普通段落，检查是否有格式
            md_text = self._extract_inline_format(paragraph)
            self.md_content.append(md_text)
    
    def _extract_inline_format(self, paragraph) -> str:
        """提取行内格式"""
        parts = []
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue
            
            # 应用格式标记
            if run.bold and run.italic:
                text = f'***{text}***'
            elif run.bold:
                text = f'**{text}**'
            elif run.italic:
                text = f'*{text}*'
            
            parts.append(text)
        
        return ''.join(parts)
    
    def _convert_table(self, table):
        """转换表格"""
        if not table.rows:
            return
        
        rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(row_data)
        
        if rows:
            # 表头
            header = '| ' + ' | '.join(rows[0]) + ' |'
            self.md_content.append(header)
            
            # 分隔符
            separator = '|' + '|'.join([' --- ' for _ in rows[0]]) + '|'
            self.md_content.append(separator)
            
            # 数据行
            for row in rows[1:]:
                row_text = '| ' + ' | '.join(row) + ' |'
                self.md_content.append(row_text)


# 便捷函数
def md_to_word(md_file: str, output_path: str = None, title: str = None) -> str:
    """Markdown 文件转 Word"""
    converter = MarkdownToWordConverter()
    return converter.convert_file(md_file, output_path, title)


def word_to_md(docx_file: str, output_path: str = None) -> str:
    """Word 文件转 Markdown"""
    converter = WordToMarkdownConverter()
    return converter.convert(docx_file, output_path)


if __name__ == '__main__':
    # 测试代码
    test_md = """# 测试文档

这是**粗体**和*斜体*的测试。

## 列表演示

- 项目1
- 项目2
- 项目3

## 表格演示

| 姓名 | 年龄 | 城市 |
|------|------|------|
| 张三 | 25 | 北京 |
| 李四 | 30 | 上海 |

## 代码块

```python
print("Hello World")
```

## 引用

> 这是一段引用文本。

---

普通段落文本。
"""
    
    converter = MarkdownToWordConverter()
    output = '/tmp/test_output.docx'
    converter.convert(test_md, output, '测试文档')
    print(f'转换完成: {output}')
