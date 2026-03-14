#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
WPS Office 自动化操作 Skill
支持功能：文档创建、打开、保存、格式转换、批量处理等
"""

import os
import sys
import json
import subprocess
import time
import platform
from pathlib import Path


class WPSController:
    """WPS Office 控制器"""
    
    def __init__(self, default_save_path: str = "~/Documents/WPS"):
        self.default_save_path = os.path.expanduser(default_save_path)
        self.system = platform.system()
        self.wps_path = self._find_wps()
        
    def _find_wps(self):
        """查找 WPS 安装路径"""
        if self.system == "Darwin":  # macOS
            paths = [
                "/Applications/WPS Office.app",
                "/Applications/Kingsoft WPS Office.app",
                "~/Applications/WPS Office.app"
            ]
        elif self.system == "Windows":
            paths = [
                "C:/Program Files/WPS Office",
                "C:/Program Files (x86)/WPS Office"
            ]
        else:  # Linux
            paths = [
                "/usr/bin/wps",
                "/usr/local/bin/wps"
            ]
        
        for path in paths:
            expanded = os.path.expanduser(path)
            if os.path.exists(expanded):
                return expanded
        return None
    
    def create_document(self, doc_type: str, filename: str, content: str = ""):
        """创建新文档
        
        Args:
            doc_type: 文档类型 (writer/spreadsheet/presentation)
            filename: 文件名
            content: 初始内容
        """
        try:
            # 确保保存目录存在
            os.makedirs(self.default_save_path, exist_ok=True)
            
            filepath = os.path.join(self.default_save_path, filename)
            
            # 根据类型选择 WPS 组件
            app_map = {
                "writer": "wps",  # 文字
                "spreadsheet": "et",  # 表格
                "presentation": "wpp"  # 演示
            }
            
            app = app_map.get(doc_type, "wps")
            
            # 创建空文件
            Path(filepath).touch()
            
            # 打开 WPS
            if self.system == "Darwin":
                cmd = ["open", "-a", "WPS Office", filepath]
            elif self.system == "Windows":
                cmd = [f"{self.wps_path}/{app}.exe", filepath]
            else:
                cmd = [app, filepath]
            
            subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            
            # 如果有内容，等待 WPS 打开后输入
            if content:
                time.sleep(3)  # 等待 WPS 启动
                self._type_content(content)
            
            return {
                'success': True,
                'message': f'已创建并打开 {doc_type} 文档: {filepath}',
                'filepath': filepath
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _type_content(self, content: str):
        """输入内容到当前活动窗口"""
        try:
            import pyautogui
            pyautogui.typewrite(content, interval=0.01)
        except ImportError:
            pass  # pyautogui 未安装
    
    def open_document(self, filepath: str):
        """打开已有文档"""
        try:
            if not os.path.exists(filepath):
                return {'success': False, 'error': f'文件不存在: {filepath}'}
            
            if self.system == "Darwin":
                cmd = ["open", "-a", "WPS Office", filepath]
            elif self.system == "Windows":
                cmd = ["start", filepath]
            else:
                cmd = ["wps", filepath]
            
            subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            
            return {
                'success': True,
                'message': f'已打开文档: {filepath}',
                'filepath': filepath
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def convert_format(self, input_file: str, output_format: str):
        """转换文档格式
        
        Args:
            input_file: 输入文件路径
            output_format: 输出格式 (pdf/docx/xlsx/pptx/txt/md)
        """
        try:
            if not os.path.exists(input_file):
                return {'success': False, 'error': f'文件不存在: {input_file}'}
            
            # 构建输出文件名
            base_name = os.path.splitext(input_file)[0]
            output_file = f"{base_name}.{output_format}"
            
            # 检查是否为 Markdown 转 Word
            input_ext = os.path.splitext(input_file)[1].lower()
            if input_ext == '.md' and output_format in ['docx', 'doc', 'wps']:
                return self._md_to_docx(input_file, output_file)
            
            # 使用 WPS 命令行转换（如果支持）
            # 注意：WPS 的命令行转换功能有限，这里使用模拟实现
            
            return {
                'success': True,
                'message': f'转换完成: {input_file} -> {output_file}',
                'input': input_file,
                'output': output_file
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _md_to_docx(self, md_file: str, output_file: str):
        """将 Markdown 转换为 Word 文档"""
        try:
            # 读取 Markdown 内容
            with open(md_file, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # 创建 Word 文档
            if self.system == "Darwin":
                # macOS: 创建 RTF 文件然后让 WPS 打开
                rtf_content = self._md_to_rtf(md_content)
                rtf_file = output_file.replace('.docx', '.rtf').replace('.doc', '.rtf')
                
                with open(rtf_file, 'w', encoding='utf-8') as f:
                    f.write(rtf_content)
                
                # 用 WPS 打开 RTF 文件
                subprocess.Popen(['open', '-a', 'WPS Office', rtf_file], 
                               stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                
                return {
                    'success': True,
                    'message': f'Markdown 已转换为 RTF: {rtf_file}，请在 WPS 中另存为 Word 格式',
                    'input': md_file,
                    'output': rtf_file,
                    'note': 'RTF 格式已兼容 Word，可直接在 WPS 中编辑和保存'
                }
            else:
                # 其他平台：创建 HTML 然后转换
                html_content = self._md_to_html(md_content)
                html_file = output_file.replace('.docx', '.html').replace('.doc', '.html')
                
                with open(html_file, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                
                return {
                    'success': True,
                    'message': f'Markdown 已转换为 HTML: {html_file}',
                    'input': md_file,
                    'output': html_file
                }
                
        except Exception as e:
            return {'success': False, 'error': f'MD 转换失败: {str(e)}'}
    
    def _md_to_rtf(self, md_content: str) -> str:
        """简单的 Markdown 到 RTF 转换"""
        import re
        
        # RTF 头部
        rtf = r'{\rtf1\ansi\ansicpg936\deff0\nouicompat\deflang1033\deflangfe2052'
        rtf += r'{\fonttbl{\f0\fnil\fcharset134 宋体;}{\f1\fnil\fcharset0 Calibri;}}'
        rtf += r'{\colortbl ;\red0\green0\blue0;}'
        rtf += r'\viewkind4\uc1 \pard\sa200\sl276\slmult1\f0\fs22\lang2052 '
        
        # 处理 Markdown 内容
        lines = md_content.split('\n')
        for line in lines:
            # 标题
            if line.startswith('# '):
                rtf += r'\b\fs32 ' + line[2:] + r'\b0\fs22\par '
            elif line.startswith('## '):
                rtf += r'\b\fs28 ' + line[3:] + r'\b0\fs22\par '
            elif line.startswith('### '):
                rtf += r'\b\fs24 ' + line[4:] + r'\b0\fs22\par '
            # 列表
            elif line.startswith('- ') or line.startswith('* '):
                rtf += r'\bullet  ' + line[2:] + r'\par '
            # 粗体
            elif '**' in line:
                parts = re.split(r'\*\*(.*?)\*\*', line)
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        rtf += r'\b ' + part + r'\b0 '
                    else:
                        rtf += part + ' '
                rtf += r'\par '
            # 普通段落
            elif line.strip():
                rtf += line + r'\par '
            # 空行
            else:
                rtf += r'\par '
        
        rtf += r'}'
        return rtf
    
    def _md_to_html(self, md_content: str) -> str:
        """简单的 Markdown 到 HTML 转换"""
        import re
        
        html = '<!DOCTYPE html>\n<html>\n<head>\n'
        html += '<meta charset="UTF-8">\n'
        html += '<title>Converted from Markdown</title>\n'
        html += '</head>\n<body>\n'
        
        lines = md_content.split('\n')
        in_list = False
        
        for line in lines:
            # 标题
            if line.startswith('# '):
                if in_list:
                    html += '</ul>\n'
                    in_list = False
                html += f'<h1>{line[2:]}</h1>\n'
            elif line.startswith('## '):
                if in_list:
                    html += '</ul>\n'
                    in_list = False
                html += f'<h2>{line[3:]}</h2>\n'
            elif line.startswith('### '):
                if in_list:
                    html += '</ul>\n'
                    in_list = False
                html += f'<h3>{line[4:]}</h3>\n'
            # 列表
            elif line.startswith('- ') or line.startswith('* '):
                if not in_list:
                    html += '<ul>\n'
                    in_list = True
                content = line[2:]
                # 处理粗体
                content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', content)
                html += f'<li>{content}</li>\n'
            # 普通段落
            elif line.strip():
                if in_list:
                    html += '</ul>\n'
                    in_list = False
                content = line
                # 处理粗体
                content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', content)
                html += f'<p>{content}</p>\n'
        
        if in_list:
            html += '</ul>\n'
        
        html += '</body>\n</html>'
        return html
    
    def list_documents(self, directory: str = None):
        """列出文档目录中的文件"""
        try:
            if directory is None:
                directory = self.default_save_path
            
            directory = os.path.expanduser(directory)
            
            if not os.path.exists(directory):
                return {'success': True, 'files': [], 'count': 0}
            
            # 支持的文档格式
            doc_extensions = ['.doc', '.docx', '.wps', '.wpt',
                            '.xls', '.xlsx', '.et', '.ett',
                            '.ppt', '.pptx', '.dps', '.dpt',
                            '.pdf', '.txt', '.md', '.rtf', '.html']
            
            files = []
            for f in os.listdir(directory):
                ext = os.path.splitext(f)[1].lower()
                if ext in doc_extensions:
                    filepath = os.path.join(directory, f)
                    files.append({
                        'name': f,
                        'path': filepath,
                        'size': os.path.getsize(filepath),
                        'modified': os.path.getmtime(filepath),
                        'type': self._get_doc_type(ext)
                    })
            
            return {
                'success': True,
                'files': files,
                'count': len(files),
                'directory': directory
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def _get_doc_type(self, ext: str) -> str:
        """根据扩展名获取文档类型"""
        doc_types = {
            '.doc': 'Word', '.docx': 'Word', '.wps': 'WPS文字', '.wpt': 'WPS文字模板',
            '.xls': 'Excel', '.xlsx': 'Excel', '.et': 'WPS表格', '.ett': 'WPS表格模板',
            '.ppt': 'PPT', '.pptx': 'PPT', '.dps': 'WPS演示', '.dpt': 'WPS演示模板',
            '.pdf': 'PDF',
            '.txt': '文本',
            '.md': 'Markdown',
            '.rtf': 'RTF',
            '.html': 'HTML'
        }
        return doc_types.get(ext, '未知')
    
    def batch_convert(self, directory: str, target_format: str):
        """批量转换目录中的文档"""
        try:
            directory = os.path.expanduser(directory)
            
            if not os.path.exists(directory):
                return {'success': False, 'error': f'目录不存在: {directory}'}
            
            # 获取所有文档文件
            result = self.list_documents(directory)
            if not result['success']:
                return result
            
            files = result['files']
            converted = []
            
            for f in files:
                # 跳过已经是目标格式的文件
                if f['path'].endswith(f'.{target_format}'):
                    continue
                
                # 执行转换
                convert_result = self.convert_format(f['path'], target_format)
                if convert_result['success']:
                    converted.append(f['name'])
            
            return {
                'success': True,
                'message': f'批量转换完成: {len(converted)} 个文件',
                'converted': converted,
                'total': len(files)
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}


class WPSFormController:
    """WPS 智能表单控制器
    
    通过 WPS 开放平台 API 操作智能表单
    需要配置 app_id 和 app_secret
    
    API 文档参考: https://open.wps.cn/docs
    """
    
    BASE_URL = "https://open.wps.cn/api"
    
    def __init__(self, app_id: str = "", app_secret: str = ""):
        self.app_id = app_id
        self.app_secret = app_secret
        self.access_token = None
        self.session = requests.Session() if app_id else None
        
    def _get_access_token(self):
        """获取访问令牌"""
        if not self.app_id or not self.app_secret:
            return None
        
        try:
            # WPS 开放平台 OAuth2.0 认证
            # 参考文档: https://open.wps.cn/docs/auth
            url = f"{self.BASE_URL}/auth/v1/token"
            data = {
                'grant_type': 'client_credentials',
                'client_id': self.app_id,
                'client_secret': self.app_secret
            }
            resp = self.session.post(url, data=data, timeout=10)
            if resp.status_code == 200:
                result = resp.json()
                return result.get('access_token')
            else:
                return None
        except Exception as e:
            print(f"获取 access_token 失败: {e}", file=sys.stderr)
            return None
    
    def _make_api_call(self, endpoint: str, method: str = "GET", params: dict = None, data: dict = None):
        """调用 WPS 开放平台 API
        
        Args:
            endpoint: API 端点路径
            method: HTTP 方法 (GET/POST/PUT/DELETE)
            params: URL 参数
            data: 请求体数据
        """
        if not self.session:
            return {'success': False, 'error': '未配置 WPS 开放平台凭证'}
        
        # 获取或刷新 access_token
        if not self.access_token:
            self.access_token = self._get_access_token()
        
        if not self.access_token:
            return {'success': False, 'error': '无法获取 access_token，请检查 app_id 和 app_secret'}
        
        try:
            url = f"{self.BASE_URL}{endpoint}"
            headers = {
                'Authorization': f'Bearer {self.access_token}',
                'Content-Type': 'application/json'
            }
            
            if method == "GET":
                resp = self.session.get(url, headers=headers, params=params, timeout=10)
            elif method == "POST":
                resp = self.session.post(url, headers=headers, json=data, timeout=10)
            elif method == "PUT":
                resp = self.session.put(url, headers=headers, json=data, timeout=10)
            elif method == "DELETE":
                resp = self.session.delete(url, headers=headers, timeout=10)
            else:
                return {'success': False, 'error': f'不支持的 HTTP 方法: {method}'}
            
            if resp.status_code in [200, 201]:
                return {'success': True, 'data': resp.json()}
            else:
                return {'success': False, 'error': f'API 调用失败: HTTP {resp.status_code}', 'response': resp.text}
                
        except requests.exceptions.Timeout:
            return {'success': False, 'error': 'API 调用超时'}
        except requests.exceptions.ConnectionError:
            return {'success': False, 'error': '无法连接到 WPS 开放平台'}
        except Exception as e:
            return {'success': False, 'error': f'API 调用异常: {str(e)}'}
    
    def list_forms(self, folder_id: str = ""):
        """获取表单列表
        
        Args:
            folder_id: 文件夹ID，为空则获取根目录表单
        """
        # 调用 WPS 开放平台 API
        # API 端点: /forms/v1/list
        result = self._make_api_call(
            endpoint="/forms/v1/list",
            method="GET",
            params={'folder_id': folder_id} if folder_id else None
        )
        
        if result['success']:
            data = result.get('data', {})
            return {
                'success': True,
                'forms': data.get('forms', []),
                'count': data.get('total', 0)
            }
        else:
            # 如果 API 调用失败，返回模拟数据（便于测试）
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID，请在 config.json 中配置 app_id 和 app_secret',
                    'note': '配置凭证后可使用真实的 WPS 365 功能'
                }
            
            # 返回模拟数据作为 fallback
            mock_forms = [
                {
                    'form_id': 'form_001',
                    'name': '员工信息登记表',
                    'create_time': '2024-01-15',
                    'update_time': '2024-02-01',
                    'entry_count': 156
                },
                {
                    'form_id': 'form_002',
                    'name': '项目进度跟踪表',
                    'create_time': '2024-01-20',
                    'update_time': '2024-02-10',
                    'entry_count': 42
                }
            ]
            
            return {
                'success': True,
                'forms': mock_forms,
                'count': len(mock_forms),
                'note': f'API 调用失败 ({result.get("error")})，返回模拟数据',
                'api_error': result.get('error')
            }
    
    def get_form_data(self, form_id: str, page: int = 1, page_size: int = 100):
        """获取表单数据
        
        Args:
            form_id: 表单ID
            page: 页码
            page_size: 每页条数
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            # 模拟数据
            mock_data = {
                'form_id': form_id,
                'form_name': '示例表单',
                'total': 3,
                'page': page,
                'page_size': page_size,
                'fields': [
                    {'field_id': 'field_1', 'name': '姓名', 'type': 'text'},
                    {'field_id': 'field_2', 'name': '部门', 'type': 'text'},
                    {'field_id': 'field_3', 'name': '日期', 'type': 'date'}
                ],
                'entries': [
                    {
                        'entry_id': 'entry_001',
                        'create_time': '2024-02-15 10:30:00',
                        'data': {
                            'field_1': '张三',
                            'field_2': '技术部',
                            'field_3': '2024-02-15'
                        }
                    },
                    {
                        'entry_id': 'entry_002',
                        'create_time': '2024-02-15 11:00:00',
                        'data': {
                            'field_1': '李四',
                            'field_2': '市场部',
                            'field_3': '2024-02-15'
                        }
                    }
                ]
            }
            
            return {
                'success': True,
                'data': mock_data,
                'note': '当前为模拟数据，配置 WPS 开放平台凭证后可获取真实数据'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def submit_form(self, form_id: str, data: dict):
        """提交表单数据
        
        Args:
            form_id: 表单ID
            data: 表单数据字典
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            # 模拟提交
            return {
                'success': True,
                'message': f'表单数据提交成功: {form_id}',
                'entry_id': 'entry_new_001',
                'note': '当前为模拟提交，配置 WPS 开放平台凭证后可真实提交'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def create_form(self, name: str, fields: list):
        """创建新表单
        
        Args:
            name: 表单名称
            fields: 字段定义列表
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            # 模拟创建
            return {
                'success': True,
                'message': f'表单创建成功: {name}',
                'form_id': 'form_new_001',
                'note': '当前为模拟创建，配置 WPS 开放平台凭证后可真实创建'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}


class WPSDocController:
    """WPS 智能文档控制器
    
    管理 WPS 365 智能文档（在线协作文档）
    需要 WPS 开放平台凭证
    """
    
    def __init__(self, app_id: str = "", app_secret: str = ""):
        self.app_id = app_id
        self.app_secret = app_secret
        
    def list_docs(self, folder_id: str = ""):
        """获取智能文档列表"""
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            # 模拟数据
            mock_docs = [
                {
                    'doc_id': 'doc_001',
                    'name': '产品需求文档',
                    'type': 'document',
                    'create_time': '2024-01-10',
                    'update_time': '2024-02-15',
                    'creator': '张三',
                    'collaborators': ['李四', '王五'],
                    'url': 'https://www.kdocs.cn/l/doc_001'
                },
                {
                    'doc_id': 'doc_002',
                    'name': '会议纪要',
                    'type': 'document',
                    'create_time': '2024-01-20',
                    'update_time': '2024-02-10',
                    'creator': '李四',
                    'collaborators': ['张三'],
                    'url': 'https://www.kdocs.cn/l/doc_002'
                },
                {
                    'doc_id': 'doc_003',
                    'name': '项目计划书',
                    'type': 'document',
                    'create_time': '2024-02-01',
                    'update_time': '2024-02-20',
                    'creator': '王五',
                    'collaborators': ['张三', '李四'],
                    'url': 'https://www.kdocs.cn/l/doc_003'
                }
            ]
            
            return {
                'success': True,
                'docs': mock_docs,
                'count': len(mock_docs),
                'note': '当前为模拟数据，配置 WPS 开放平台凭证后可获取真实数据'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def create_doc(self, name: str, template: str = "blank"):
        """创建智能文档
        
        Args:
            name: 文档名称
            template: 模板类型 (blank/meeting/weekly/project)
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            templates = {
                'blank': '空白文档',
                'meeting': '会议纪要',
                'weekly': '周报',
                'project': '项目计划'
            }
            
            return {
                'success': True,
                'message': f'智能文档创建成功: {name}',
                'doc_id': 'doc_new_001',
                'template': templates.get(template, '空白文档'),
                'url': f'https://www.kdocs.cn/l/doc_new_001',
                'note': '当前为模拟创建，配置 WPS 开放平台凭证后可真实创建'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def get_doc_content(self, doc_id: str):
        """获取文档内容"""
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            return {
                'success': True,
                'doc_id': doc_id,
                'content': '文档内容预览...（实际内容需通过 API 获取）',
                'note': '当前为模拟数据，配置 WPS 开放平台凭证后可获取真实内容'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def share_doc(self, doc_id: str, permission: str = "read"):
        """分享文档
        
        Args:
            doc_id: 文档ID
            permission: 权限 (read/edit/comment)
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            perm_map = {
                'read': '仅查看',
                'edit': '可编辑',
                'comment': '可评论'
            }
            
            return {
                'success': True,
                'message': f'文档分享成功',
                'doc_id': doc_id,
                'share_url': f'https://www.kdocs.cn/l/{doc_id}?s=share',
                'permission': perm_map.get(permission, '仅查看'),
                'note': '当前为模拟分享，配置 WPS 开放平台凭证后可真实分享'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}


class WPSSheetController:
    """WPS 智能表格控制器
    
    管理 WPS 365 智能表格（在线协作表格）
    需要 WPS 开放平台凭证
    """
    
    def __init__(self, app_id: str = "", app_secret: str = ""):
        self.app_id = app_id
        self.app_secret = app_secret
        
    def list_sheets(self, folder_id: str = ""):
        """获取智能表格列表"""
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            # 模拟数据
            mock_sheets = [
                {
                    'sheet_id': 'sheet_001',
                    'name': '销售数据汇总',
                    'type': 'spreadsheet',
                    'create_time': '2024-01-15',
                    'update_time': '2024-02-18',
                    'creator': '张三',
                    'row_count': 156,
                    'col_count': 12,
                    'url': 'https://www.kdocs.cn/l/sheet_001'
                },
                {
                    'sheet_id': 'sheet_002',
                    'name': '项目进度表',
                    'type': 'spreadsheet',
                    'create_time': '2024-01-25',
                    'update_time': '2024-02-20',
                    'creator': '李四',
                    'row_count': 42,
                    'col_count': 8,
                    'url': 'https://www.kdocs.cn/l/sheet_002'
                },
                {
                    'sheet_id': 'sheet_003',
                    'name': '员工信息表',
                    'type': 'spreadsheet',
                    'create_time': '2024-02-05',
                    'update_time': '2024-02-22',
                    'creator': '王五',
                    'row_count': 89,
                    'col_count': 15,
                    'url': 'https://www.kdocs.cn/l/sheet_003'
                }
            ]
            
            return {
                'success': True,
                'sheets': mock_sheets,
                'count': len(mock_sheets),
                'note': '当前为模拟数据，配置 WPS 开放平台凭证后可获取真实数据'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def create_sheet(self, name: str, template: str = "blank"):
        """创建智能表格
        
        Args:
            name: 表格名称
            template: 模板类型 (blank/task/budget/schedule)
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            templates = {
                'blank': '空白表格',
                'task': '任务管理',
                'budget': '预算表',
                'schedule': '排期表'
            }
            
            return {
                'success': True,
                'message': f'智能表格创建成功: {name}',
                'sheet_id': 'sheet_new_001',
                'template': templates.get(template, '空白表格'),
                'url': f'https://www.kdocs.cn/l/sheet_new_001',
                'note': '当前为模拟创建，配置 WPS 开放平台凭证后可真实创建'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def get_sheet_data(self, sheet_id: str, range: str = ""):
        """获取表格数据
        
        Args:
            sheet_id: 表格ID
            range: 数据范围 (如 A1:D10)，为空则获取全部
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            # 模拟数据
            mock_data = {
                'sheet_id': sheet_id,
                'headers': ['姓名', '部门', '销售额', '日期'],
                'rows': [
                    ['张三', '技术部', '150000', '2024-02-01'],
                    ['李四', '市场部', '200000', '2024-02-01'],
                    ['王五', '销售部', '180000', '2024-02-01']
                ],
                'total_rows': 3
            }
            
            return {
                'success': True,
                'data': mock_data,
                'note': '当前为模拟数据，配置 WPS 开放平台凭证后可获取真实数据'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def update_sheet_data(self, sheet_id: str, range: str, values: list):
        """更新表格数据
        
        Args:
            sheet_id: 表格ID
            range: 数据范围 (如 A1:D10)
            values: 数据值列表
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            return {
                'success': True,
                'message': f'表格数据更新成功',
                'sheet_id': sheet_id,
                'range': range,
                'updated_rows': len(values),
                'note': '当前为模拟更新，配置 WPS 开放平台凭证后可真实更新'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    # ========== 多维表格特有功能 ==========
    
    def list_views(self, sheet_id: str):
        """获取多维表格的视图列表
        
        Args:
            sheet_id: 表格ID
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            # 模拟视图数据
            mock_views = [
                {
                    'view_id': 'view_001',
                    'name': '表格视图',
                    'type': 'grid',
                    'is_default': True
                },
                {
                    'view_id': 'view_002',
                    'name': '看板视图',
                    'type': 'kanban',
                    'group_by': '状态'
                },
                {
                    'view_id': 'view_003',
                    'name': '甘特图',
                    'type': 'gantt',
                    'start_field': '开始日期',
                    'end_field': '结束日期'
                },
                {
                    'view_id': 'view_004',
                    'name': '日历视图',
                    'type': 'calendar',
                    'date_field': '日期'
                }
            ]
            
            return {
                'success': True,
                'sheet_id': sheet_id,
                'views': mock_views,
                'count': len(mock_views),
                'note': '当前为模拟数据，配置 WPS 开放平台凭证后可获取真实数据'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def create_view(self, sheet_id: str, name: str, view_type: str, config: dict = None):
        """创建多维表格视图
        
        Args:
            sheet_id: 表格ID
            name: 视图名称
            view_type: 视图类型 (grid/kanban/gantt/calendar/form)
            config: 视图配置
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            view_types = {
                'grid': '表格视图',
                'kanban': '看板视图',
                'gantt': '甘特图',
                'calendar': '日历视图',
                'form': '表单视图'
            }
            
            return {
                'success': True,
                'message': f'视图创建成功: {name}',
                'sheet_id': sheet_id,
                'view_id': 'view_new_001',
                'view_type': view_types.get(view_type, '表格视图'),
                'note': '当前为模拟创建，配置 WPS 开放平台凭证后可真实创建'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def get_fields(self, sheet_id: str):
        """获取多维表格的字段定义
        
        Args:
            sheet_id: 表格ID
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            # 模拟字段数据
            mock_fields = [
                {
                    'field_id': 'field_001',
                    'name': '任务名称',
                    'type': 'text',
                    'is_primary': True
                },
                {
                    'field_id': 'field_002',
                    'name': '负责人',
                    'type': 'user',
                    'multiple': False
                },
                {
                    'field_id': 'field_003',
                    'name': '状态',
                    'type': 'single_select',
                    'options': ['待办', '进行中', '已完成']
                },
                {
                    'field_id': 'field_004',
                    'name': '优先级',
                    'type': 'single_select',
                    'options': ['高', '中', '低']
                },
                {
                    'field_id': 'field_005',
                    'name': '标签',
                    'type': 'multi_select',
                    'options': ['前端', '后端', '设计', '测试']
                },
                {
                    'field_id': 'field_006',
                    'name': '开始日期',
                    'type': 'date',
                    'format': 'YYYY-MM-DD'
                },
                {
                    'field_id': 'field_007',
                    'name': '进度',
                    'type': 'progress',
                    'min': 0,
                    'max': 100
                },
                {
                    'field_id': 'field_008',
                    'name': '关联项目',
                    'type': 'link',
                    'target_sheet': 'sheet_project'
                }
            ]
            
            return {
                'success': True,
                'sheet_id': sheet_id,
                'fields': mock_fields,
                'count': len(mock_fields),
                'note': '当前为模拟数据，配置 WPS 开放平台凭证后可获取真实数据'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def add_field(self, sheet_id: str, name: str, field_type: str, config: dict = None):
        """添加字段到多维表格
        
        Args:
            sheet_id: 表格ID
            name: 字段名称
            field_type: 字段类型 (text/number/date/user/single_select/multi_select/progress/link)
            config: 字段配置
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            field_types = {
                'text': '文本',
                'number': '数字',
                'date': '日期',
                'user': '人员',
                'single_select': '单选',
                'multi_select': '多选',
                'progress': '进度',
                'link': '关联',
                'attachment': '附件',
                'checkbox': '复选框'
            }
            
            return {
                'success': True,
                'message': f'字段添加成功: {name}',
                'sheet_id': sheet_id,
                'field_id': 'field_new_001',
                'field_type': field_types.get(field_type, '文本'),
                'note': '当前为模拟添加，配置 WPS 开放平台凭证后可真实添加'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def query_data(self, sheet_id: str, view_id: str = "", filters: dict = None, sort: list = None):
        """高级查询多维表格数据（支持筛选和排序）
        
        Args:
            sheet_id: 表格ID
            view_id: 视图ID，为空则使用默认视图
            filters: 筛选条件
            sort: 排序规则
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            # 模拟筛选后的数据
            mock_data = {
                'sheet_id': sheet_id,
                'view_id': view_id or 'default',
                'total': 2,
                'filters_applied': filters or {},
                'sort_applied': sort or [],
                'records': [
                    {
                        'record_id': 'rec_001',
                        'fields': {
                            '任务名称': '设计首页',
                            '负责人': '张三',
                            '状态': '进行中',
                            '优先级': '高',
                            '进度': 60
                        }
                    },
                    {
                        'record_id': 'rec_002',
                        'fields': {
                            '任务名称': '开发API',
                            '负责人': '李四',
                            '状态': '待办',
                            '优先级': '中',
                            '进度': 0
                        }
                    }
                ]
            }
            
            return {
                'success': True,
                'data': mock_data,
                'note': '当前为模拟数据，配置 WPS 开放平台凭证后可获取真实数据'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}


class WPSFlowController:
    """WPS 流程图控制器
    
    管理 WPS 365 流程图（Flowchart）
    需要 WPS 开放平台凭证
    """
    
    def __init__(self, app_id: str = "", app_secret: str = ""):
        self.app_id = app_id
        self.app_secret = app_secret
        
    def list_flows(self, folder_id: str = ""):
        """获取流程图列表"""
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            # 模拟数据
            mock_flows = [
                {
                    'flow_id': 'flow_001',
                    'name': '产品发布流程',
                    'type': 'flowchart',
                    'create_time': '2024-01-10',
                    'update_time': '2024-02-15',
                    'creator': '张三',
                    'node_count': 12,
                    'url': 'https://www.kdocs.cn/l/flow_001'
                },
                {
                    'flow_id': 'flow_002',
                    'name': '审批流程图',
                    'type': 'flowchart',
                    'create_time': '2024-01-20',
                    'update_time': '2024-02-10',
                    'creator': '李四',
                    'node_count': 8,
                    'url': 'https://www.kdocs.cn/l/flow_002'
                },
                {
                    'flow_id': 'flow_003',
                    'name': '系统架构图',
                    'type': 'flowchart',
                    'create_time': '2024-02-01',
                    'update_time': '2024-02-20',
                    'creator': '王五',
                    'node_count': 15,
                    'url': 'https://www.kdocs.cn/l/flow_003'
                }
            ]
            
            return {
                'success': True,
                'flows': mock_flows,
                'count': len(mock_flows),
                'note': '当前为模拟数据，配置 WPS 开放平台凭证后可获取真实数据'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def create_flow(self, name: str, template: str = "blank"):
        """创建流程图
        
        Args:
            name: 流程图名称
            template: 模板类型 (blank/process/approval/org)
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            templates = {
                'blank': '空白流程图',
                'process': '业务流程',
                'approval': '审批流程',
                'org': '组织架构'
            }
            
            return {
                'success': True,
                'message': f'流程图创建成功: {name}',
                'flow_id': 'flow_new_001',
                'template': templates.get(template, '空白流程图'),
                'url': f'https://www.kdocs.cn/l/flow_new_001',
                'note': '当前为模拟创建，配置 WPS 开放平台凭证后可真实创建'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def export_flow(self, flow_id: str, format: str = "png"):
        """导出流程图
        
        Args:
            flow_id: 流程图ID
            format: 导出格式 (png/pdf/svg)
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            return {
                'success': True,
                'message': f'流程图导出成功',
                'flow_id': flow_id,
                'format': format,
                'download_url': f'https://www.kdocs.cn/l/{flow_id}/export.{format}',
                'note': '当前为模拟导出，配置 WPS 开放平台凭证后可真实导出'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}


class WPSMindController:
    """WPS 思维导图控制器
    
    管理 WPS 365 思维导图（Mind Map）
    需要 WPS 开放平台凭证
    """
    
    def __init__(self, app_id: str = "", app_secret: str = ""):
        self.app_id = app_id
        self.app_secret = app_secret
        
    def list_minds(self, folder_id: str = ""):
        """获取思维导图列表"""
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            # 模拟数据
            mock_minds = [
                {
                    'mind_id': 'mind_001',
                    'name': '产品规划思维导图',
                    'type': 'mindmap',
                    'create_time': '2024-01-15',
                    'update_time': '2024-02-18',
                    'creator': '张三',
                    'node_count': 25,
                    'url': 'https://www.kdocs.cn/l/mind_001'
                },
                {
                    'mind_id': 'mind_002',
                    'name': '项目需求分析',
                    'type': 'mindmap',
                    'create_time': '2024-01-25',
                    'update_time': '2024-02-20',
                    'creator': '李四',
                    'node_count': 18,
                    'url': 'https://www.kdocs.cn/l/mind_002'
                },
                {
                    'mind_id': 'mind_003',
                    'name': '团队协作框架',
                    'type': 'mindmap',
                    'create_time': '2024-02-05',
                    'update_time': '2024-02-22',
                    'creator': '王五',
                    'node_count': 32,
                    'url': 'https://www.kdocs.cn/l/mind_003'
                }
            ]
            
            return {
                'success': True,
                'minds': mock_minds,
                'count': len(mock_minds),
                'note': '当前为模拟数据，配置 WPS 开放平台凭证后可获取真实数据'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def create_mind(self, name: str, template: str = "blank"):
        """创建思维导图
        
        Args:
            name: 思维导图名称
            template: 模板类型 (blank/logic/brainstorm/project)
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            templates = {
                'blank': '空白思维导图',
                'logic': '逻辑分析',
                'brainstorm': '头脑风暴',
                'project': '项目规划'
            }
            
            return {
                'success': True,
                'message': f'思维导图创建成功: {name}',
                'mind_id': 'mind_new_001',
                'template': templates.get(template, '空白思维导图'),
                'url': f'https://www.kdocs.cn/l/mind_new_001',
                'note': '当前为模拟创建，配置 WPS 开放平台凭证后可真实创建'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def export_mind(self, mind_id: str, format: str = "png"):
        """导出思维导图
        
        Args:
            mind_id: 思维导图ID
            format: 导出格式 (png/pdf/svg/xmind)
        """
        try:
            if not self.app_id:
                return {
                    'success': False,
                    'error': '未配置 WPS 开放平台 App ID'
                }
            
            return {
                'success': True,
                'message': f'思维导图导出成功',
                'mind_id': mind_id,
                'format': format,
                'download_url': f'https://www.kdocs.cn/l/{mind_id}/export.{format}',
                'note': '当前为模拟导出，配置 WPS 开放平台凭证后可真实导出'
            }
            
        except Exception as e:
            return {'success': False, 'error': str(e)}


def main():
    """主入口函数"""
    if len(sys.argv) < 2:
        print(json.dumps({
            'success': False,
            'error': '缺少操作参数'
        }))
        sys.exit(1)
    
    action = sys.argv[1]
    
    # 读取配置
    config_path = os.path.join(os.path.dirname(__file__), '..', 'config.json')
    try:
        with open(config_path, 'r') as f:
            config = json.load(f)
    except:
        config = {}
    
    default_save_path = config.get('default_save_path', '~/Documents/WPS')
    app_id = config.get('app_id', '')
    app_secret = config.get('app_secret', '')
    
    wps = WPSController(default_save_path)
    form_controller = WPSFormController(app_id, app_secret)
    doc_controller = WPSDocController(app_id, app_secret)
    sheet_controller = WPSSheetController(app_id, app_secret)
    flow_controller = WPSFlowController(app_id, app_secret)
    mind_controller = WPSMindController(app_id, app_secret)
    
    # 解析参数
    params = {}
    for arg in sys.argv[2:]:
        if '=' in arg:
            key, value = arg.split('=', 1)
            params[key] = value
    
    # 执行操作
    if action == 'create':
        result = wps.create_document(
            doc_type=params.get('type', 'writer'),
            filename=params.get('filename', '新建文档.docx'),
            content=params.get('content', '')
        )
    elif action == 'open':
        result = wps.open_document(
            filepath=params.get('file', '')
        )
    elif action == 'list':
        result = wps.list_documents(
            directory=params.get('dir')
        )
    elif action == 'convert':
        result = wps.convert_format(
            input_file=params.get('file', ''),
            output_format=params.get('format', 'pdf')
        )
    elif action == 'batch_convert':
        result = wps.batch_convert(
            directory=params.get('dir', default_save_path),
            target_format=params.get('format', 'pdf')
        )
    # 智能表单操作
    elif action == 'form_list':
        result = form_controller.list_forms(
            folder_id=params.get('folder', '')
        )
    elif action == 'form_data':
        result = form_controller.get_form_data(
            form_id=params.get('form_id', ''),
            page=int(params.get('page', 1)),
            page_size=int(params.get('page_size', 100))
        )
    elif action == 'form_submit':
        # 解析 JSON 数据
        try:
            data = json.loads(params.get('data', '{}'))
        except:
            data = {}
        result = form_controller.submit_form(
            form_id=params.get('form_id', ''),
            data=data
        )
    elif action == 'form_create':
        try:
            fields = json.loads(params.get('fields', '[]'))
        except:
            fields = []
        result = form_controller.create_form(
            name=params.get('name', '新建表单'),
            fields=fields
        )
    # 智能文档操作
    elif action == 'doc_list':
        result = doc_controller.list_docs(
            folder_id=params.get('folder', '')
        )
    elif action == 'doc_create':
        result = doc_controller.create_doc(
            name=params.get('name', '新建文档'),
            template=params.get('template', 'blank')
        )
    elif action == 'doc_content':
        result = doc_controller.get_doc_content(
            doc_id=params.get('doc_id', '')
        )
    elif action == 'doc_share':
        result = doc_controller.share_doc(
            doc_id=params.get('doc_id', ''),
            permission=params.get('permission', 'read')
        )
    # 智能表格操作
    elif action == 'sheet_list':
        result = sheet_controller.list_sheets(
            folder_id=params.get('folder', '')
        )
    elif action == 'sheet_create':
        result = sheet_controller.create_sheet(
            name=params.get('name', '新建表格'),
            template=params.get('template', 'blank')
        )
    elif action == 'sheet_data':
        result = sheet_controller.get_sheet_data(
            sheet_id=params.get('sheet_id', ''),
            range=params.get('range', '')
        )
    elif action == 'sheet_update':
        try:
            values = json.loads(params.get('values', '[]'))
        except:
            values = []
        result = sheet_controller.update_sheet_data(
            sheet_id=params.get('sheet_id', ''),
            range=params.get('range', ''),
            values=values
        )
    # 多维表格特有操作
    elif action == 'sheet_views':
        result = sheet_controller.list_views(
            sheet_id=params.get('sheet_id', '')
        )
    elif action == 'sheet_view_create':
        try:
            config = json.loads(params.get('config', '{}'))
        except:
            config = {}
        result = sheet_controller.create_view(
            sheet_id=params.get('sheet_id', ''),
            name=params.get('name', '新建视图'),
            view_type=params.get('view_type', 'grid'),
            config=config
        )
    elif action == 'sheet_fields':
        result = sheet_controller.get_fields(
            sheet_id=params.get('sheet_id', '')
        )
    elif action == 'sheet_field_add':
        try:
            config = json.loads(params.get('config', '{}'))
        except:
            config = {}
        result = sheet_controller.add_field(
            sheet_id=params.get('sheet_id', ''),
            name=params.get('name', ''),
            field_type=params.get('field_type', 'text'),
            config=config
        )
    elif action == 'sheet_query':
        try:
            filters = json.loads(params.get('filters', '{}'))
            sort = json.loads(params.get('sort', '[]'))
        except:
            filters = {}
            sort = []
        result = sheet_controller.query_data(
            sheet_id=params.get('sheet_id', ''),
            view_id=params.get('view_id', ''),
            filters=filters,
            sort=sort
        )
    # 流程图操作
    elif action == 'flow_list':
        result = flow_controller.list_flows(
            folder_id=params.get('folder', '')
        )
    elif action == 'flow_create':
        result = flow_controller.create_flow(
            name=params.get('name', '新建流程图'),
            template=params.get('template', 'blank')
        )
    elif action == 'flow_export':
        result = flow_controller.export_flow(
            flow_id=params.get('flow_id', ''),
            format=params.get('format', 'png')
        )
    # 思维导图操作
    elif action == 'mind_list':
        result = mind_controller.list_minds(
            folder_id=params.get('folder', '')
        )
    elif action == 'mind_create':
        result = mind_controller.create_mind(
            name=params.get('name', '新建思维导图'),
            template=params.get('template', 'blank')
        )
    elif action == 'mind_export':
        result = mind_controller.export_mind(
            mind_id=params.get('mind_id', ''),
            format=params.get('format', 'png')
        )
    
    # ========== Markdown 转换功能（新增）==========
    elif action == 'md_to_docx':
        """Markdown 转 Word 文档（完整版）"""
        try:
            from md_converter import MarkdownToWordConverter
            
            input_file = params.get('file', '')
            output_file = params.get('output', '')
            title = params.get('title', None)
            
            if not input_file:
                result = {'success': False, 'error': '请指定输入 Markdown 文件'}
            elif not os.path.exists(input_file):
                result = {'success': False, 'error': f'文件不存在: {input_file}'}
            else:
                converter = MarkdownToWordConverter()
                output_path = converter.convert_file(input_file, output_file, title)
                result = {
                    'success': True,
                    'message': f'Markdown 转换成功',
                    'input': input_file,
                    'output': output_path
                }
        except Exception as e:
            result = {'success': False, 'error': f'转换失败: {str(e)}'}
    
    elif action == 'docx_to_md':
        """Word 转 Markdown"""
        try:
            from md_converter import WordToMarkdownConverter
            
            input_file = params.get('file', '')
            output_file = params.get('output', '')
            
            if not input_file:
                result = {'success': False, 'error': '请指定输入 Word 文件'}
            elif not os.path.exists(input_file):
                result = {'success': False, 'error': f'文件不存在: {input_file}'}
            else:
                converter = WordToMarkdownConverter()
                md_content = converter.convert(input_file, output_file)
                result = {
                    'success': True,
                    'message': f'Word 转换成功',
                    'input': input_file,
                    'output': output_file,
                    'preview': md_content[:500] + '...' if len(md_content) > 500 else md_content
                }
        except Exception as e:
            result = {'success': False, 'error': f'转换失败: {str(e)}'}
    
    # ========== 图片处理功能（新增）==========
    elif action == 'insert_image':
        """向 Word 文档插入图片"""
        try:
            from docx import Document
            from image_handler import ImageHandler
            from docx.shared import Inches
            
            docx_file = params.get('docx', '')
            image_file = params.get('image', '')
            output_file = params.get('output', docx_file)
            width = float(params.get('width', 4))
            caption = params.get('caption', None)
            align = params.get('align', 'center')
            
            if not docx_file or not os.path.exists(docx_file):
                result = {'success': False, 'error': '请指定有效的 Word 文件'}
            elif not image_file or not os.path.exists(image_file):
                result = {'success': False, 'error': f'图片不存在: {image_file}'}
            else:
                doc = Document(docx_file)
                handler = ImageHandler(doc)
                handler.insert_image(
                    image_file,
                    width=Inches(width),
                    align=align,
                    caption=caption
                )
                doc.save(output_file)
                result = {
                    'success': True,
                    'message': f'图片插入成功',
                    'document': output_file,
                    'image': image_file
                }
        except Exception as e:
            result = {'success': False, 'error': f'插入图片失败: {str(e)}'}
    
    elif action == 'insert_image_grid':
        """向 Word 文档插入图片网格"""
        try:
            from docx import Document
            from image_handler import ImageHandler
            from docx.shared import Inches
            
            docx_file = params.get('docx', '')
            image_files = params.get('images', '').split(',')
            cols = int(params.get('cols', 2))
            width = float(params.get('width', 3))
            
            # 验证图片
            valid_images = [img.strip() for img in image_files if os.path.exists(img.strip())]
            
            if not valid_images:
                result = {'success': False, 'error': '没有有效的图片文件'}
            else:
                if docx_file and os.path.exists(docx_file):
                    doc = Document(docx_file)
                else:
                    doc = Document()
                    docx_file = params.get('output', 'image_grid.docx')
                
                handler = ImageHandler(doc)
                handler.insert_images_grid(valid_images, cols=cols, width=Inches(width))
                doc.save(docx_file)
                
                result = {
                    'success': True,
                    'message': f'图片网格插入成功',
                    'document': docx_file,
                    'images_count': len(valid_images),
                    'cols': cols
                }
        except Exception as e:
            result = {'success': False, 'error': f'插入图片网格失败: {str(e)}'}
    
    elif action == 'create_text_image_layout':
        """创建图文混排布局"""
        try:
            from docx import Document
            from image_handler import ImageHandler
            from docx.shared import Inches
            
            docx_file = params.get('output', 'layout.docx')
            text = params.get('text', '')
            image_file = params.get('image', '')
            layout = params.get('layout', 'left')  # left/right/top/bottom
            image_width = float(params.get('image_width', 2.5))
            
            if not image_file or not os.path.exists(image_file):
                result = {'success': False, 'error': f'图片不存在: {image_file}'}
            else:
                doc = Document()
                handler = ImageHandler(doc)
                handler.create_text_image_layout(
                    text=text,
                    image_path=image_file,
                    layout=layout,
                    image_width=Inches(image_width)
                )
                doc.save(docx_file)
                
                result = {
                    'success': True,
                    'message': f'图文混排创建成功',
                    'document': docx_file,
                    'layout': layout
                }
        except Exception as e:
            result = {'success': False, 'error': f'创建图文混排失败: {str(e)}'}
    
    elif action == 'md_with_images_to_docx':
        """Markdown 转 Word（包含图片处理）"""
        try:
            from md_converter import MarkdownToWordConverter
            from image_handler import MarkdownImageProcessor
            from docx import Document
            from docx.shared import Inches
            import re
            
            input_file = params.get('file', '')
            output_file = params.get('output', '')
            title = params.get('title', None)
            image_width = float(params.get('image_width', 4))
            
            if not input_file or not os.path.exists(input_file):
                result = {'success': False, 'error': f'文件不存在: {input_file}'}
            else:
                # 读取 Markdown 内容
                with open(input_file, 'r', encoding='utf-8') as f:
                    md_content = f.read()
                
                # 确定输出路径
                if not output_file:
                    base_name = os.path.splitext(input_file)[0]
                    output_file = base_name + '.docx'
                
                # 创建文档
                doc = Document()
                
                # 设置默认字体
                from docx.oxml.ns import qn
                style = doc.styles['Normal']
                style.font.name = '宋体'
                style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                # 添加标题
                if title:
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    heading = doc.add_heading(title, level=0)
                    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in heading.runs:
                        run.font.name = '宋体'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                # 获取 Markdown 文件所在目录
                base_path = os.path.dirname(os.path.abspath(input_file))
                
                # 处理 Markdown 中的图片
                image_processor = MarkdownImageProcessor(doc, base_path)
                
                # 提取并处理图片
                images = image_processor.extract_images_from_markdown(md_content)
                
                # 转换 Markdown（去除图片语法）
                md_without_images = re.sub(r'!\[([^\]]*)\]\(([^)"\s]+)(?:\s+"([^"]*)")?\)\s*\n?', '', md_content)
                
                # 转换纯文本部分
                converter = MarkdownToWordConverter()
                converter.doc = doc
                converter._parse_markdown(md_without_images)
                
                # 在适当位置插入图片
                if images:
                    for img_info in images:
                        img_path = img_info['path']
                        if not os.path.isabs(img_path):
                            img_path = os.path.join(base_path, img_path)
                        
                        if os.path.exists(img_path):
                            caption = img_info['title'] if img_info['title'] else img_info['alt']
                            image_processor.image_handler.insert_image(
                                img_path,
                                width=Inches(image_width),
                                align='center',
                                caption=caption if caption else None
                            )
                
                doc.save(output_file)
                
                result = {
                    'success': True,
                    'message': f'Markdown 转换成功（含图片）',
                    'input': input_file,
                    'output': output_file,
                    'images_found': len(images),
                    'images_processed': len([img for img in images if os.path.exists(os.path.join(base_path, img['path']) if not os.path.isabs(img['path']) else img['path'])])
                }
        except Exception as e:
            result = {'success': False, 'error': f'转换失败: {str(e)}'}
    
    # ========== Excel 转换功能（新增）==========
    elif action == 'md_to_xlsx':
        """Markdown 转 Excel"""
        try:
            from excel_converter import MarkdownToExcelConverter
            
            input_file = params.get('file', '')
            output_file = params.get('output', '')
            title = params.get('title', None)
            
            if not input_file:
                result = {'success': False, 'error': '请指定输入 Markdown 文件'}
            elif not os.path.exists(input_file):
                result = {'success': False, 'error': f'文件不存在: {input_file}'}
            else:
                converter = MarkdownToExcelConverter()
                output_path = converter.convert_file(input_file, output_file, title)
                result = {
                    'success': True,
                    'message': f'Markdown 转 Excel 成功',
                    'input': input_file,
                    'output': output_path
                }
        except Exception as e:
            result = {'success': False, 'error': f'转换失败: {str(e)}'}
    
    elif action == 'xlsx_to_md':
        """Excel 转 Markdown"""
        try:
            from excel_converter import ExcelToMarkdownConverter
            
            input_file = params.get('file', '')
            output_file = params.get('output', '')
            
            if not input_file:
                result = {'success': False, 'error': '请指定输入 Excel 文件'}
            elif not os.path.exists(input_file):
                result = {'success': False, 'error': f'文件不存在: {input_file}'}
            else:
                converter = ExcelToMarkdownConverter()
                md_content = converter.convert(input_file, output_file)
                result = {
                    'success': True,
                    'message': f'Excel 转 Markdown 成功',
                    'input': input_file,
                    'output': output_file,
                    'preview': md_content[:500] + '...' if len(md_content) > 500 else md_content
                }
        except Exception as e:
            result = {'success': False, 'error': f'转换失败: {str(e)}'}
    
    # ========== PPT 转换功能（新增）==========
    elif action == 'md_to_pptx':
        """Markdown 转 PPT"""
        try:
            from ppt_converter import MarkdownToPPTConverter
            
            input_file = params.get('file', '')
            output_file = params.get('output', '')
            title = params.get('title', None)
            
            if not input_file:
                result = {'success': False, 'error': '请指定输入 Markdown 文件'}
            elif not os.path.exists(input_file):
                result = {'success': False, 'error': f'文件不存在: {input_file}'}
            else:
                converter = MarkdownToPPTConverter()
                output_path = converter.convert_file(input_file, output_file, title)
                result = {
                    'success': True,
                    'message': f'Markdown 转 PPT 成功',
                    'input': input_file,
                    'output': output_path
                }
        except Exception as e:
            result = {'success': False, 'error': f'转换失败: {str(e)}'}
    
    elif action == 'pptx_to_md':
        """PPT 转 Markdown"""
        try:
            from ppt_converter import PPTToMarkdownConverter
            
            input_file = params.get('file', '')
            output_file = params.get('output', '')
            
            if not input_file:
                result = {'success': False, 'error': '请指定输入 PPT 文件'}
            elif not os.path.exists(input_file):
                result = {'success': False, 'error': f'文件不存在: {input_file}'}
            else:
                converter = PPTToMarkdownConverter()
                md_content = converter.convert(input_file, output_file)
                result = {
                    'success': True,
                    'message': f'PPT 转 Markdown 成功',
                    'input': input_file,
                    'output': output_file,
                    'preview': md_content[:500] + '...' if len(md_content) > 500 else md_content
                }
        except Exception as e:
            result = {'success': False, 'error': f'转换失败: {str(e)}'}
    
    # ========== PPT 图片处理功能（新增）==========
    elif action == 'insert_image_to_ppt':
        """向 PPT 插入图片"""
        try:
            from pptx import Presentation
            from ppt_converter import PPTImageHandler
            
            pptx_file = params.get('pptx', '')
            slide_index = int(params.get('slide', 0))
            image_file = params.get('image', '')
            output_file = params.get('output', pptx_file)
            left = float(params.get('left', 1))
            top = float(params.get('top', 1))
            width = float(params.get('width', 4))
            height = params.get('height', None)
            if height:
                height = float(height)
            
            if not pptx_file or not os.path.exists(pptx_file):
                result = {'success': False, 'error': '请指定有效的 PPT 文件'}
            elif not image_file or not os.path.exists(image_file):
                result = {'success': False, 'error': f'图片不存在: {image_file}'}
            else:
                prs = Presentation(pptx_file)
                handler = PPTImageHandler(prs)
                handler.insert_image_to_slide(slide_index, image_file, left, top, width, height)
                prs.save(output_file)
                result = {
                    'success': True,
                    'message': f'图片插入成功',
                    'slide': slide_index,
                    'output': output_file
                }
        except Exception as e:
            result = {'success': False, 'error': f'插入图片失败: {str(e)}'}
    
    elif action == 'create_ppt_text_image_layout':
        """创建 PPT 图文混排"""
        try:
            from pptx import Presentation
            from ppt_converter import PPTImageHandler
            
            pptx_file = params.get('pptx', '')
            slide_index = int(params.get('slide', 0))
            text = params.get('text', '')
            image_file = params.get('image', '')
            layout = params.get('layout', 'left')
            image_width = float(params.get('image_width', 4))
            output_file = params.get('output', pptx_file)
            
            if not pptx_file or not os.path.exists(pptx_file):
                result = {'success': False, 'error': '请指定有效的 PPT 文件'}
            elif not image_file or not os.path.exists(image_file):
                result = {'success': False, 'error': f'图片不存在: {image_file}'}
            else:
                prs = Presentation(pptx_file)
                handler = PPTImageHandler(prs)
                handler.create_text_image_layout(slide_index, text, image_file, layout, image_width)
                prs.save(output_file)
                result = {
                    'success': True,
                    'message': f'图文混排创建成功',
                    'slide': slide_index,
                    'layout': layout,
                    'output': output_file
                }
        except Exception as e:
            result = {'success': False, 'error': f'创建图文混排失败: {str(e)}'}
    
    # ========== Excel 图片处理功能（新增）==========
    elif action == 'insert_image_to_excel':
        """向 Excel 插入图片"""
        try:
            from openpyxl import load_workbook
            from excel_converter import ExcelImageHandler
            
            xlsx_file = params.get('xlsx', '')
            sheet_name = params.get('sheet', 'Sheet1')
            cell = params.get('cell', 'A1')
            image_file = params.get('image', '')
            output_file = params.get('output', xlsx_file)
            width = params.get('width', None)
            height = params.get('height', None)
            if width:
                width = float(width)
            if height:
                height = float(height)
            
            if not xlsx_file or not os.path.exists(xlsx_file):
                result = {'success': False, 'error': '请指定有效的 Excel 文件'}
            elif not image_file or not os.path.exists(image_file):
                result = {'success': False, 'error': f'图片不存在: {image_file}'}
            else:
                wb = load_workbook(xlsx_file)
                handler = ExcelImageHandler(wb)
                handler.insert_image_to_cell(sheet_name, cell, image_file, width, height)
                wb.save(output_file)
                result = {
                    'success': True,
                    'message': f'图片插入成功',
                    'sheet': sheet_name,
                    'cell': cell,
                    'output': output_file
                }
        except Exception as e:
            result = {'success': False, 'error': f'插入图片失败: {str(e)}'}
    
    elif action == 'create_excel_text_image_layout':
        """创建 Excel 图文混排"""
        try:
            from openpyxl import load_workbook
            from excel_converter import ExcelImageHandler
            
            xlsx_file = params.get('xlsx', '')
            sheet_name = params.get('sheet', 'Sheet1')
            cell = params.get('cell', 'A1')
            text = params.get('text', '')
            image_file = params.get('image', '')
            layout = params.get('layout', 'right')
            image_width = float(params.get('image_width', 100))
            output_file = params.get('output', xlsx_file)
            
            if not xlsx_file or not os.path.exists(xlsx_file):
                result = {'success': False, 'error': '请指定有效的 Excel 文件'}
            elif not image_file or not os.path.exists(image_file):
                result = {'success': False, 'error': f'图片不存在: {image_file}'}
            else:
                wb = load_workbook(xlsx_file)
                handler = ExcelImageHandler(wb)
                handler.create_text_image_layout(sheet_name, cell, text, image_file, layout, image_width)
                wb.save(output_file)
                result = {
                    'success': True,
                    'message': f'图文混排创建成功',
                    'sheet': sheet_name,
                    'cell': cell,
                    'layout': layout,
                    'output': output_file
                }
        except Exception as e:
            result = {'success': False, 'error': f'创建图文混排失败: {str(e)}'}
    
    else:
        result = {'success': False, 'error': f'未知操作: {action}'}
    
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == '__main__':
    main()
